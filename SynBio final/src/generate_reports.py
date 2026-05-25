import os
import sys
import subprocess
import pandas as pd
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
import shutil
import time

# Define directories
LATEX_DIR = "reports/latex"
WORD_DIR = "reports/word"
TABLES_DIR = "results/tables"
FIGURES_DIR = "reports/latex/shared/figures"

# Helper to load data
def load_table(name):
    path = os.path.join(TABLES_DIR, name)
    if os.path.exists(path):
        return pd.read_csv(path)
    return None

# Load results tables
and_sweep = load_table("and_gate_parameter_sweep.csv")
and_perf = load_table("and_gate_performance.csv")
cv_results = load_table("cross_validation_results.csv")
de_discovery = load_table("differential_expression_discovery.csv")
qc_summary = load_table("expression_qc_summary.csv")
ext_val = load_table("external_validation_final_pair.csv")
final_pair = load_table("final_candidate_pair.csv")
gene_scores = load_table("gene_pair_scores.csv")
model_perf = load_table("model_performance_summary.csv")
random_pair = load_table("random_pair_control.csv")
sensitivity = load_table("threshold_sensitivity.csv")
shap_imp = load_table("shap_feature_importance.csv")
shap_thresh = load_table("shap_threshold_candidates.csv")
top_candidates = load_table("top_tumor_high_candidates.csv")

# ----------------- TABLE PREPARATIONS -----------------
# Table 1: Dataset Summary
t1_headers = ["Dataset / Cohort", "Platform / Type", "Tumor Samples", "Normal Samples", "Total Samples"]
t1_data = [
    ["TCGA-PAAD", "RNA-seq (RSEM TPM)", "178", "0", "178"],
    ["GTEx Pancreas", "RNA-seq (RSEM TPM)", "0", "167", "167"],
    ["Discovery Cohort (Combined)", "RNA-seq (RSEM TPM)", "178", "167", "345"],
    ["GSE62452 (External Validation)", "Microarray (Affymetrix GPL6244)", "69", "61", "130"]
]

# Table 2: Top 10 DE Genes (sorted by specificity score)
t2_headers = ["Gene", "Mean Tumor (log2 TPM)", "Mean Normal (log2 TPM)", "log2FC", "FDR", "ROC-AUC", "Specificity Score"]
t2_data = []
if top_candidates is not None:
    df_sorted = top_candidates.sort_values(by=["specificity_score", "mean_pdac"], ascending=[False, False])
    for _, row in df_sorted.head(10).iterrows():
        gene = str(row["gene"])
        rec = de_discovery[de_discovery["gene"] == gene] if de_discovery is not None else []
        if len(rec) > 0:
            mean_p = f"{rec.iloc[0]['mean_pdac']:.3f}"
            mean_n = f"{rec.iloc[0]['mean_normal']:.3f}"
            log2fc = f"{rec.iloc[0]['log2fc']:.3f}"
            fdr = f"{rec.iloc[0]['fdr']:.2e}"
            auc = f"{rec.iloc[0]['auc']:.3f}"
            spec = f"{row['specificity_score']:.3f}"
            t2_data.append([gene, mean_p, mean_n, log2fc, fdr, auc, spec])

# Table 3: Model Performance
t3_headers = ["Classifier Model", "Mean 5-Fold CV AUC", "Test AUC", "Test Accuracy", "Sensitivity", "Specificity", "F1 Score"]
t3_data = []
if cv_results is not None and model_perf is not None:
    for _, row in model_perf.iterrows():
        model_name = str(row["Model"])
        cv_rec = cv_results[cv_results["Model"] == model_name]
        cv_auc = f"{cv_rec.iloc[0]['Mean_CV_AUC']:.3f}" if len(cv_rec) > 0 else "N/A"
        test_auc = f"{row['Test_AUC']:.3f}"
        test_acc = f"{row['Test_Accuracy']*100:.1f}%"
        sens = f"{row['Test_Recall']:.3f}"
        tn, fp = row['TN'], row['FP']
        spec = f"{tn / (tn + fp):.3f}" if (tn + fp) > 0 else "N/A"
        f1 = f"{row['Test_F1']:.3f}"
        t3_data.append([model_name.replace("_", " "), cv_auc, test_auc, test_acc, sens, spec, f1])

# Table 4: Top 10 SHAP Features
t4_headers = ["Rank", "Gene Symbol", "Mean Absolute SHAP Value", "Inferred Threshold (log2 TPM)", "95% Confidence Interval"]
t4_data = []
if shap_imp is not None and shap_thresh is not None:
    for rank, row in enumerate(shap_imp.head(10).iterrows(), 1):
        gene = str(row[1]["gene"])
        shap_val = f"{row[1]['mean_abs_shap']:.4f}"
        thresh_rec = shap_thresh[shap_thresh["gene"] == gene]
        if len(thresh_rec) > 0:
            thresh = f"{thresh_rec.iloc[0]['inferred_threshold']:.3f}"
            ci = f"[{thresh_rec.iloc[0]['ci_95_lower']:.3f}, {thresh_rec.iloc[0]['ci_95_upper']:.3f}]"
        else:
            thresh, ci = "N/A", "N/A"
        t4_data.append([str(rank), gene, shap_val, thresh, ci])

# Table 5: Selected Candidate Pair
t5_headers = ["Property", "Gene A (UBE2S)", "Gene B (CCR6)", "Composite AND Gate"]
t5_data = []
if final_pair is not None and and_perf is not None:
    r = final_pair.iloc[0]
    ap = and_perf.iloc[0]
    t5_data = [
        ["Individual ROC-AUC", f"{r['gene_A_auc']:.4f}", f"{r['gene_B_auc']:.4f}", f"{ap['ROC_AUC']:.4f}"],
        ["log2 Fold Change (log2FC)", f"{r['gene_A_log2fc']:.3f}", f"{r['gene_B_log2fc']:.3f}", "—"],
        ["Pairwise Spearman Correlation (r_s)", f"{r['correlation']:.4f}", "—", "—"],
        ["Inferred Activation Threshold (K)", f"{ap['K_A']:.4f}", f"{ap['K_B']:.4f}", "—"],
        ["Tumor AND Activation Rate", f"{r['tumor_AND_activation']*100:.1f}%", "—", f"{r['tumor_AND_activation']*100:.1f}%"],
        ["Normal AND Activation Rate", f"{r['normal_AND_activation']*100:.1f}%", "—", f"{r['normal_AND_activation']*100:.1f}%"]
    ]

# Table 6: AND Gate Performance Metrics (Discovery Cohort)
t6_headers = ["Metric", "Formula / Definition", "Value"]
t6_data = []
if and_perf is not None:
    ap = and_perf.iloc[0]
    t6_data = [
        ["ROC-AUC", "Area Under the ROC Curve", f"{ap['ROC_AUC']:.4f}"],
        ["Accuracy", "(TP + TN) / Total", f"{ap['Accuracy']*100:.2f}%"],
        ["Sensitivity", "TP / (TP + FN)", f"{ap['Sensitivity']*100:.2f}%"],
        ["Specificity", "TN / (TN + FP)", f"{ap['Specificity']*100:.2f}%"],
        ["Decision Threshold", "Hill output threshold for ON state", f"{ap['Optimal_Threshold']:.2f}"]
    ]

# Table 7: External Validation Results (GSE62452)
t7_headers = ["Metric", "Value in Discovery (TCGA+GTEx)", "Value in Validation (GSE62452)", "Interpretation"]
t7_data = []
if ext_val is not None and and_perf is not None:
    ap = and_perf.iloc[0]
    ev = ext_val.iloc[0]
    t7_data = [
        ["Sample Size (N)", "345", f"{ev['sample_size']}", "Microarray vs RNA-seq"],
        ["ROC-AUC", f"{ap['ROC_AUC']:.4f}", f"{ev['ROC_AUC']:.4f}", "Reduced discrimination power"],
        ["Accuracy", f"{ap['Accuracy']*100:.1f}%", f"{ev['Accuracy']*100:.1f}%", "Impacted by platform shift"],
        ["Sensitivity", f"{ap['Sensitivity']*100:.1f}%", f"{ev['Sensitivity']*100:.1f}%", "Significant drop (platform scale mismatch)"],
        ["Specificity", f"{ap['Specificity']*100:.1f}%", f"{ev['Specificity']*100:.1f}%", "Highly conserved (retains safety profile)"]
    ]

# Table 8: Threshold Sensitivity (Subset of Perturbations)
t8_headers = ["UBE2S Threshold (K_A)", "CCR6 Threshold (K_B)", "ROC-AUC", "Accuracy", "Sensitivity", "Specificity"]
t8_data = []
if sensitivity is not None:
    # Baseline
    baseline_rec = and_perf.iloc[0]
    t8_data.append([
        "0.760 (Baseline)", "0.464 (Baseline)",
        f"{baseline_rec['ROC_AUC']:.4f}", f"{baseline_rec['Accuracy']*100:.1f}%",
        f"{baseline_rec['Sensitivity']*100:.1f}%", f"{baseline_rec['Specificity']*100:.1f}%"
    ])
    for _, row in sensitivity.iterrows():
        p_A = row["Perturbation_A"]
        p_B = row["Perturbation_B"]
        if p_A in [-0.5, -0.1, 0.1, 0.5] and p_B in [-0.5, -0.1, 0.1, 0.5] and p_A == p_B:
            k_A_val = f"{0.760 * (1 + p_A):.3f} ({p_A*100:+.0f}%)"
            k_B_val = f"{0.464 * (1 + p_B):.3f} ({p_B*100:+.0f}%)"
            auc = f"{row['ROC_AUC']:.4f}"
            acc = f"{row['Accuracy']*100:.1f}%"
            t8_data.append([k_A_val, k_B_val, auc, acc, "N/A", "N/A"])

en_sections = {'Abstract': 'Pancreatic ductal adenocarcinoma (PDAC) remains a highly lethal malignancy with a 5-year survival rate '
             'below 12%, primarily due to late-stage diagnosis and a lack of specific tumor biomarkers. In this study, '
             'we present an unbiased, data-driven computational pipeline to identify optimal candidate gene pairs for '
             'a synthetic biology AND-gate biosensor designed to discriminate PDAC from normal tissue. We utilized '
             'transcriptomic data from the TCGA-PAAD (n=178 tumors) and GTEx Normal Pancreas (n=167 normal tissues) '
             'cohorts, analyzing a filtered set of 58,581 genes. Differential expression and machine learning '
             'classification using an L1-regularized Logistic Regression model achieved a 5-fold cross-validation AUC '
             'of 1.000, identifying the most predictive features. Explainable AI (SHAP) was applied to infer '
             'expression thresholds for the top-ranked genes. Candidate gene pairs were selected based on '
             'orthogonality and activation rates. The optimal pair, UBE2S and CCR6, demonstrates a pairwise Pearson '
             'correlation of 0.714. Mathematical modeling using the Hill equation optimized the AND-gate parameters, '
             'yielding a discovery cohort AUC of 0.9986, an accuracy of 98.6%, a sensitivity of 97.8%, and a '
             'specificity of 99.4% at an activation threshold of 0.25. Permutation testing against 1,000 random gene '
             'pairs confirmed statistical significance (p < 0.0001). External validation on the GSE62452 microarray '
             'dataset (n=130) showed moderate AUC (0.648) with high specificity (98.4%) but low sensitivity (4.3%), '
             'reflecting cross-platform transferability challenges. Our pipeline provides a robust framework for '
             'logic-gated synthetic sensor design, bridging the gap between bioinformatics and wet-lab implementation.',
 'Affiliation': '$^1$Department of Life Science, National Taiwan University\\\\$^2$Department of Biochemical Science and Technology, National Taiwan University',
 'Author': 'SHIH, Chen-Jung$^1$, SU, Te-Fang$^1$, LIAO, Xuan-You$^2$, LIN, Chia-I$^2$',
 'Conclusion': 'In conclusion, we have developed a data-driven computational framework for the selection of input gene '
               'pairs for a synthetic biology AND-gate biosensor in pancreatic cancer. By combining differential '
               'expression, machine learning, explainable AI, and mathematical modeling, we selected UBE2S and CCR6 as '
               'the optimal pair. This combination achieves excellent classification performance and robustness in '
               'silico, while capturing distinct biological hallmarks of PDAC (mitotic progression and '
               'microenvironmental inflammation). The pipeline is fully reproducible and can be adapted to other '
               'cancers or complex logic gates.',
 'DE': 'We performed differential expression (DE) analysis comparing the 178 PDAC tumor samples against the 167 normal '
       "pancreas samples. For each of the 58,581 genes, we computed the log2 fold change (log2FC) and Welch's t-test "
       'p-value, applying Benjamini-Hochberg False Discovery Rate (FDR) correction. Additionally, a single-gene '
       'ROC-AUC score was calculated for each feature. Tumor-high genes were defined using the thresholds: log2FC >= '
       '1.0 and FDR < 0.05, identifying 19,399 candidate genes. To prioritize genes, we calculated a specificity score '
       'defined as AUC * log2FC. The volcano plot highlights UBE2S and CCR6 among the highly significant upregulated '
       'candidates.',
 'DataSources': 'To establish a robust discovery cohort, we integrated transcriptomic data from two primary public '
                'sources: the Cancer Genome Atlas (TCGA-PAAD, representing primary pancreatic tumors, n=178) and the '
                'Genotype-Tissue Expression (GTEx, representing healthy normal pancreas tissue, n=167). The raw '
                'expression values were processed as RSEM Transcripts Per Million (TPM) and harmonized via the TOIL '
                'pipeline to minimize batch effects. For external validation, we retrieved the GSE62452 dataset from '
                'the Gene Expression Omnibus (GEO), containing 130 pancreatic tissue samples (69 tumor and 61 normal '
                'adjacent tissues) analyzed using the Affymetrix GPL6244 microarray platform. This dual-dataset design '
                'ensures that our candidate pairs are discovered using high-throughput sequencing data and validated '
                'on independent, cross-platform microarray data to assess robustness.',
 'Date': 'May 2026',
 'HillModeling': "The logic gate's response was modeled using a dual-input Hill equation:\n"
                 '\n'
                 '\\[\n'
                 'Output = P_{\\text{basal}} + V_{\\max} \\left( \\frac{[A]^n}{K_A^n + [A]^n} \\right) \\left( '
                 '\\frac{[B]^n}{K_B^n + [B]^n} \\right)\n'
                 '\\]\n'
                 '\n'
                 'where [A] and [B] are the min-max scaled expressions of UBE2S and CCR6, $K_A$ and $K_B$ are the '
                 'activation thresholds, $n$ is the Hill coefficient (steepness), and $P_{\\text{basal}}$ is the '
                 'leakiness. Parameter optimization via grid search determined the optimal settings: $n = 1$, '
                 '$P_{\\text{basal}} = 0.0$, $K_A = 0.760$ (for UBE2S), and $K_B = 0.464$ (for CCR6). The optimal '
                 'output decision threshold was set to 0.25 to maximize AUC and classification accuracy.',
 'InSilico': 'Simulating the optimized Hill equation AND gate on the 345-sample discovery cohort yielded outstanding '
             'classification performance. The AND-gate output achieved a ROC-AUC of 0.9986, an overall accuracy of '
             '98.55%, a sensitivity of 97.75%, and a specificity of 99.40%. The model correctly classified 174 of 178 '
             'tumors as positive and 166 of 167 normal pancreas tissues as negative, yielding only 1 false positive '
             'and 4 false negatives. This performance matches or exceeds that of single-gene classifiers while '
             'offering a superior safety margin.',
 'Introduction': 'Pancreatic ductal adenocarcinoma (PDAC) is characterized by a silent progression, aggressive '
                 'metastatic potential, and a dense desmoplastic tumor microenvironment (TME) that acts as a physical '
                 'and immunological barrier. Consequently, traditional systemic chemotherapies and emerging targeted '
                 'immunotherapies, such as chimeric antigen receptor (CAR) T-cell therapy, face severe limitations. '
                 'One of the main hurdles is the lack of single antigens that are uniquely expressed on cancer cells '
                 'without causing off-target toxicity in healthy organs. Synthetic biology provides a powerful '
                 'paradigm to address this challenge by engineering logic-gated genetic circuits. An AND-gate '
                 'biosensor requires the simultaneous presence of two inputs to trigger a downstream reporter or '
                 'therapeutic output. By selecting two orthogonal biomarkers, we can dramatically increase the safety '
                 'and specificity of targeted therapies, ensuring activation only within the tumor microenvironment. '
                 'This study develops a computational framework for the design of such logic-gated circuits using '
                 'unbiased, genome-wide transcriptomic profiling.',
 'Limitations': 'Several critical limitations must be acknowledged. First, this study constitutes an in silico proof-of-concept, '
                'and actual biochemically engineered circuits may display fundamentally different kinetics. Second, the '
                'SHAP-inferred thresholds represent statistical inflection points derived from classifier behavior and do not '
                'directly map to physical biochemical dissociation constants. Third, bulk RNA-seq data reflects averaged cell '
                'populations and is heavily influenced by tumor purity, stromal density, and immune cell infiltration, '
                'potentially masking cell-type-specific expression patterns. Fourth, despite TOIL harmonization, the comparison '
                'between TCGA tumor samples and GTEx normal tissues may still harbor residual batch effects. Fifth, the external '
                'validation cohort demonstrated extremely low sensitivity (4.3%), highlighting a significant challenge in '
                'cross-platform threshold transfer from RNA-seq to microarray data. Sixth, the selected candidate pair '
                '(UBE2S + CCR6) is not strictly statistically orthogonal, exhibiting a Spearman correlation of 0.714 in bulk '
                'data. Seventh, transcriptomic abundance does not guarantee equivalent sensor accessibility or protein-level '
                'expression. Eighth, translating these candidates into a functional synthetic circuit requires promoter '
                'engineering or RNA-based sensor design, each introducing additional layers of design complexity. Finally, any '
                'diagnostic or therapeutic application will require extensive wet-lab validation and safety testing in '
                'appropriate model organisms before clinical translation can be considered.',
 'ML': 'To identify the most predictive and sparse gene signature, we trained three classifiers on a stratified 80/20 '
       'train-test split of the discovery cohort: L1-regularized Logistic Regression (Lasso), Random Forest (100 '
       'estimators), and XGBoost. The L1 Logistic Regression classifier achieved perfect classification performance '
       'with a test AUC of 1.000 and accuracy of 100.0%, matching the 5-fold cross-validation AUC (1.000 +- 0.000). '
       'The Random Forest model achieved a test AUC of 0.9983 and XGBoost achieved 1.000. Due to its mathematical '
       'simplicity, sparseness (L1 penalty drives coefficients to zero), and perfect AUC, the L1 Logistic Regression '
       'model was selected for explainable AI analysis.',
 'Orthogonality': 'A key design requirement for AND-gate sensors is that the two inputs must be orthogonal. UBE2S and '
                  'CCR6 have a Pearson correlation of r = 0.714, which is higher than the ideal threshold of |r| <= '
                  '0.4. Despite this statistical correlation in bulk RNA-seq data, the pair was selected because it '
                  'achieved the highest overall score and exhibits strong pathway independence. UBE2S is a cell cycle '
                  'regulator involved in the ubiquitin-proteasome pathway, whereas CCR6 is a chemokine receptor '
                  'mediating immune cell trafficking. Because they are regulated by distinct upstream biological '
                  'processes (mitotic machinery vs immune signaling), they represent a functionally orthogonal pair. '
                  'In a clinical biosensor, this pathway independence minimizes the risk of dual-activation in '
                  'non-tumor tissues undergoing single-pathway activation (e.g., local inflammation).',
 'PairSelection': 'To select the optimal input pair from the top 100 SHAP features, we computed a composite Pair Score '
                  'for all pairwise combinations. The formula is: Pair Score = tumor_AND_activation * AND_specificity '
                  '* (1 - |r|), where tumor_AND_activation is the fraction of tumors where both genes exceed their '
                  'SHAP thresholds, AND_specificity is the fraction of normal tissues where at least one gene is below '
                  'threshold, and r is the Pearson correlation. The UBE2S and CCR6 pair scored highest (0.264), '
                  'exhibiting individual AUCs of 0.9959 and 0.9964, tumor activation of 93.3%, and normal activation '
                  'of 0.6%.',
 'Pipeline': 'The computational design pipeline was built in Python 3.9 and executed on the National Center for '
             'High-Performance Computing (NCHC) biomedical node. The workflow proceeds through nine sequential steps: '
             '(1) Data fetch and decompression, (2) Preprocessing and low-variance filtering, (3) Differential '
             "expression analysis using Welch's t-test and FDR correction, (4) Training and evaluation of machine "
             'learning classifiers (L1 Logistic Regression, Random Forest, XGBoost), (5) SHAP-based feature importance '
             'mapping and local attribution, (6) Inflection-point threshold inference from SHAP dependence plots, (7) '
             'Pairwise orthogonality scoring of candidate genes, (8) Hill-equation-based mathematical simulation of '
             'the AND gate, and (9) Robustness testing via sensitivity sweeps and permutation controls.',
 'QC': 'Quality control (QC) was performed on the raw TOIL TPM matrix. Out of 60,498 annotated genes, near-zero '
       'variance genes (retained variance > 0.0) were filtered, resulting in a clean dataset of 58,581 features. Group '
       'balances were verified to be adequate (178 tumor vs 167 normal). Principal component analysis (PCA) and '
       'density checks were conducted. Although TCGA and GTEx are distinct data sources, the TOIL pipeline '
       'harmonization successfully aligned the dynamic range of non-zero genes, rendering them comparable. Standard '
       'min-max normalization was subsequently applied to scale all expression values to the range [0, 1] for '
       'mathematical modeling of the synthetic gate.',
 'Rationale': 'The desmoplastic stroma of PDAC contains diverse cell populations, including cancer-associated '
              'fibroblasts (CAFs), extracellular matrix components, and infiltrating immune cells, which collectively '
              'modulate tumor progression and therapy resistance. Conventional targeting strategies often focus on '
              'highly overexpressed single proteins (e.g., mesothelin or CEA), which are frequently present at lower '
              "levels in normal tissues (e.g., lung, pleura, or gastrointestinal tract), leading to 'on-target, "
              "off-tumor' toxicity. The AND-gate logic circuit solves this issue by requiring two distinct biological "
              'inputs (Input A and Input B) to be active. If only one input is present, the gate remains closed (OFF). '
              'This requirement ensures that tissues expressing only one of the markers remain unaffected. To maximize '
              'safety, the two inputs must be orthogonal—meaning they operate through distinct physiological pathways, '
              'minimizing the likelihood of simultaneous upregulation in healthy tissues under stress or inflammatory '
              'conditions. This study leverages unbiased transcriptomics to identify and model such orthogonal inputs.',
 'References': 'Siegel, R. L., Giaquinto, A. N., & Jemal, A. (2024). Cancer Statistics, 2024. CA: A Cancer Journal for '
               'Clinicians, 74(1), 12-49.\n'
               'Ho, W. J., Jaffee, E. M., & Zheng, L. (2020). The tumour microenvironment in pancreatic cancer — '
               'clinical challenges and opportunities. Nature Reviews Clinical Oncology, 17(9), 527-540.\n'
               'Neesse, A., Algül, H., Tuveson, D. A., & Gress, T. M. (2015). Stromal biology and therapy in '
               'pancreatic cancer: a changing paradigm. Gut, 64(9), 1476-1484.\n'
               'Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances '
               'in Neural Information Processing Systems, 30, 4765-4774.\n'
               'Hill, A. V. (1910). The possible effects of the aggregation of the molecules of haemoglobin on its '
               'dissociation curves. The Journal of Physiology, 40, iv-vii.',
 'Robustness': 'We performed two negative controls to validate the UBE2S + CCR6 AND gate. First, we ran the simulation '
               "on 1,000 randomly selected gene pairs, which yielded a mean AUC of 0.594. The selected pair's AUC "
               '(0.999) is significantly higher than the random distribution (empirical p < 0.0001). Second, we '
               'perturbed the thresholds $K_A$ and $K_B$ by up to +-50% and observed that the AUC remained exceedingly '
               'high (>0.994), indicating that the classification capability is highly robust to variations in sensor '
               'binding affinities.',
 'SHAP': 'We applied SHAP (SHapley Additive exPlanations) to interpret the selected L1 Logistic Regression classifier. '
         'The SHAP summary plot shows the top 15 features driving the classifier. Non-coding transcripts such as '
         'RP11-40C6.2 and AC009065.4 had the highest mean absolute SHAP values. Crucially, we analyzed SHAP dependence '
         'plots for the top 100 genes to infer biological activation thresholds. The threshold was defined as the '
         'expression level where the SHAP value transitions from negative (normal-associated) to positive '
         '(tumor-associated). Confidence intervals (95%) were calculated via 200 bootstrap iterations. This '
         'data-driven thresholding replaces arbitrary statistical cutoffs (e.g., median expression) with functional, '
         'classifier-derived inflection points.',
 'SingleCell': 'Because the discovery analysis is based on bulk RNA-seq, it remains unclear whether UBE2S and CCR6 are '
               'co-expressed within the same malignant cell population or arise from distinct compartments of the '
               'tumor microenvironment. This distinction has direct consequences for circuit architecture. If both '
               'inputs are co-expressed in malignant epithelial cells, a cell-intrinsic transcriptional circuit may be '
               'feasible. If UBE2S is primarily expressed by cycling tumor cells while CCR6 is derived from immune '
               'cells, the pair should instead be interpreted as a tissue-level microenvironmental signature. Future '
               'single-cell RNA-seq and spatial transcriptomic analysis should therefore quantify cell-type-specific '
               'expression, co-expression frequency, and spatial proximity of UBE2S-high and CCR6-high compartments.',
 'Supplementary': 'Supplementary Table 1: Parameter sweep results. The sweep evaluated Hill coefficients n from 1 to 4 '
                  'and leakiness basal values from 0.0 to 0.1. The optimal setting was n=1, P_basal=0.0. Supplementary '
                  'Table 2: List of top 100 SHAP genes and their inferred thresholds. The top genes are primarily '
                  'non-coding RNAs, but protein-coding genes (MMP12, MISP, UBE2S) were prioritized for wet-lab '
                  'accessibility.',
 'Title': 'Data-Driven Design of a Logic-Gated Biosensor via Unbiased Transcriptomic Profiling of Pancreatic Tumor '
          'Microenvironment',
 'WetLab': 'Several experimental avenues merit exploration to translate the computational findings of this study into '
           'functional synthetic biology constructs. One promising direction involves the engineering of synthetic '
           'promoter systems, wherein the upstream regulatory regions of UBE2S and CCR6 would be cloned to drive '
           'orthogonal transcription factors (e.g., tTA and LhG4) in a split-transactivator configuration, enabling '
           'AND-gate logic at the transcriptional level. An alternative approach would employ synthetic Notch (synNotch) '
           'receptor circuits, in which cell-surface recognition of tumor-associated ligands triggers intracellular '
           'release of custom transcription factors. Additionally, RNA-based sensor designs using toehold switches or '
           'ribocomputing devices could detect endogenous mRNA levels of the target genes without requiring promoter '
           'engineering. Functional validation should be conducted in PDAC cell lines (e.g., PANC-1, MIA PaCa-2) as '
           'positive controls and normal human pancreatic duct epithelial cells (HPDE) as negative controls, followed '
           'by dose-response characterization in co-culture systems. Longer-term directions include in vivo validation '
           'using patient-derived xenograft (PDX) mouse models, assessment of circuit stability under metabolic stress, '
           'and exploration of multi-input logic gates (e.g., three-input AND or AND-NOT) to further improve tumor '
           'specificity and reduce off-target activation.'}

zh_sections = {'Abstract': '胰臟導管腺癌 (pancreatic ductal adenocarcinoma, PDAC) '
             '仍然是致死率極高的惡性腫瘤，五年存活率低於12%，主要原因在於診斷較晚且缺乏具特異性的腫瘤生物標記。本研究提出一個無偏差、數據驅動的運算分析管線，旨在篩選最適合用於合成生物學 logic-gated 及閘 '
             '(AND-gate) 生物感測器的候選基因組合，以精準區分胰臟癌與正常胰臟組織。我們整合了 TCGA-PAAD (n=178 腫瘤) 與 GTEx 正常胰臟組織 (n=167 正常) '
             '世代的轉錄體數據，共分析了 58,581 個基因。利用 L1 正則化邏輯斯迴歸 (L1-regularized Logistic Regression) 機器學習模型進行分類，其五折交叉驗證 AUC 達到 '
             '1.000，藉此篩選出具高度預測能力的特徵。接著，利用可解釋型人工智慧 (explainable artificial intelligence, XAI) 技術中的 SHAP '
             '值分析，推估出候選基因在腫瘤與正常組織轉換的「模型推估活化閾值 (model-inferred activation threshold)」。根據基因之間的正交性 (orthogonality) '
             '與組合活化率，我們選定 UBE2S 與 CCR6 為最終候選基因組合，兩者皮爾森相關係數 (Pearson correlation) 為 0.714。利用希爾方程式 (Hill equation) '
             '對該及閘進行數學建模與參數優化，結果顯示該及閘在發現世代中的 AUC 達 0.9986，準確度 (Accuracy) 為 98.6%、敏感度 (Sensitivity) 為 97.8%、特異度 '
             '(Specificity) 為 99.4% (輸出活化閾值設為 0.25)。經由 1,000 次隨機基因組合的排列測試 (permutation test)，證實此結果具有極顯著的統計學意義 (p < '
             '0.0001)。在外在驗證世代 (GSE62452 微陣列數據，n=130) 中，該及閘維持了極高的特異度 (98.4%)，但敏感度降至 4.3%，顯示跨平台轉換 (RNA-seq vs '
             'Microarray) 在閾值轉移上仍具挑戰。本研究建立的分析管線為合成生物感測器的邏輯閘設計提供了一套穩健且具可重現性的運算框架。',
 'Affiliation': '$^1$國立臺灣大學 生命科學系\\\\$^2$國立臺灣大學 生化科技學系',
 'Author': '施貞蓉$^1$、宿淂芳$^1$、廖軒佑$^2$、林家誼$^2$',
 'Conclusion': '總結而言，本研究成功建立了一套數據驅動的運算框架，用於篩選胰臟癌及閘生物感測器的最優輸入基因對。透過結合差異表達、機器學習、可解釋型人工智慧與數學模擬，選定 UBE2S 與 CCR6 '
               '作為雙輸入特徵。此組合不僅在電腦模擬中表現出極高的分類準確度與特異度，更代表了胰臟癌的兩個核心特徵：細胞週期失控與腫瘤微環境發炎。本分析管線具備高度的可重現性與擴展性，可推廣至其他癌症類型或多輸入邏輯閘的設計中。',
 'DE': "我們對 178 個胰臟癌樣本與 167 個正常組織進行了差異表現分析。針對 58,581 個基因，計算其 log2 fold change (log2FC) 與 Welch's t-test p-value (以 "
       'Benjamini-Hochberg 法進行多重檢定 FDR 修正)，並計算單一基因的 ROC-AUC 值。腫瘤高表達候選基因定義為 log2FC >= 1.0 且 FDR < 0.05，共篩選出 19,399 '
       '個符合條件的基因。為進一步排序，我們計算了特徵特異性得分 (Specificity Score = AUC * log2FC)。火山圖 (Volcano plot) 清楚標示出 UBE2S 與 CCR6 '
       '等基因在胰臟癌中呈現顯著上調。',
 'DataSources': '為建立具代表性的發現世代 (discovery cohort)，我們整合了兩個大型公共資料庫：癌症基因體圖譜 (TCGA-PAAD，包含 178 個胰臟癌腫瘤樣本) 與健康型態基因體表達雙向庫 '
                '(GTEx，包含 167 個正常胰臟組織樣本)。數據格式採用一致化處理的 RSEM TPM，並透過 UCSC Xena TOIL 管線進行批次效應修正。在外在驗證方面，我們自 GEO 下載了獨立的 '
                'GSE62452 世代 (共 130 個樣本，包含 69 個腫瘤樣本與 61 個相鄰正常組織樣本)，其平台為 Affymetrix GPL6244 微陣列 '
                '(microarray)。此種發現與驗證的雙重資料庫設計，能有效評估所篩選基因在跨世代、跨平台技術下的表現與穩定性。',
 'Date': '2026年5月',
 'HillModeling': '我們基於雙輸入希爾方程式 (Hill equation) 對該及閘進行定量建模：\n'
                 '\n'
                 '\\[\n'
                 'Output = P_{\\text{basal}} + V_{\\max} \\left( \\frac{[A]^n}{K_A^n + [A]^n} \\right) \\left( '
                 '\\frac{[B]^n}{K_B^n + [B]^n} \\right)\n'
                 '\\]\n'
                 '\n'
                 '式中，[A] 與 [B] 分別代表 UBE2S 與 CCR6 歸一化後的轉錄體豐度；$K_A$ 與 $K_B$ 為模型推估之活化閾值；$n$ 為反應陡峭度 (Hill '
                 'coefficient)；$P_{\\text{basal}}$ 為基底洩漏量 (basal leakiness)；$V_{\\max}$ 為最大輸出量。經網格掃描，最優參數為：$n = '
                 '1$、$P_{\\text{basal}} = 0.0$、$K_A = 0.760$ (UBE2S 閾值)、$K_B = 0.464$ (CCR6 閾值)。及閘輸出判斷的決策閾值 (decision '
                 'threshold) 設定為 0.25，以最大化分類準確度。',
 'InSilico': '在 345 個樣本的發現世代中進行希爾方程式 AND gate 電腦模擬驗證，結果展現出近乎完美的分類效能。及閘輸出之 ROC-AUC 達到 0.9986，整體準確度為 98.55%，敏感度為 '
             '97.75%，特異度為 99.40%。在 178 個胰臟癌樣本中，有 174 個被正確預測為陽性；在 167 個正常組織中，僅有 1 個發生假陽性活化 (假陽性率僅 '
             '0.6%)。這證實了雙輸入邏輯及閘在維持極高特異度的同時，並未顯著犧牲敏感度。',
 'Introduction': '胰臟導管腺癌 (PDAC) 的病理特徵包括隱匿性病程發展、極強的早期轉移能力，以及由多種細胞與細胞外基質構成的高度促結締組織增生基質 (desmoplastic stroma) 與胰臟腫瘤微環境 '
                 '(tumor microenvironment, TME)。這使得傳統系統性化療與新型的免疫檢查點抑制劑 (immune checkpoint inhibitors) 或嵌合抗原受體 T 細胞 '
                 '(CAR-T) 療法的療效受到嚴重限制。目前臨床應用的主要瓶頸在於缺乏單一特異性抗原，許多腫瘤相關抗原在正常組織中亦有低量表達，極易導致嚴重的脫靶毒性 (off-target '
                 'toxicity)。合成生物學 (synthetic biology) 提供了解決此一困境的新路徑。透過在細胞或基因層次建構邏輯閘 (logic gates) 電路，例如 AND 閘 (及閘) '
                 '電路，只有在兩種輸入訊號 (Input A 與 Input B) 同時高於特定閾值時，才會觸發 downstream '
                 '報導基因或治療性載荷的釋放。這種雙輸入設計能呈指數級地提升對腫瘤細胞的辨識特異度，降低正常組織的假陽性活化率。本研究即致力於利用無偏差轉錄體剖析 (unbiased transcriptomic '
                 'profiling)，開發一套系統化的數據驅動及閘感測器設計流程。',
 'Limitations': '本研究存在若干關鍵限制，在此必須加以說明。其一，本研究僅為電腦模擬層面的概念驗證，實際生物化學工程建構之合成電路可能展現截然不同的反應動力學。其二，SHAP 推估之活化閾值為分類器行為的統計拐點，並非實驗量測之生化解離常數。其三，組織層級的 bulk RNA-seq 數據反映的是細胞群體的平均表現，極易受到腫瘤純度、基質密度與免疫細胞浸潤的影響，可能掩蓋了細胞類型特異性的表達模式。其四，儘管已透過 TOIL 管線進行數據標準化，TCGA 腫瘤樣本與 GTEx 正常組織的比較仍可能存在殘餘的批次效應。其五，外部驗證世代顯示極低的敏感度（4.3%），凸顯了從 RNA-seq 到微陣列的跨平台閾值轉移具有重大挑戰。其六，選定的候選基因對（UBE2S + CCR6）在統計上並非嚴格正交，其在 bulk 數據中的 Spearman 相關係數達 0.714。其七，轉錄本豐度並不保證感測器在蛋白質層級的可及性或等量的蛋白質翻譯。其八，將這些候選基因轉化為功能性合成電路，需要啟動子工程或 RNA 感測器設計，各自引入額外的設計複雜度。最後，任何診斷或治療性的臨床應用，都需要在適當的模式生物中進行廣泛的濕實驗驗證與安全性測試，方能考慮臨床轉譯。',
 'ML': '為獲得最稀疏且最具預測力的基因特徵，我們將數據按 80/20 比例劃分訓練集與測試集，訓練了三種機器學習模型：L1 正則化邏輯斯迴歸、隨機森林 (Random Forest) 以及 XGBoost。結果顯示，L1 '
       '邏輯斯迴歸在測試集上取得了 1.000 的完美 AUC 以及 100.0% 的準確度，其五折交叉驗證 (5-fold CV) 的 AUC 亦為 1.000 +- 0.000。隨機森林測試集 AUC 為 '
       '0.9983，XGBoost 則為 1.000。基於 L1 正則化能將無效權重歸零的稀疏性特徵，以及優異的分類效能，我們選擇 L1 邏輯斯迴歸模型作為後續 SHAP 解釋分析的基礎。',
 'Orthogonality': '及閘設計的核心在於輸入基因對的「正交性 (orthogonality)」。雖然 UBE2S 與 CCR6 在 bulk RNA-seq 數據中表現出中度相關 (r = 0.714)，高於設定的 '
                  '|r| <= 0.4 理想篩選標準，但本分析管線在考量整體綜合得分最優的情況下觸發了 fallback 機制。在生物學機制層面，UBE2S (泛素結合酶 E2 S) 參與細胞分裂週期 (Cell '
                  'cycle / Ubiquitin-proteasome pathway)，而 CCR6 (C-C 趨化因子受體 6) 則參與趨化因子與免疫微環境訊號傳遞 (Chemokine signaling '
                  '/ Immune '
                  'microenvironment)。兩者由完全不同的上游信號通路控制，具備功能上的正交性。這在臨床應用上至關重要，因為正常組織幾乎不可能同時發生細胞週期加速與趨化因子受體上調，從而確保了生物感測器的極高安全性。',
 'PairSelection': '為從前 100 個 SHAP 基因中選出最佳的及閘輸入，我們對所有可能的兩兩組合計算了綜合評分：Pair Score = tumor_AND_activation * AND_specificity '
                  '* (1 - |r|)。其中 tumor_AND_activation 代表兩基因同時高於各自閾值的腫瘤比例，AND_specificity 代表至少有一基因低於閾值的正常組織比例，r '
                  '為皮爾森相關係數。經掃描，UBE2S + CCR6 組合脫穎而出，得分最高 (0.264)。該基因對個別 AUC 分別為 0.9959 與 0.9964，在腫瘤中的同時活化率為 '
                  '93.3%，在正常組織中的活化率僅為 0.6%。',
 'Pipeline': '本研究的運算分析管線完全使用 Python 3.9 開發，並在台灣國家高速網路與計算中心 (NCHC) 的生物醫學節點上執行。管線包含九個核心步驟：(1) 數據自動下載與解壓縮，(2) '
             "樣本分組提取與近零變異基因過濾，(3) 基於 Welch's t-test 與 Benjamini-Hochberg FDR 修正的差異表達分析，(4) 機器學習分類器 (L1 "
             '正則化邏輯斯迴歸、隨機森林、XGBoost) 的訓練與交叉驗證，(5) 應用 SHAP 分析模型進行特徵重要性 (feature importance) 排序，(6) 由 SHAP 依賴圖推估特徵的活化閾值 '
             '(Inflection point)，(7) 進行候選基因對的正交性評估與綜合評分，(8) 基於希爾方程式對及閘進行雙輸入數學建模與參數掃描，以及 (9) 藉由閾值敏感度分析與 1,000 '
             '次隨機基因對排列控制進行穩健性評估。',
 'QC': '我們對 TOIL 平台的 TPM 表達矩陣進行了嚴格的品質控制 (QC)。在 60,498 個註冊基因中，過濾掉在所有樣本中變異度為零的基因，最終保留 58,581 個基因。樣本分組分布均衡 (178 個腫瘤樣本，167 '
       '個健康樣本)。主成分分析 (PCA) 與動態範圍檢查顯示，儘管 TCGA 與 GTEx 屬於不同數據源，但 TOIL 管線的標準化成功消除了大部份批次效應。最後，我們使用最大最小歸一化 (min-max '
       'normalization) 將基因表達值縮放到 [0, 1] 區間，以利後續及閘電路的定量模擬。',
 'Rationale': '胰臟導管腺癌的腫瘤微環境 (TME) 極其複雜，包含癌症相關纖維母細胞 (CAFs)、免疫細胞及血管系統。傳統單抗原靶向治療 (如靶向 Mesothelin 或 CEA) 常因這些蛋白在健康組織 '
              "(如間皮細胞或腸道上皮) 的低量表達而引發 'on-target, off-tumor' 副作用。及閘 (AND-gate) 邏輯感測器則要求細胞必須同時具備兩個特徵 (特徵 A 與特徵 B) "
              '才能激活。若細胞僅表達單一特徵，感測器則保持關閉 (OFF) 狀態。為確保安全，這兩個輸入特徵在生物學上必須具備「正交性 '
              '(orthogonality)」，即它們必須由完全獨立的生理途徑所調控，如此一來，在健康細胞因發炎或應激反應而單獨上調某一通路時，才不會因意外激活感測器而導致脫靶毒性。本研究藉由無偏差轉錄體大數據，精準篩選並定量模擬此類具備正交性的輸入特徵組合。',
 'References': 'Siegel, R. L., Giaquinto, A. N., & Jemal, A. (2024). Cancer Statistics, 2024. CA: A Cancer Journal for '
               'Clinicians, 74(1), 12-49.\n'
               'Ho, W. J., Jaffee, E. M., & Zheng, L. (2020). The tumour microenvironment in pancreatic cancer — '
               'clinical challenges and opportunities. Nature Reviews Clinical Oncology, 17(9), 527-540.\n'
               'Neesse, A., Algül, H., Tuveson, D. A., & Gress, T. M. (2015). Stromal biology and therapy in '
               'pancreatic cancer: a changing paradigm. Gut, 64(9), 1476-1484.\n'
               'Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances '
               'in Neural Information Processing Systems, 30, 4765-4774.\n'
               'Hill, A. V. (1910). The possible effects of the aggregation of the molecules of haemoglobin on its '
               'dissociation curves. The Journal of Physiology, 40, iv-vii.',
 'Robustness': '我們進行了兩項穩健性與負控制分析。首先，我們評估了 1,000 對隨機挑選的基因組合，其平均及閘 AUC 僅為 0.594，而 UBE2S + CCR6 組合的 AUC (0.999) 顯著高於隨機分布 '
               '(p < 0.0001)，排除偶然性。其次，我們對 $K_A$ 與 $K_B$ 參數進行了 +-10%、+-25% 與 +-50% 的微擾分析，結果顯示，即使在 +-50% 的極端微擾下，及閘的 AUC '
               '仍維持在 0.994 以上，證明該感測器對生化親和力 (affinity) 的波動具有極高的容錯能力。',
 'SHAP': '我們採用 SHAP 歸因方法對選定的 L1 邏輯斯迴歸模型進行了解釋。SHAP 摘要圖顯示前 15 個對模型決策貢獻最大的基因，其中非編碼 transcripts (如 RP11-40C6.2 與 '
         'AC009065.4) 排名最前。為了濕實驗的可行性，我們針對前 100 個 SHAP 基因進行了「依賴性分析」，尋找 SHAP 值從負值(代表健康組織特徵) 轉為正值 (代表腫瘤特徵) '
         '的交叉點，以此拐點作為「模型推估活化閾值 (model-inferred activation threshold)」。此種數據驅動的閾值推估，避免了傳統使用中位數或均值等主觀劃分方式，更能反映分類器的功能邊界。',
 'SingleCell': '由於本研究的發現分析建立於 bulk RNA-seq，尚無法判斷 UBE2S 與 CCR6 '
               '是否共同表現於同一群惡性上皮細胞，或是分別來自腫瘤微環境中的不同細胞區室。這項區分會直接影響電路架構的設計：若兩個輸入在惡性上皮細胞中共同表現，則細胞內部轉錄電路較具可行性；若 UBE2S '
               '主要來自增殖中的癌細胞，而 CCR6 主要反映免疫細胞或發炎微環境，則此組合應被解讀為組織層級的微環境特徵，而非單一細胞內部的 AND gate。後續應利用 single-cell RNA-seq 與 '
               'spatial transcriptomics 評估兩基因的細胞類型特異性、共同表現比例與空間鄰近關係。',
 'Subtitle': 'Data-Driven Design of a Logic-Gated Biosensor via Unbiased Transcriptomic Profiling of Pancreatic Tumor '
             'Microenvironment',
 'Supplementary': '補充表格 1：參數掃描結果。此掃描評估了希爾係數 n 從 1 到 4，以及基底洩漏量從 0.0 到 0.1 的影響，最優值為 n=1, P_basal=0.0。\n'
                  '補充表格 2：前 100 個 SHAP 候選基因及其推估閾值。雖然許多非編碼基因排在前列，但考量到啟動子設計的可行性，我們優先挑選了 UBE2S 與 CCR6 等蛋白編碼基因。',
 'Title': '以無偏差轉錄體剖析進行胰臟腫瘤微環境之邏輯閘生物感測器的數據驅動設計',
 'WetLab': '為將本研究的運算結果轉化為功能性合成生物線路，未來有數個實驗方向值得深入探索。首先，可嘗試建構合成啟動子系統，將 UBE2S 與 CCR6 的上游調控區段分別克隆至驅動正交轉錄因子（如 tTA 與 LhG4）的載體中，以 split-transactivator 架構實現轉錄層級的 AND 閘邏輯。其次，可利用合成 Notch（synNotch）受體線路，藉由細胞表面對腫瘤相關配體的辨識，觸發細胞內部客製化轉錄因子的釋放。此外，基於 RNA 的感測器設計（如 toehold switches 或 ribocomputing devices），可直接偵測目標基因的內源性 mRNA 濃度，無需啟動子工程。功能驗證應先以胰臟癌細胞株（如 PANC-1、MIA PaCa-2）做為陽性對照，人類正常胰管上皮細胞（HPDE）做為陰性對照，進行劑量反應特性分析。中長期方向則包含在患者來源異種移植（PDX）小鼠模型中進行體內驗證、評估電路在代謝壓力下的穩定性，以及探索多輸入邏輯閘（如三輸入 AND 閘或 AND-NOT 閘）以進一步提升腫瘤特異度與降低脫靶活化風險。'}

# ----------------- LATEX GENERATORS -----------------

def escape_text_for_latex(text):
    import re
    pattern = r"(\$.*?\$|\\\[.*?\\\])"
    tokens = re.split(pattern, text, flags=re.DOTALL)
    escaped_tokens = []
    for i, token in enumerate(tokens):
        if i % 2 == 0:
            escaped_token = token.replace("_", "\\_").replace("%", "\\%")
            escaped_tokens.append(escaped_token)
        else:
            escaped_tokens.append(token)
    return "".join(escaped_tokens)

def clean_math_for_word(text):
    if not isinstance(text, str):
        return text
    import re
    # Replace display equations \\[ ... \\] with simpler text representation
    def repl_display(m):
        eq = m.group(1).strip()
        eq = eq.replace(r"\text{basal}", "basal")
        eq = eq.replace(r"\max", "max")
        eq = eq.replace(r"\left(", "(")
        eq = eq.replace(r"\right)", ")")
        eq = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1/\2)", eq)
        eq = eq.replace("{", "").replace("}", "")
        eq = eq.replace("_", "")
        return f"\n{eq}\n"
    
    text = re.sub(r"\\\[(.*?)\\\]", repl_display, text, flags=re.DOTALL)
    
    # Replace inline equations $ ... $ with simplified text
    def repl_inline(m):
        eq = m.group(1).strip()
        eq = eq.replace(r"\text{basal}", "basal")
        eq = eq.replace(r"\max", "max")
        eq = eq.replace(r"\pm", "±")
        eq = eq.replace(r"\in", "in")
        eq = eq.replace(r"\le", "≤")
        eq = eq.replace(r"\ge", "≥")
        eq = eq.replace(r"\langle", "<")
        eq = eq.replace(r"\rangle", ">")
        eq = eq.replace(r"\infty", "∞")
        eq = eq.replace(r"\{", "{").replace(r"\}", "}")
        eq = eq.replace("\\", "")
        return eq
        
    text = re.sub(r"\$(.*?)\$", repl_inline, text)
    
    text = text.replace(r"\pm", "±")
    text = text.replace(r"\&", "&")
    text = text.replace(r"\%", "%")
    text = text.replace(r"\_", "_")
    return text

def generate_latex_en(is_revised=False):
    filename = "pdac_biosensor_report_en_revised.tex" if is_revised else "pdac_biosensor_report_en.tex"
    filepath = os.path.join(LATEX_DIR, "en", filename)
    
    en_sections_esc = {k: (escape_text_for_latex(v) if isinstance(v, str) else v) for k, v in globals()['en_sections'].items()}
    en_secs = en_sections_esc

    def make_latex_table(headers, data, caption, label, spec=None, span_columns=False):
        actual_span = span_columns and is_revised
        if spec is None:
            spec = "l" + "c"*(len(headers)-1)
        env = "table*" if actual_span else "table"
        pos = "[t]" if actual_span else "[H]"
        latex = f"\\begin{{{env}}}{pos}\n\\centering\n"
        latex += f"\\caption{{{caption}}}\\label{{{label}}}\n"
        if actual_span:
            latex += f"\\resizebox{{\\textwidth}}{{!}}{{\n"
        else:
            latex += f"\\resizebox{{\\linewidth}}{{!}}{{\n"
        latex += f"\\begin{{tabular}}{{{spec}}}\n\\toprule\n"
        escaped_headers = [h.replace("%", "\\%").replace("_", "\\_") for h in headers]
        latex += " & ".join([f"\\textbf{{{h}}}" for h in escaped_headers]) + " \\\\\n\\midrule\n"
        for row in data:
            escaped_row = [str(x).replace("%", "\\%").replace("_", "\\_").replace("&", "\&").replace("+-", "$\\pm$") for x in row]
            latex += " & ".join(escaped_row) + " \\\\\n"
        latex += f"\\bottomrule\n\\end{{tabular}}\n}}\n\\end{{{env}}}\n"
        return latex

    fig_width = "0.95\\linewidth" if is_revised else "0.7\\textwidth"
    fig_placement = "[htbp]" if is_revised else "[H]"
    
    if is_revised:
        intro_text = en_secs['Introduction']
        if intro_text.startswith("Pancreatic"):
            intro_text = "\\IEEEPARstart{P}{ancreatic}" + intro_text[10:]
        target_sentence = "Synthetic biology provides a powerful paradigm to address this challenge by engineering logic-gated genetic circuits. "
        replacement_sentence = target_sentence + "\\IEEEpubidadjcol "
        intro_text = intro_text.replace(target_sentence, replacement_sentence)
        
        preamble = f"""% !TEX program = xelatex
% !BIB program = biber
% Auto-generated English report for PDAC Biosensor Project
\\documentclass[12pt, journal]{{IEEEtran}}

\\usepackage{{fontspec}}
\\usepackage{{graphicx}}
\\graphicspath{{{{../shared/figures/}}}}
\\usepackage{{float}}
\\usepackage{{booktabs}}
\\usepackage{{amsmath}}
\\usepackage{{amssymb}}
\\usepackage{{siunitx}}
\\usepackage{{xcolor}}
\\usepackage{{url}}
\\usepackage{{hyperref}}
\\usepackage{{subcaption}}
\\usepackage{{array}}
\\usepackage{{multirow}}
\\usepackage{{tabularx}}

\\setmainfont{{Times New Roman}}
\\setsansfont{{Arial}}
\\setmonofont{{Courier New}}

\\renewcommand{{\\IEEEPARstart}}[2]{{\\noindent\\textbf{{\\huge #1}}\\textsc{{#2}}}}

\\usepackage[style=numeric,backend=biber]{{biblatex}}
\\addbibresource{{references_en.bib}}

\\markboth{{Cognitive Security Vol. X Issue. X}}{{}}
\\IEEEpubid{{XXXXXXX/csip.XXXXXXXX  ~\\copyright~2026 CSI Press}}

\\title{{{en_secs['Title']}}}

\\author{{\\IEEEauthorblockN{{SHIH, CHEN-JUNG\\IEEEauthorrefmark{{1}}, SU, TE-FANG\\IEEEauthorrefmark{{1}}, LIAO, XUAN-YOU\\IEEEauthorrefmark{{2}}, and LIN, CHIA-I\\IEEEauthorrefmark{{2}}}} \\\\
\\vspace{{4pt}}
\\IEEEauthorblockA{{\\footnotesize
\\IEEEauthorrefmark{{1}}Department of Life Science, National Taiwan University, Taipei, Taiwan \\\\
\\IEEEauthorrefmark{{2}}Department of Biochemical Science and Technology, National Taiwan University, Taipei, Taiwan}}
\\thanks{{\\hrule \\vspace{{4pt}} \\noindent Manuscript received May 25, 2026; revised May 25, 2026. \\vspace{{3pt}} \\\\
Corresponding Author Email: \\href{{mailto:email@example.com}}{{email@example.com}} \\vspace{{3pt}}}}
}}

\\IEEEaftertitletext{{\\vspace{{-1\\baselineskip}}\\noindent\\begin{{abstract}}
{en_secs['Abstract']}
\\end{{abstract}}
\\noindent\\begin{{IEEEkeywords}}
Pancreatic ductal adenocarcinoma, AND-gate biosensor, transcriptomics, explainable AI, UBE2S, CCR6
\\end{{IEEEkeywords}}
\\vspace{{1\\baselineskip}}}}

\\begin{{document}}
\\maketitle
"""
    else:
        intro_text = en_secs['Introduction']
        doc_class = "\\documentclass[12pt, a4paper]{article}"
        style_file = "report_style.tex"
        
        preamble = f"""% !TEX program = xelatex
% !BIB program = biber
% Auto-generated English report for PDAC Biosensor Project
{doc_class}

\\input{{../shared/{style_file}}}
\\input{{../shared/macros.tex}}

\\setmainfont{{Times New Roman}}
\\setsansfont{{Arial}}
\\setmonofont{{Courier New}}

\\usepackage[style=apa,backend=biber]{{biblatex}}
\\addbibresource{{references_en.bib}}

\\title{{{en_secs['Title']}}}
\\author{{{en_secs['Author']}}}
\\affil{{{en_secs['Affiliation']}}}
\\date{{{en_secs['Date']}}}

\\begin{{document}}
\\maketitle
\\newpage

\\begin{{abstract}}
{en_secs['Abstract']}
\\end{{abstract}}
\\newpage
"""

    content = preamble + f"""
\\section{{Introduction}}
{intro_text}

\\section{{Scientific Rationale and Unmet Need}}
{en_secs['Rationale']}

\\section{{Data Sources}}
{en_secs['DataSources']}

{make_latex_table(t1_headers, t1_data, "Data Cohorts and Sample Size Distribution", "tab:datasets", span_columns=False)}

\\section{{Computational Pipeline}}
{en_secs['Pipeline']}

\\section{{Quality Control and Batch-Effect Assessment}}
{en_secs['QC']}

\\section{{Differential Expression Analysis}}
{en_secs['DE']}

\\begin{{figure}}{fig_placement}
\\centering
\\includegraphics[width={fig_width}]{{volcano_discovery.png}}
\\caption{{Volcano plot highlighting significantly differentially expressed genes in the discovery cohort (TCGA-PAAD vs GTEx Normal Pancreas). UBE2S and CCR6 are annotated as significant upregulated candidates.}}
\\label{{fig:volcano}}
\\end{{figure}}

{make_latex_table(t2_headers, t2_data, "Top 10 Differentially Expressed Genes Sorted by Specificity Score", "tab:top_de", span_columns=True)}

\\section{{Machine Learning Classifier Performance}}
{en_secs['ML']}

{make_latex_table(t3_headers, t3_data, "Machine Learning Classifier Performance and Cross-Validation Summary", "tab:ml_perf", span_columns=True)}

\\section{{SHAP-Based Explainable AI Analysis}}
{en_secs['SHAP']}

\\begin{{figure}}{fig_placement}
\\centering
\\includegraphics[width={fig_width}]{{shap_summary.png}}
\\caption{{SHAP summary plot showing feature importances for the top-ranked genes driving the L1 Logistic Regression classifier.}}
\\label{{fig:shap_summary}}
\\end{{figure}}

{make_latex_table(t4_headers, t4_data, "Top 10 SHAP Feature Importance and Inferred Expression Thresholds", "tab:shap_thresh", span_columns=True)}

\\section{{Candidate Gene Pair Selection}}
{en_secs['PairSelection']}

{make_latex_table(t5_headers, t5_data, "Detailed Molecular Profile of Selected Candidate Pair", "tab:candidate_pair", span_columns=True)}

\\section{{Orthogonality Assessment}}
{en_secs['Orthogonality']}

\\begin{{figure}}{fig_placement}
\\centering
\\includegraphics[width={fig_width}]{{gene_pair_scatter_final.png}}
\\caption{{Scatter plot of UBE2S vs CCR6 rescaled expression in discovery cohort, demonstrating decision boundary quadrants and sample clustering.}}
\\label{{fig:scatter}}
\\end{{figure}}

\\section{{Need for Single-Cell and Spatial Validation}}
{en_secs['SingleCell']}

\\section{{Hill-Equation-Based AND Gate Modeling}}
{en_secs['HillModeling']}

\\section{{In Silico Validation}}
{en_secs['InSilico']}

\\begin{{figure}}{fig_placement}
\\centering
\\includegraphics[width={fig_width}]{{and_gate_heatmap_final.png}}
\\caption{{2D Contour heatmap demonstrating simulation output of UBE2S AND CCR6 logical AND gate, with rescaled expression and tumor/normal sample overlays.}}
\\label{{fig:heatmap}}
\\end{{figure}}

{make_latex_table(t6_headers, t6_data, "AND Gate Simulation Performance Summary in Discovery Cohort", "tab:and_perf", span_columns=False)}

\\begin{{figure}}{fig_placement}
\\centering
\\includegraphics[width={fig_width}]{{roc_curves.png}}
\\caption{{ROC curves comparing individual inputs (UBE2S, CCR6) against the combined logic gate output.}}
\\label{{fig:roc}}
\\end{{figure}}

{make_latex_table(t7_headers, t7_data, "External Validation Results on GSE62452 Dataset", "tab:ext_val", span_columns=True)}

\\section{{Robustness and Negative Controls}}
{en_secs['Robustness']}

\\section{{Limitations}}
{en_secs['Limitations']}

\\section{{Future Experimental Directions}}
{en_secs['WetLab']}

\\section{{Conclusion}}
{en_secs['Conclusion']}

"""
    bib_prefix = "" if is_revised else "\\newpage\n"
    content += f"""{bib_prefix}\\nocite{{*}}
\\printbibliography[title={{References}}]

"""
    supp_prefix = "\\newpage\n\\onecolumn\n" if is_revised else "\\newpage\n"
    
    content += f"""{supp_prefix}\\section{{Supplementary Tables and Figures}}
{en_secs['Supplementary']}

\\subsection{{Supplementary Table 1: Parameter Sweep}}
{make_latex_table(["Hill Coefficient (n)", "Basal Leakiness (P_basal)", "ROC-AUC", "Accuracy", "Sensitivity", "Specificity"], 
                  [[str(x[0]), str(x[1]), f"{x[2]:.4f}", f"{x[4]*100:.1f}\\%", f"{x[5]*100:.1f}\\%", f"{x[6]*100:.1f}\\%"] for x in and_sweep.values[:6]] if and_sweep is not None else [], 
                  "Hill Equation Grid Search Parameter Sweep Performance", "tab:supp_sweep", span_columns=True)}

\\subsection{{Supplementary Table 2: Threshold Sensitivity}}
{make_latex_table(t8_headers, t8_data, "Sensitivity Analysis of K Parameter Perturbations", "tab:supp_sensitivity", span_columns=True)}

\\subsection{{Supplementary Figures}}
\\begin{{figure}}[H]
\\centering
\\begin{{subfigure}}[b]{{0.48\\textwidth}}
\\centering
\\includegraphics[width=\\textwidth]{{shap_dependence_top_genes/UBE2S_shap_dependence.png}}
\\caption{{UBE2S SHAP dependence}}
\\end{{subfigure}}
\\hfill
\\begin{{subfigure}}[b]{{0.48\\textwidth}}
\\centering
\\includegraphics[width=\\textwidth]{{shap_dependence_top_genes/CCR6_shap_dependence.png}}
\\caption{{CCR6 SHAP dependence}}
\\end{{subfigure}}
\\caption{{SHAP dependence plots for selected candidate genes UBE2S and CCR6 showing threshold transition inflection points.}}
\\label{{fig:supp_shap_dependence}}
\\end{{figure}}

\\end{{document}}
"""
    with open(filepath, "w") as f:
        f.write(content)
    print(f"Generated {filepath}")


def generate_latex_zh(is_revised=False):
    filename = "pdac_biosensor_report_zh_revised.tex" if is_revised else "pdac_biosensor_report_zh.tex"
    filepath = os.path.join(LATEX_DIR, "zh", filename)
    
    zh_sections_esc = {k: (escape_text_for_latex(v) if isinstance(v, str) else v) for k, v in globals()['zh_sections'].items()}
    zh_secs = zh_sections_esc

    def make_latex_table(headers, data, caption, label, spec=None, span_columns=False):
        actual_span = False  # Chinese is always single-column, no table* needed
        if spec is None:
            spec = "l" + "c"*(len(headers)-1)
        env = "table"
        pos = "[H]"
        latex = f"\\begin{{{env}}}{pos}\n\\centering\n"
        latex += f"\\caption{{{caption}}}\\label{{{label}}}\n"
        latex += f"\\resizebox{{\\linewidth}}{{!}}{{\n"
        latex += f"\\begin{{tabular}}{{{spec}}}\n\\toprule\n"
        escaped_headers = [h.replace("%", "\\%").replace("_", "\\_") for h in headers]
        latex += " & ".join([f"\\textbf{{{h}}}" for h in escaped_headers]) + " \\\\\n\\midrule\n"
        for row in data:
            escaped_row = [str(x).replace("%", "\\%").replace("_", "\\_").replace("&", "\\&").replace("+-", "$\\pm$") for x in row]
            latex += " & ".join(escaped_row) + " \\\\\n"
        latex += f"\\bottomrule\n\\end{{tabular}}\n}}\n\\end{{{env}}}\n"
        return latex

    # Chinese: always single-column for both revised and non-revised
    doc_class = "\\documentclass[12pt, a4paper]{article}"
    style_file = "report_style.tex"
    fig_width = "0.7\\textwidth"

    # Both revised and non-revised Chinese use single-column title block
    title_block = f"""
\\maketitle
\\newpage

\\begin{{abstract}}
{zh_secs['Abstract']}
\\end{{abstract}}
\\newpage
"""

    content = f"""% !TEX program = xelatex
% !BIB program = biber
% Auto-generated Chinese report for PDAC Biosensor Project
{doc_class}

\\input{{../shared/{style_file}}}
\\input{{../shared/macros.tex}}
\\usepackage{{xeCJK}}
\\setCJKmainfont{{DFKai-SB}}
\\setCJKsansfont{{DFKai-SB}}

\\setmainfont{{Times New Roman}}
\\setsansfont{{Arial}}
\\setmonofont{{Courier New}}

\\usepackage[style=apa,backend=biber]{{biblatex}}
\\addbibresource{{references_zh.bib}}

\\title{{{zh_secs['Title']}}}
\\author{{{zh_secs['Author']}}}
\\affil{{{zh_secs['Affiliation']}}}
\\date{{{zh_secs['Date']}}}

\\begin{{document}}

{title_block}

\\section{{前言}}
{zh_secs['Introduction']}

\\section{{科學背景與未滿足之臨床需求}}
{zh_secs['Rationale']}

\\section{{資料來源}}
{zh_secs['DataSources']}

{make_latex_table(t1_headers, t1_data, "數據世代與樣本量分布", "tab:datasets_zh", span_columns=False)}

\\section{{運算分析管線}}
{zh_secs['Pipeline']}

\\section{{品質控制與批次效應評估}}
{zh_secs['QC']}

\\section{{差異表現分析}}
{zh_secs['DE']}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width={fig_width}]{{volcano_discovery.png}}
\\caption{{發現世代中的火山圖，標記顯著差異表現基因。其中 UBE2S 與 CCR6 被註記為顯著上調之候選基因。}}
\\label{{fig:volcano_zh}}
\\end{{figure}}

{make_latex_table(t2_headers, t2_data, "前 10 個差異表現基因 (按特異性得分排序)", "tab:top_de_zh", span_columns=True)}

\\section{{機器學習分類器表現}}
{zh_secs['ML']}

{make_latex_table(t3_headers, t3_data, "機器學習分類器表現與五折交叉驗證結果摘要", "tab:ml_perf_zh", span_columns=True)}

\\section{{基於 SHAP 的可解釋型人工智慧分析}}
{zh_secs['SHAP']}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width={fig_width}]{{shap_summary.png}}
\\caption{{SHAP 摘要圖，顯示推動 L1 邏輯斯迴歸分類器決策的前 15 個特徵重要性排名。}}
\\label{{fig:shap_summary_zh}}
\\end{{figure}}

{make_latex_table(t4_headers, t4_data, "前 10 個 SHAP 特徵重要性與模型推估表達閾值", "tab:shap_thresh_zh", span_columns=True)}

\\section{{候選基因組合篩選}}
{zh_secs['PairSelection']}

{make_latex_table(t5_headers, t5_data, "選定候選基因組合的詳細分子譜描述", "tab:candidate_pair_zh", span_columns=True)}

\\section{{正交性評估}}
{zh_secs['Orthogonality']}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width={fig_width}]{{gene_pair_scatter_final.png}}
\\caption{{UBE2S 與 CCR6 在發現世代中的表達量散佈圖，呈現決策邊界與樣本分布象限。}}
\\label{{fig:scatter_zh}}
\\end{{figure}}

\\section{{單細胞與空間轉錄體驗證之必要性}}
{zh_secs['SingleCell']}

\\section{{基於希爾方程式的 AND gate 建模}}
{zh_secs['HillModeling']}

\\section{{電腦模擬驗證}}
{zh_secs['InSilico']}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width={fig_width}]{{and_gate_heatmap_final.png}}
\\caption{{UBE2S 與 CCR6 及閘模擬輸出之二維等高線熱圖，並疊加腫瘤與健康樣本。}}
\\label{{fig:heatmap_zh}}
\\end{{figure}}

{make_latex_table(t6_headers, t6_data, "及閘模擬在發現世代中的分類效能指標摘要", "tab:and_perf_zh", span_columns=False)}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width={fig_width}]{{roc_curves.png}}
\\caption{{比較單一基因輸入 (UBE2S, CCR6) 與雙輸入邏輯及閘輸出的 ROC 曲線。}}
\\label{{fig:roc_zh}}
\\end{{figure}}

{make_latex_table(t7_headers, t7_data, "GSE62452 外部驗證結果摘要", "tab:ext_val_zh", span_columns=True)}

\\section{{穩健性分析與負控制}}
{zh_secs['Robustness']}

\\section{{研究限制}}
{zh_secs['Limitations']}

\\section{{未來實驗方向}}
{zh_secs['WetLab']}

\\section{{結論}}
{zh_secs['Conclusion']}

\\newpage
\\nocite{{*}}
\\printbibliography[title={{參考文獻}}]

\\newpage
\\section{{補充表格與圖}}
{zh_secs['Supplementary']}

\\subsection{{補充表格 1：參數掃描結果}}
{make_latex_table(["Hill 係數 (n)", "基底洩漏量 (P_basal)", "ROC-AUC", "準確度", "敏感度", "特異度"], 
                  [[str(x[0]), str(x[1]), f"{x[2]:.4f}", f"{x[4]*100:.1f}\\%", f"{x[5]*100:.1f}\\%", f"{x[6]*100:.1f}\\%"] for x in and_sweep.values[:6]] if and_sweep is not None else [], 
                  "希爾方程式網格掃描參數表現結果", "tab:supp_sweep_zh", span_columns=True)}

\\subsection{{補充表格 2：K 參數微擾對分類效能影響之敏感度分析}}
{make_latex_table(t8_headers, t8_data, "K 參數微擾對分類效能影響之敏感度分析", "tab:supp_sensitivity_zh", span_columns=True)}

\\subsection{{補充圖}}
\\begin{{figure}}[H]
\\centering
\\begin{{subfigure}}[b]{{0.48\\textwidth}}
\\centering
\\includegraphics[width=\\textwidth]{{shap_dependence_top_genes/UBE2S_shap_dependence.png}}
\\caption{{UBE2S SHAP 依賴性}}
\\end{{subfigure}}
\\hfill
\\begin{{subfigure}}[b]{{0.48\\textwidth}}
\\centering
\\includegraphics[width=\\textwidth]{{shap_dependence_top_genes/CCR6_shap_dependence.png}}
\\caption{{CCR6 SHAP 依賴性}}
\\end{{subfigure}}
\\caption{{選定候選基因 UBE2S 與 CCR6 之 SHAP 依賴圖，呈現出拐點轉換閾值。}}
\\label{{fig:supp_shap_dependence_zh}}
\\end{{figure}}

\\end{{document}}
"""
    with open(filepath, "w") as f:
        f.write(content)
    print(f"Generated {filepath}")

def add_word_paragraph(doc, text, style='Normal', space_after=6, line_spacing=1.5, bold=False, italic=False, font_name='Times New Roman', font_size=12, is_chinese=False):
    text = clean_math_for_word(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if is_chinese:
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    return p

def add_word_heading(doc, text, level, is_chinese=False):
    text = clean_math_for_word(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    font_size = 18 if level == 1 else (14 if level == 2 else 12)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial' if not is_chinese else 'DFKai-SB'
    run.font.size = Pt(font_size)
    if is_chinese:
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    return p

def add_word_table(doc, headers, data, is_chinese=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        title_clean = clean_math_for_word(title)
        hdr_cells[i].text = title_clean
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.name = 'Arial' if not is_chinese else 'DFKai-SB'
        run.font.size = Pt(10)
        if is_chinese:
            run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
            
    for row in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            val_clean = clean_math_for_word(val)
            row_cells[i].text = str(val_clean)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Times New Roman' if not is_chinese else 'DFKai-SB'
                run.font.size = Pt(10)
                if is_chinese:
                    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def add_word_image(doc, filepath, caption, is_chinese=False, width=Inches(5.0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    if os.path.exists(filepath):
        p.add_run().add_picture(filepath, width=width)
    else:
        r = p.add_run(f"[Image placeholder: {filepath}]")
        r.italic = True
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(f"Figure: {clean_math_for_word(caption)}")
    run_cap.italic = True
    run_cap.font.name = 'Times New Roman' if not is_chinese else 'DFKai-SB'
    run_cap.font.size = Pt(10)
    if is_chinese:
        run_cap._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')

def set_section_columns(section, num_cols, col_space=None):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    if num_cols > 1:
        cols.set(qn('w:equalWidth'), '1')
        if col_space is not None:
            cols.set(qn('w:space'), str(col_space))

def generate_word_en(is_revised=False):
    filename = "pdac_biosensor_report_en_revised.docx" if is_revised else "pdac_biosensor_report_en.docx"
    filepath = os.path.join(WORD_DIR, "en", filename)
    doc = docx.Document()
    
    sec1 = doc.sections[0]
    sec1.top_margin = Inches(0.98)
    sec1.bottom_margin = Inches(0.98)
    sec1.left_margin = Inches(0.98)
    sec1.right_margin = Inches(0.98)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(80)
    p_title.paragraph_format.space_after = Pt(18)
    run_title = p_title.add_run(clean_math_for_word(en_sections["Title"]))
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    
    en_word_author = 'SHIH, Chen-Jung\u00b9, SU, Te-Fang\u00b9, LIAO, Xuan-You\u00b2, LIN, Chia-I\u00b2'
    en_word_affiliation = '\u00b9Department of Life Science, National Taiwan University\n\u00b2Department of Biochemical Science and Technology, National Taiwan University'
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(6)
    run_author = p_author.add_run(en_word_author)
    run_author.font.name = 'Times New Roman'
    run_author.font.size = Pt(12)
    
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(18)
    run_aff = p_aff.add_run(en_word_affiliation)
    run_aff.font.name = 'Times New Roman'
    run_aff.font.size = Pt(10)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(clean_math_for_word(en_sections["Date"]))
    run_date.font.name = 'Times New Roman'
    run_date.font.size = Pt(11)
    
    doc.add_page_break()
    
    add_word_heading(doc, "Abstract", 1)
    add_word_paragraph(doc, en_sections["Abstract"])
    
    if is_revised:
        sec2 = doc.add_section(start_type=WD_SECTION.NEW_PAGE)
        set_section_columns(sec2, 2, col_space=720)
        sec2.top_margin = Inches(0.7)
        sec2.bottom_margin = Inches(0.7)
        sec2.left_margin = Inches(0.7)
        sec2.right_margin = Inches(0.7)
        main_fig_width = Inches(3.0)
    else:
        doc.add_page_break()
        main_fig_width = Inches(5.0)
    
    en_limitations_paragraph = (
        "Several critical limitations must be acknowledged. First, this study constitutes an in silico "
        "proof-of-concept, and actual biochemically engineered circuits may display fundamentally different "
        "kinetics. Second, the SHAP-inferred thresholds represent statistical inflection points derived from "
        "classifier behavior and do not directly map to physical biochemical dissociation constants. Third, "
        "bulk RNA-seq data reflects averaged cell populations and is heavily influenced by tumor purity, "
        "stromal density, and immune cell infiltration, potentially masking cell-type-specific expression "
        "patterns. Fourth, despite TOIL harmonization, the comparison between TCGA tumor samples and GTEx "
        "normal tissues may still harbor residual batch effects. Fifth, the external validation cohort "
        "demonstrated extremely low sensitivity (4.3%), highlighting a significant challenge in cross-platform "
        "threshold transfer from RNA-seq to microarray data. Sixth, the selected candidate pair (UBE2S + CCR6) "
        "is not strictly statistically orthogonal, exhibiting a Spearman correlation of 0.714 in bulk data. "
        "Seventh, transcriptomic abundance does not guarantee equivalent sensor accessibility or protein-level "
        "expression. Eighth, translating these candidates into a functional synthetic circuit requires promoter "
        "engineering or RNA-based sensor design, each introducing additional layers of design complexity. "
        "Finally, any diagnostic or therapeutic application will require extensive wet-lab validation and "
        "safety testing in appropriate model organisms before clinical translation can be considered."
    )
        
    sections_order = [
        ("I. Introduction", en_sections["Introduction"]),
        ("II. Scientific Rationale and Unmet Need", en_sections["Rationale"]),
        ("III. Data Sources", en_sections["DataSources"]),
        ("IV. Computational Pipeline", en_sections["Pipeline"]),
        ("V. Quality Control and Batch-Effect Assessment", en_sections["QC"]),
        ("VI. Differential Expression Analysis", en_sections["DE"]),
        ("VII. Machine Learning Classifier Performance", en_sections["ML"]),
        ("VIII. SHAP-Based Explainable AI Analysis", en_sections["SHAP"]),
        ("IX. Candidate Gene Pair Selection", en_sections["PairSelection"]),
        ("X. Orthogonality Assessment", en_sections["Orthogonality"]),
        ("XI. Need for Single-Cell and Spatial Validation", en_sections["SingleCell"]),
        ("XII. Hill-Equation-Based AND Gate Modeling", en_sections["HillModeling"]),
        ("XIII. In Silico Validation", en_sections["InSilico"]),
        ("XIV. Robustness and Controls", en_sections["Robustness"])
    ]
    
    for title, text in sections_order:
        add_word_heading(doc, title, 1)
        add_word_paragraph(doc, text)
        
        if "III. Data Sources" in title:
            add_word_table(doc, t1_headers, t1_data)
        elif "VI. Differential Expression" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "volcano_discovery.png"), "Volcano plot of discovery cohort", width=main_fig_width)
            add_word_table(doc, t2_headers, t2_data)
        elif "VII. Machine Learning" in title:
            add_word_table(doc, t3_headers, t3_data)
        elif "VIII. SHAP-Based" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "shap_summary.png"), "SHAP Feature Importance Summary Bar Plot", width=main_fig_width)
            add_word_table(doc, t4_headers, t4_data)
        elif "IX. Candidate Gene" in title:
            add_word_table(doc, t5_headers, t5_data)
        elif "X. Orthogonality" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "gene_pair_scatter_final.png"), "UBE2S vs CCR6 rescaled expression scatter plot", width=main_fig_width)
        elif "XIII. In Silico Validation" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "and_gate_heatmap_final.png"), "AND gate logical activation surface heatmap", width=main_fig_width)
            add_word_table(doc, t6_headers, t6_data)
            add_word_image(doc, os.path.join(FIGURES_DIR, "roc_curves.png"), "ROC curves comparison: single vs logical combination", width=main_fig_width)
            add_word_heading(doc, "Table 7: External Validation Results (GSE62452)", 2)
            add_word_table(doc, t7_headers, t7_data)
            
    add_word_heading(doc, "XV. Limitations", 1)
    add_word_paragraph(doc, en_limitations_paragraph)
        
    add_word_heading(doc, "XVI. Future Experimental Directions", 1)
    add_word_paragraph(doc, en_sections["WetLab"])
    
    add_word_heading(doc, "XVII. Conclusion", 1)
    add_word_paragraph(doc, en_sections["Conclusion"])
    
    add_word_heading(doc, "XVIII. References", 1)
    refs = en_sections["References"].split("\n")
    for ref in refs:
        add_word_paragraph(doc, ref, space_after=4)
        
    if is_revised:
        sec3 = doc.add_section(start_type=WD_SECTION.NEW_PAGE)
        set_section_columns(sec3, 1)
        sec3.top_margin = Inches(0.98)
        sec3.bottom_margin = Inches(0.98)
        sec3.left_margin = Inches(0.98)
        sec3.right_margin = Inches(0.98)
    else:
        doc.add_page_break()
        
    add_word_heading(doc, "XIX. Supplementary Tables and Figures", 1)
    add_word_paragraph(doc, en_sections["Supplementary"])
    
    add_word_heading(doc, "Supplementary Table 1: Grid Search Sweep", 2)
    add_word_table(doc, ["Hill Coefficient (n)", "Basal Leakiness (P_basal)", "ROC-AUC", "Accuracy", "Sensitivity", "Specificity"],
                   [[str(x[0]), str(x[1]), f"{x[2]:.4f}", f"{x[4]*100:.1f}%", f"{x[5]*100:.1f}%", f"{x[6]*100:.1f}%"] for x in and_sweep.values[:6]] if and_sweep is not None else [])
    
    add_word_heading(doc, "Supplementary Table 2: Sensitivity analysis", 2)
    add_word_table(doc, t8_headers, t8_data)
    
    add_word_heading(doc, "Supplementary Figures: SHAP dependence plots", 2)
    add_word_image(doc, os.path.join(FIGURES_DIR, "shap_dependence_top_genes/UBE2S_shap_dependence.png"), "UBE2S SHAP dependence plot", width=Inches(4.5))
    add_word_image(doc, os.path.join(FIGURES_DIR, "shap_dependence_top_genes/CCR6_shap_dependence.png"), "CCR6 SHAP dependence plot", width=Inches(4.5))
    
    doc.save(filepath)
    print(f"Generated {filepath}")

def generate_word_zh(is_revised=False):
    filename = "pdac_biosensor_report_zh_revised.docx" if is_revised else "pdac_biosensor_report_zh.docx"
    filepath = os.path.join(WORD_DIR, "zh", filename)
    doc = docx.Document()
    
    sec1 = doc.sections[0]
    sec1.top_margin = Inches(0.98)
    sec1.bottom_margin = Inches(0.98)
    sec1.left_margin = Inches(0.98)
    sec1.right_margin = Inches(0.98)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(80)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(clean_math_for_word(zh_sections["Title"]))
    run_title.bold = True
    run_title.font.name = 'DFKai-SB'
    run_title.font.size = Pt(20)
    run_title._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run(clean_math_for_word(zh_sections["Subtitle"]))
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(12)
    
    zh_word_author = '施貞蓉\u00b9、宿淂芳\u00b9、廖軒佑\u00b2、林家誼\u00b2'
    zh_word_affiliation = '\u00b9國立臺灣大學 生命科學系\n\u00b2國立臺灣大學 生化科技學系'
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(6)
    run_author = p_author.add_run(zh_word_author)
    run_author.font.name = 'DFKai-SB'
    run_author.font.size = Pt(12)
    run_author._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(18)
    run_aff = p_aff.add_run(zh_word_affiliation)
    run_aff.font.name = 'DFKai-SB'
    run_aff.font.size = Pt(10)
    run_aff._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(clean_math_for_word(zh_sections["Date"]))
    run_date.font.name = 'DFKai-SB'
    run_date.font.size = Pt(11)
    run_date._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    
    doc.add_page_break()
    
    add_word_heading(doc, "摘要", 1, is_chinese=True)
    add_word_paragraph(doc, zh_sections["Abstract"], is_chinese=True, font_name='DFKai-SB')
    
    # Chinese Word: always single-column for both revised and non-revised
    doc.add_page_break()
    main_fig_width = Inches(5.0)
    
    zh_limitations_paragraph = (
        '本研究存在若干關鍵限制，在此必須加以說明。其一，本研究僅為電腦模擬層面的概念驗證，實際生物化學工程建構之合成電路可能展現截然不同的反應動力學。'
        '其二，SHAP 推估之活化閾值為分類器行為的統計拐點，並非實驗量測之生化解離常數。'
        '其三，組織層級的 bulk RNA-seq 數據反映的是細胞群體的平均表現，極易受到腫瘤純度、基質密度與免疫細胞浸潤的影響，可能掩蓋了細胞類型特異性的表達模式。'
        '其四，儘管已透過 TOIL 管線進行數據標準化，TCGA 腫瘤樣本與 GTEx 正常組織的比較仍可能存在殘餘的批次效應。'
        '其五，外部驗證世代顯示極低的敏感度（4.3%），凸顯了從 RNA-seq 到微陣列的跨平台閾值轉移具有重大挑戰。'
        '其六，選定的候選基因對（UBE2S + CCR6）在統計上並非嚴格正交，其在 bulk 數據中的 Spearman 相關係數達 0.714。'
        '其七，轉錄本豐度並不保證感測器在蛋白質層級的可及性或等量的蛋白質翻譯。'
        '其八，將這些候選基因轉化為功能性合成電路，需要啟動子工程或 RNA 感測器設計，各自引入額外的設計複雜度。'
        '最後，任何診斷或治療性的臨床應用，都需要在適當的模式生物中進行廣泛的濕實驗驗證與安全性測試，方能考慮臨床轉譯。'
    )
        
    sections_order_zh = [
        ("一、前言", zh_sections["Introduction"]),
        ("二、科學背景與未滿足之臨床需求", zh_sections["Rationale"]),
        ("三、資料來源", zh_sections["DataSources"]),
        ("四、運算分析管線", zh_sections["Pipeline"]),
        ("五、品質控制與批次效應評估", zh_sections["QC"]),
        ("六、差異表現分析", zh_sections["DE"]),
        ("七、機器學習分類器表現", zh_sections["ML"]),
        ("八、基於 SHAP 的可解釋型人工智慧分析", zh_sections["SHAP"]),
        ("九、候選基因組合篩選", zh_sections["PairSelection"]),
        ("十、正交性評估", zh_sections["Orthogonality"]),
        ("十一、單細胞與空間轉錄體驗證之必要性", zh_sections["SingleCell"]),
        ("十二、基於希爾方程式的 AND gate 建模", zh_sections["HillModeling"]),
        ("十三、電腦模擬驗證", zh_sections["InSilico"]),
        ("十四、穩健性分析與負控制", zh_sections["Robustness"])
    ]
    
    for title, text in sections_order_zh:
        add_word_heading(doc, title, 1, is_chinese=True)
        add_word_paragraph(doc, text, is_chinese=True, font_name='DFKai-SB')
        
        if "三、資料來源" in title:
            add_word_table(doc, t1_headers, t1_data, is_chinese=True)
        elif "六、差異表現分析" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "volcano_discovery.png"), "發現世代之火山圖", is_chinese=True, width=main_fig_width)
            add_word_table(doc, t2_headers, t2_data, is_chinese=True)
        elif "七、機器學習" in title:
            add_word_table(doc, t3_headers, t3_data, is_chinese=True)
        elif "八、基於 SHAP" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "shap_summary.png"), "SHAP 特徵重要性摘要條形圖", is_chinese=True, width=main_fig_width)
            add_word_table(doc, t4_headers, t4_data, is_chinese=True)
        elif "九、候選基因" in title:
            add_word_table(doc, t5_headers, t5_data, is_chinese=True)
        elif "十、正交性評估" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "gene_pair_scatter_final.png"), "UBE2S 與 CCR6 表達量散佈圖", is_chinese=True, width=main_fig_width)
        elif "十三、電腦模擬驗證" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "and_gate_heatmap_final.png"), "邏輯及閘活化熱圖", is_chinese=True, width=main_fig_width)
            add_word_table(doc, t6_headers, t6_data, is_chinese=True)
            add_word_image(doc, os.path.join(FIGURES_DIR, "roc_curves.png"), "ROC 曲線比較圖", is_chinese=True, width=main_fig_width)
            add_word_heading(doc, "表格 7：外部驗證結果 (GSE62452)", 2, is_chinese=True)
            add_word_table(doc, t7_headers, t7_data, is_chinese=True)
            
    add_word_heading(doc, "十五、研究限制", 1, is_chinese=True)
    add_word_paragraph(doc, zh_limitations_paragraph, is_chinese=True, font_name='DFKai-SB')
        
    add_word_heading(doc, "十六、未來實驗方向", 1, is_chinese=True)
    add_word_paragraph(doc, zh_sections["WetLab"], is_chinese=True, font_name='DFKai-SB')
    
    add_word_heading(doc, "十七、結論", 1, is_chinese=True)
    add_word_paragraph(doc, zh_sections["Conclusion"], is_chinese=True, font_name='DFKai-SB')
    
    add_word_heading(doc, "十八、參考文獻", 1, is_chinese=True)
    refs = zh_sections["References"].split("\n")
    for ref in refs:
        add_word_paragraph(doc, ref, is_chinese=True, font_name='DFKai-SB', space_after=4)
        
    doc.add_page_break()
        
    add_word_heading(doc, "十九、補充表格與圖", 1, is_chinese=True)
    add_word_paragraph(doc, zh_sections["Supplementary"], is_chinese=True, font_name='DFKai-SB')
    
    add_word_heading(doc, "補充表格 1：希爾方程式網格掃描參數表現結果", 2, is_chinese=True)
    add_word_table(doc, ["Hill 係數 (n)", "基底洩漏量 (P_basal)", "ROC-AUC", "準確度", "敏感度", "特異度"],
                   [[str(x[0]), str(x[1]), f"{x[2]:.4f}", f"{x[4]*100:.1f}%", f"{x[5]*100:.1f}%", f"{x[6]*100:.1f}%"] for x in and_sweep.values[:6]] if and_sweep is not None else [], is_chinese=True)
    
    add_word_heading(doc, "補充表格 2：K 參數微擾對分類效能影響之敏感度分析", 2, is_chinese=True)
    add_word_table(doc, t8_headers, t8_data, is_chinese=True)
    
    add_word_heading(doc, "補充圖：SHAP 依賴圖", 2, is_chinese=True)
    add_word_image(doc, os.path.join(FIGURES_DIR, "shap_dependence_top_genes/UBE2S_shap_dependence.png"), "UBE2S SHAP 依賴圖", is_chinese=True, width=Inches(4.5))
    add_word_image(doc, os.path.join(FIGURES_DIR, "shap_dependence_top_genes/CCR6_shap_dependence.png"), "CCR6 SHAP 依賴圖", is_chinese=True, width=Inches(4.5))
    
    doc.save(filepath)
    print(f"Generated {filepath}")

import shutil
import time

def make_backup(path):
    if os.path.exists(path):
        ts = time.strftime("%Y%m%d_%H%M%S")
        backup_path = f"{path}.backup_{ts}"
        shutil.copy2(path, backup_path)
        print(f"Created backup of {path} at {backup_path}")

def main():
    original_files = [
        os.path.join(LATEX_DIR, "en", "pdac_biosensor_report_en.tex"),
        os.path.join(LATEX_DIR, "zh", "pdac_biosensor_report_zh.tex"),
        os.path.join(WORD_DIR, "en", "pdac_biosensor_report_en.docx"),
        os.path.join(WORD_DIR, "zh", "pdac_biosensor_report_zh.docx")
    ]
    for path in original_files:
        make_backup(path)
        
    print("Writing English LaTeX files...")
    generate_latex_en(is_revised=False)
    generate_latex_en(is_revised=True)
    
    print("Writing Chinese LaTeX files...")
    generate_latex_zh(is_revised=False)
    generate_latex_zh(is_revised=True)
    
    print("Writing English Word documents...")
    generate_word_en(is_revised=False)
    generate_word_en(is_revised=True)
    
    print("Writing Chinese Word documents...")
    generate_word_zh(is_revised=False)
    generate_word_zh(is_revised=True)
    
    print("\nAll files written successfully!")

if __name__ == "__main__":
    main()
