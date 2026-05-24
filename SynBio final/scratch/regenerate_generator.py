import os
import marshal
import pprint

def main():
    pyc_path = '/Users/Janet/Documents/Antigravity/SynBio final/src/__pycache__/generate_reports.cpython-313.pyc'
    target_path = '/Users/Janet/Documents/Antigravity/SynBio final/src/generate_reports.py'
    
    # 1. Execute the pyc to get the clean dictionaries
    with open(pyc_path, 'rb') as f:
        f.read(16)
        code_obj = marshal.load(f)
        
    globals_dict = {}
    exec(code_obj, globals_dict)
    
    en_sections = globals_dict['en_sections']
    zh_sections = globals_dict['zh_sections']
    
    # Ensure CJK typo is fixed in the dictionary directly
    if 'QC' in zh_sections:
        zh_sections['QC'] = zh_sections['QC'].replace('蛋白質翻译', '蛋白質翻譯')
    for k in zh_sections:
        if isinstance(zh_sections[k], str):
            zh_sections[k] = zh_sections[k].replace('蛋白質翻译', '蛋白質翻譯')

    # Fix author list to the specified four authors if not already there
    en_sections['Author'] = 'Chen-Jung Shih, Hsuan-Yu Liao, De-Fang Su, Chia-Yi Lin'
    zh_sections['Author'] = '施貞蓉、廖軒佑、宿淂芳、林家誼'
    
    # Add SingleCell section text
    en_sections['SingleCell'] = (
        "Because the discovery analysis is based on bulk RNA-seq, it remains unclear whether UBE2S and CCR6 "
        "are co-expressed within the same malignant cell population or arise from distinct compartments of the "
        "tumor microenvironment. This distinction has direct consequences for circuit architecture. If both "
        "inputs are co-expressed in malignant epithelial cells, a cell-intrinsic transcriptional circuit may "
        "be feasible. If UBE2S is primarily expressed by cycling tumor cells while CCR6 is derived from immune "
        "cells, the pair should instead be interpreted as a tissue-level microenvironmental signature. Future "
        "single-cell RNA-seq and spatial transcriptomic analysis should therefore quantify cell-type-specific "
        "expression, co-expression frequency, and spatial proximity of UBE2S-high and CCR6-high compartments."
    )
    zh_sections['SingleCell'] = (
        "由於本研究的發現分析建立於 bulk RNA-seq，尚無法判斷 UBE2S 與 CCR6 是否共同表現於同一群惡性上皮細胞，"
        "或是分別來自腫瘤微環境中的不同細胞區室。這項區分會直接影響電路架構的設計：若兩個輸入在惡性上皮細胞中"
        "共同表現，則細胞內部轉錄電路較具可行性；若 UBE2S 主要來自增殖中的癌細胞，而 CCR6 主要反映免疫細胞"
        "或發炎微環境，則此組合應被解讀為組織層級的微環境特徵，而非單一細胞內部的 AND gate。後續應利用 "
        "single-cell RNA-seq 與 spatial transcriptomics 評估兩基因的細胞類型特異性、共同表現比例與空間鄰近關係。"
    )
    
    # 2. Write the file content
    with open(target_path, 'w', encoding='utf-8') as f:
        # Imports & Setup
        f.write('''import os
import sys
import subprocess
import pandas as pd
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
import shutil
import time

# Define directories
LATEX_DIR = "reports/latex"
WORD_DIR = "reports/word"
TABLES_DIR = "results/tables"
FIGURES_DIR = "reports/latex/shared/figures"

# Helper to load data
def load_table(name):
    path = os.path.join(TABLES_DIR, name)
    if os.path.exists(path):
        return pd.read_csv(path)
    return None

# Load results tables
and_sweep = load_table("and_gate_parameter_sweep.csv")
and_perf = load_table("and_gate_performance.csv")
cv_results = load_table("cross_validation_results.csv")
de_discovery = load_table("differential_expression_discovery.csv")
qc_summary = load_table("expression_qc_summary.csv")
ext_val = load_table("external_validation_final_pair.csv")
final_pair = load_table("final_candidate_pair.csv")
gene_scores = load_table("gene_pair_scores.csv")
model_perf = load_table("model_performance_summary.csv")
random_pair = load_table("random_pair_control.csv")
sensitivity = load_table("threshold_sensitivity.csv")
shap_imp = load_table("shap_feature_importance.csv")
shap_thresh = load_table("shap_threshold_candidates.csv")
top_candidates = load_table("top_tumor_high_candidates.csv")

# ----------------- TABLE PREPARATIONS -----------------
# Table 1: Dataset Summary
t1_headers = ["Dataset / Cohort", "Platform / Type", "Tumor Samples", "Normal Samples", "Total Samples"]
t1_data = [
    ["TCGA-PAAD", "RNA-seq (RSEM TPM)", "178", "0", "178"],
    ["GTEx Pancreas", "RNA-seq (RSEM TPM)", "0", "167", "167"],
    ["Discovery Cohort (Combined)", "RNA-seq (RSEM TPM)", "178", "167", "345"],
    ["GSE62452 (External Validation)", "Microarray (Affymetrix GPL6244)", "69", "61", "130"]
]

# Table 2: Top 10 DE Genes (sorted by specificity score)
t2_headers = ["Gene", "Mean Tumor (log2 TPM)", "Mean Normal (log2 TPM)", "log2FC", "FDR", "ROC-AUC", "Specificity Score"]
t2_data = []
if top_candidates is not None:
    df_sorted = top_candidates.sort_values(by=["specificity_score", "mean_pdac"], ascending=[False, False])
    for _, row in df_sorted.head(10).iterrows():
        gene = str(row["gene"])
        rec = de_discovery[de_discovery["gene"] == gene] if de_discovery is not None else []
        if len(rec) > 0:
            mean_p = f"{rec.iloc[0]['mean_pdac']:.3f}"
            mean_n = f"{rec.iloc[0]['mean_normal']:.3f}"
            log2fc = f"{rec.iloc[0]['log2fc']:.3f}"
            fdr = f"{rec.iloc[0]['fdr']:.2e}"
            auc = f"{rec.iloc[0]['auc']:.3f}"
            spec = f"{row['specificity_score']:.3f}"
            t2_data.append([gene, mean_p, mean_n, log2fc, fdr, auc, spec])

# Table 3: Model Performance
t3_headers = ["Classifier Model", "Mean 5-Fold CV AUC", "Test AUC", "Test Accuracy", "Sensitivity", "Specificity", "F1 Score"]
t3_data = []
if cv_results is not None and model_perf is not None:
    for _, row in model_perf.iterrows():
        model_name = str(row["Model"])
        cv_rec = cv_results[cv_results["Model"] == model_name]
        cv_auc = f"{cv_rec.iloc[0]['Mean_CV_AUC']:.3f}" if len(cv_rec) > 0 else "N/A"
        test_auc = f"{row['Test_AUC']:.3f}"
        test_acc = f"{row['Test_Accuracy']*100:.1f}%"
        sens = f"{row['Test_Recall']:.3f}"
        tn, fp = row['TN'], row['FP']
        spec = f"{tn / (tn + fp):.3f}" if (tn + fp) > 0 else "N/A"
        f1 = f"{row['Test_F1']:.3f}"
        t3_data.append([model_name.replace("_", " "), cv_auc, test_auc, test_acc, sens, spec, f1])

# Table 4: Top 10 SHAP Features
t4_headers = ["Rank", "Gene Symbol", "Mean Absolute SHAP Value", "Inferred Threshold (log2 TPM)", "95% Confidence Interval"]
t4_data = []
if shap_imp is not None and shap_thresh is not None:
    for rank, row in enumerate(shap_imp.head(10).iterrows(), 1):
        gene = str(row[1]["gene"])
        shap_val = f"{row[1]['mean_abs_shap']:.4f}"
        thresh_rec = shap_thresh[shap_thresh["gene"] == gene]
        if len(thresh_rec) > 0:
            thresh = f"{thresh_rec.iloc[0]['inferred_threshold']:.3f}"
            ci = f"[{thresh_rec.iloc[0]['ci_95_lower']:.3f}, {thresh_rec.iloc[0]['ci_95_upper']:.3f}]"
        else:
            thresh, ci = "N/A", "N/A"
        t4_data.append([str(rank), gene, shap_val, thresh, ci])

# Table 5: Selected Candidate Pair
t5_headers = ["Property", "Gene A (UBE2S)", "Gene B (CCR6)", "Composite AND Gate"]
t5_data = []
if final_pair is not None and and_perf is not None:
    r = final_pair.iloc[0]
    ap = and_perf.iloc[0]
    t5_data = [
        ["Individual ROC-AUC", f"{r['gene_A_auc']:.4f}", f"{r['gene_B_auc']:.4f}", f"{ap['ROC_AUC']:.4f}"],
        ["log2 Fold Change (log2FC)", f"{r['gene_A_log2fc']:.3f}", f"{r['gene_B_log2fc']:.3f}", "—"],
        ["Pairwise Spearman Correlation (r_s)", f"{r['correlation']:.4f}", "—", "—"],
        ["Inferred Activation Threshold (K)", f"{ap['K_A']:.4f}", f"{ap['K_B']:.4f}", "—"],
        ["Tumor AND Activation Rate", f"{r['tumor_AND_activation']*100:.1f}%", "—", f"{r['tumor_AND_activation']*100:.1f}%"],
        ["Normal AND Activation Rate", f"{r['normal_AND_activation']*100:.1f}%", "—", f"{r['normal_AND_activation']*100:.1f}%"]
    ]

# Table 6: AND Gate Performance Metrics (Discovery Cohort)
t6_headers = ["Metric", "Formula / Definition", "Value"]
t6_data = []
if and_perf is not None:
    ap = and_perf.iloc[0]
    t6_data = [
        ["ROC-AUC", "Area Under the ROC Curve", f"{ap['ROC_AUC']:.4f}"],
        ["Accuracy", "(TP + TN) / Total", f"{ap['Accuracy']*100:.2f}%"],
        ["Sensitivity", "TP / (TP + FN)", f"{ap['Sensitivity']*100:.2f}%"],
        ["Specificity", "TN / (TN + FP)", f"{ap['Specificity']*100:.2f}%"],
        ["Decision Threshold", "Hill output threshold for ON state", f"{ap['Optimal_Threshold']:.2f}"]
    ]

# Table 7: External Validation Results (GSE62452)
t7_headers = ["Metric", "Value in Discovery (TCGA+GTEx)", "Value in Validation (GSE62452)", "Interpretation"]
t7_data = []
if ext_val is not None and and_perf is not None:
    ap = and_perf.iloc[0]
    ev = ext_val.iloc[0]
    t7_data = [
        ["Sample Size (N)", "345", f"{ev['sample_size']}", "Microarray vs RNA-seq"],
        ["ROC-AUC", f"{ap['ROC_AUC']:.4f}", f"{ev['ROC_AUC']:.4f}", "Reduced discrimination power"],
        ["Accuracy", f"{ap['Accuracy']*100:.1f}%", f"{ev['Accuracy']*100:.1f}%", "Impacted by platform shift"],
        ["Sensitivity", f"{ap['Sensitivity']*100:.1f}%", f"{ev['Sensitivity']*100:.1f}%", "Significant drop (platform scale mismatch)"],
        ["Specificity", f"{ap['Specificity']*100:.1f}%", f"{ev['Specificity']*100:.1f}%", "Highly conserved (retains safety profile)"]
    ]

# Table 8: Threshold Sensitivity (Subset of Perturbations)
t8_headers = ["UBE2S Threshold (K_A)", "CCR6 Threshold (K_B)", "ROC-AUC", "Accuracy", "Sensitivity", "Specificity"]
t8_data = []
if sensitivity is not None:
    # Baseline
    baseline_rec = and_perf.iloc[0]
    t8_data.append([
        "0.760 (Baseline)", "0.464 (Baseline)",
        f"{baseline_rec['ROC_AUC']:.4f}", f"{baseline_rec['Accuracy']*100:.1f}%",
        f"{baseline_rec['Sensitivity']*100:.1f}%", f"{baseline_rec['Specificity']*100:.1f}%"
    ])
    for _, row in sensitivity.iterrows():
        p_A = row["Perturbation_A"]
        p_B = row["Perturbation_B"]
        if p_A in [-0.5, -0.1, 0.1, 0.5] and p_B in [-0.5, -0.1, 0.1, 0.5] and p_A == p_B:
            k_A_val = f"{0.760 * (1 + p_A):.3f} ({p_A*100:+.0f}%)"
            k_B_val = f"{0.464 * (1 + p_B):.3f} ({p_B*100:+.0f}%)"
            auc = f"{row['ROC_AUC']:.4f}"
            acc = f"{row['Accuracy']*100:.1f}%"
            t8_data.append([k_A_val, k_B_val, auc, acc, "N/A", "N/A"])
''')
        
        # Write en_sections
        f.write('\nen_sections = ')
        pprint.pprint(en_sections, stream=f, width=120)
        
        # Write zh_sections
        f.write('\nzh_sections = ')
        pprint.pprint(zh_sections, stream=f, width=120)
        
        # Append updated generators and helper functions
        f.write('''
# ----------------- LATEX GENERATORS -----------------

def escape_text_for_latex(text):
    import re
    pattern = r"(\\$.*?\\$|\\\\\\[.*?\\\\\\])"
    tokens = re.split(pattern, text, flags=re.DOTALL)
    escaped_tokens = []
    for i, token in enumerate(tokens):
        if i % 2 == 0:
            escaped_token = token.replace("_", "\\\\_").replace("%", "\\\\%")
            escaped_tokens.append(escaped_token)
        else:
            escaped_tokens.append(token)
    return "".join(escaped_tokens)

def clean_math_for_word(text):
    if not isinstance(text, str):
        return text
    import re
    # Replace display equations \\\\[ ... \\\\] with simpler text representation
    def repl_display(m):
        eq = m.group(1).strip()
        eq = eq.replace(r"\\text{basal}", "basal")
        eq = eq.replace(r"\\max", "max")
        eq = eq.replace(r"\\left(", "(")
        eq = eq.replace(r"\\right)", ")")
        eq = re.sub(r"\\\\frac\\{([^{}]+)\\}\\{([^{}]+)\\}", r"(\\1/\\2)", eq)
        eq = eq.replace("{", "").replace("}", "")
        eq = eq.replace("_", "")
        return f"\\n{eq}\\n"
    
    text = re.sub(r"\\\\\\\\[(.*?)\\\\\\\\]", repl_display, text, flags=re.DOTALL)
    
    # Replace inline equations $ ... $ with simplified text
    def repl_inline(m):
        eq = m.group(1).strip()
        eq = eq.replace(r"\\text{basal}", "basal")
        eq = eq.replace(r"\\max", "max")
        eq = eq.replace(r"\\pm", "±")
        eq = eq.replace(r"\\in", "in")
        eq = eq.replace(r"\\le", "≤")
        eq = eq.replace(r"\\ge", "≥")
        eq = eq.replace(r"\\langle", "<")
        eq = eq.replace(r"\\rangle", ">")
        eq = eq.replace(r"\\infty", "∞")
        eq = eq.replace(r"\\{", "{").replace(r"\\}", "}")
        eq = eq.replace(r"\\\\", "")
        eq = re.sub(r"\\\\[a-zA-Z]+\\{([^{}]+)\\}", r"\\1", eq)
        eq = eq.replace("\\\\", "")
        return eq
        
    text = re.sub(r"\\$(.*?)\\$", repl_inline, text)
    
    text = text.replace(r"\\pm", "±")
    text = text.replace(r"\\\\&", "&")
    text = text.replace(r"\\\\%", "%")
    text = text.replace(r"\\\\_", "_")
    return text

def generate_latex_en(is_revised=False):
    filename = "pdac_biosensor_report_en_revised.tex" if is_revised else "pdac_biosensor_report_en.tex"
    filepath = os.path.join(LATEX_DIR, "en", filename)
    
    en_sections_esc = {k: (escape_text_for_latex(v) if isinstance(v, str) else v) for k, v in globals()['en_sections'].items()}
    en_sections = en_sections_esc

    def make_latex_table(headers, data, caption, label, spec=None, span_columns=False):
        actual_span = span_columns and is_revised
        if spec is None:
            spec = "l" + "c"*(len(headers)-1)
        env = "table*" if actual_span else "table"
        pos = "[t]" if actual_span else "[H]"
        latex = f"\\\\begin{{{env}}}{pos}\\n\\\\centering\\n"
        latex += f"\\\\caption{{{caption}}}\\\\label{{{label}}}\\n"
        latex += f"\\\\begin{{tabular}}{{{spec}}}\\n\\\\toprule\\n"
        escaped_headers = [h.replace("%", "\\\\%").replace("_", "\\\\_") for h in headers]
        latex += " & ".join([f"\\\\textbf{{{h}}}" for h in escaped_headers]) + " \\\\\\\\\\n\\\\midrule\\n"
        for row in data:
            escaped_row = [str(x).replace("%", "\\\\%").replace("_", "\\\\_").replace("&", "\\\\&").replace("+-", "$\\\\pm$") for x in row]
            latex += " & ".join(escaped_row) + " \\\\\\\\\\n"
        latex += f"\\\\bottomrule\\n\\\\end{{tabular}}\\n\\\\end{{{env}}}\\n"
        return latex

    doc_class = "\\\\documentclass[10pt, a4paper, twocolumn]{article}" if is_revised else "\\\\documentclass[12pt, a4paper]{article}"
    style_file = "report_style_revised.tex" if is_revised else "report_style.tex"
    fig_width = "0.95\\\\linewidth" if is_revised else "0.7\\\\textwidth"
    
    if is_revised:
        title_block = f\"\"\"
\\\\twocolumn[
  \\\\begin{{@twocolumnfalse}}
    \\\\maketitle
    \\\\begin{{abstract}}
      {en_sections['Abstract']}
    \\\\end{{abstract}}
    \\\\vspace{{1.5em}}
  \\\\end{{@twocolumnfalse}}
]
\"\"\"
    else:
        title_block = f\"\"\"
\\\\maketitle
\\\\newpage

\\\\begin{{abstract}}
{en_sections['Abstract']}
\\\\end{{abstract}}
\\\\newpage
\"\"\"

    content = f\"\"\"% Auto-generated English report for PDAC Biosensor Project
{doc_class}

\\\\input{{../shared/{style_file}}}
\\\\input{{../shared/macros.tex}}

\\\\setmainfont{{Times New Roman}}
\\\\setsansfont{{Arial}}
\\\\setmonofont{{Courier New}}

\\\\usepackage[style=apa,backend=biber]{{biblatex}}
\\\\addbibresource{{references_en.bib}}

\\\\title{{{en_sections['Title']}}}
\\\\author[1]{{{en_sections['Author']}}}
\\\\affil[1]{{{en_sections['Affiliation']}}}
\\\\date{{{en_sections['Date']}}}

\\\\begin{{document}}

{title_block}
\\\\tableofcontents
\\\\newpage

\\\\section{{Introduction}}
{en_sections['Introduction']}

\\\\section{{Scientific Rationale and Unmet Need}}
{en_sections['Rationale']}

\\\\section{{Data Sources}}
{en_sections['DataSources']}

{make_latex_table(t1_headers, t1_data, "Data Cohorts and Sample Size Distribution", "tab:datasets", span_columns=False)}

\\\\section{{Computational Pipeline}}
{en_sections['Pipeline']}

\\\\section{{Quality Control and Batch-Effect Assessment}}
{en_sections['QC']}

\\\\section{{Differential Expression Analysis}}
{en_sections['DE']}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{volcano_discovery.png}}
\\\\caption{{Volcano plot highlighting significantly differentially expressed genes in the discovery cohort (TCGA-PAAD vs GTEx Normal Pancreas). UBE2S and CCR6 are annotated as significant upregulated candidates.}}
\\\\label{{fig:volcano}}
\\\\end{{figure}}

{make_latex_table(t2_headers, t2_data, "Top 10 Differentially Expressed Genes Sorted by Specificity Score", "tab:top_de", span_columns=True)}

\\\\section{{Machine Learning Classifier Performance}}
{en_sections['ML']}

{make_latex_table(t3_headers, t3_data, "Machine Learning Classifier Performance and Cross-Validation Summary", "tab:ml_perf", span_columns=True)}

\\\\section{{SHAP-Based Explainable AI Analysis}}
{en_sections['SHAP']}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{shap_summary.png}}
\\\\caption{{SHAP summary plot showing feature importances for the top-ranked genes driving the L1 Logistic Regression classifier.}}
\\\\label{{fig:shap_summary}}
\\\\end{{figure}}

{make_latex_table(t4_headers, t4_data, "Top 10 SHAP Feature Importance and Inferred Expression Thresholds", "tab:shap_thresh", span_columns=True)}

\\\\section{{Candidate Gene Pair Selection}}
{en_sections['PairSelection']}

{make_latex_table(t5_headers, t5_data, "Detailed Molecular Profile of Selected Candidate Pair", "tab:candidate_pair", span_columns=True)}

\\\\section{{Orthogonality Assessment}}
{en_sections['Orthogonality']}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{gene_pair_scatter_final.png}}
\\\\caption{{Scatter plot of UBE2S vs CCR6 rescaled expression in discovery cohort, demonstrating decision boundary quadrants and sample clustering.}}
\\\\label{{fig:scatter}}
\\\\end{{figure}}

\\\\section{{Need for Single-Cell and Spatial Validation}}
{en_sections['SingleCell']}

\\\\section{{Hill-Equation-Based AND Gate Modeling}}
{en_sections['HillModeling']}
Here, we present the formal model parameters optimized to capture this logic.

\\\\section{{In Silico Validation}}
{en_sections['InSilico']}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{and_gate_heatmap_final.png}}
\\\\caption{{2D Contour heatmap demonstrating simulation output of UBE2S AND CCR6 logical AND gate, with rescaled expression and tumor/normal sample overlays.}}
\\\\label{{fig:heatmap}}
\\\\end{{figure}}

{make_latex_table(t6_headers, t6_data, "AND Gate Simulation Performance Summary in Discovery Cohort", "tab:and_perf", span_columns=False)}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{roc_curves.png}}
\\\\caption{{ROC curves comparing individual inputs (UBE2S, CCR6) against the combined logic gate output.}}
\\\\label{{fig:roc}}
\\\\end{{figure}}

{make_latex_table(t7_headers, t7_data, "External Validation Results on GSE62452 Dataset", "tab:ext_val", span_columns=True)}

\\\\section{{Robustness and Negative Controls}}
{en_sections['Robustness']}

\\\\section{{Limitations}}
\\\\begin{{enumerate}}
\"\"\"
    limitations_list = [
        "This is an in silico proof-of-concept; biochemical kinetics may differ.",
        "SHAP-inferred thresholds are statistical inflection points and do not map directly to biochemical dissociation constants.",
        "Bulk RNA-seq data reflects average cell populations and is highly influenced by stromal density and immune infiltration.",
        "The comparison between TCGA (tumor) and GTEx (normal) may contain subtle batch effects despite TOIL harmonization.",
        "External validation cohort demonstrated extremely low sensitivity (4.3%), indicating significant challenge in threshold transfer.",
        "The selected candidate pair (UBE2S + CCR6) is not strictly statistically orthogonal, exhibiting a correlation of 0.714.",
        "Transcriptomic abundance differences do not guarantee equivalent sensor accessibility or protein translation.",
        "The final candidates require promoter engineering or RNA sensor design, which introduces additional complexity.",
        "Any diagnostic or therapeutic application requires extensive wet-lab validation and safety testing in model organisms."
    ]
    for lim in limitations_list:
        content += f"\\\\item {lim}\\n"
        
    supp_prefix = "\\\\newpage\\n\\\\onecolumn\\n" if is_revised else "\\\\newpage\\n"
    
    content += f\"\"\"\\\\end{{enumerate}}

\\\\section{{Proposed Wet-Lab Validation}}
{en_sections['WetLab']}

\\\\section{{Conclusion}}
{en_sections['Conclusion']}

\\\\newpage
\\\\printbibliography[title={{References}}]

{supp_prefix}\\\\section{{Supplementary Tables and Figures}}
{en_sections['Supplementary']}

\\\\subsection{{Supplementary Table 1: Parameter Sweep}}
{make_latex_table(["Hill Coefficient (n)", "Basal Leakiness (P_basal)", "ROC-AUC", "Accuracy", "Sensitivity", "Specificity"], 
                  [[str(x[0]), str(x[1]), f"{x[2]:.4f}", f"{x[4]*100:.1f}\\\\%", f"{x[5]*100:.1f}\\\\%", f"{x[6]*100:.1f}\\\\%"] for x in and_sweep.values[:6]] if and_sweep is not None else [], 
                  "Hill Equation Grid Search Parameter Sweep Performance", "tab:supp_sweep", span_columns=True)}

\\\\subsection{{Supplementary Table 2: Threshold Sensitivity}}
{make_latex_table(t8_headers, t8_data, "Sensitivity Analysis of K Parameter Perturbations", "tab:supp_sensitivity", span_columns=True)}

\\\\subsection{{Supplementary Figures}}
\\\\begin{{figure}}[H]
\\\\centering
\\\\begin{{subfigure}}[b]{{0.48\\\\textwidth}}
\\\\centering
\\\\includegraphics[width=\\\\textwidth]{{shap_dependence_top_genes/UBE2S_shap_dependence.png}}
\\\\caption{{UBE2S SHAP dependence}}
\\\\end{{subfigure}}
\\\\hfill
\\\\begin{{subfigure}}[b]{{0.48\\\\textwidth}}
\\\\centering
\\\\includegraphics[width=\\\\textwidth]{{shap_dependence_top_genes/CCR6_shap_dependence.png}}
\\\\caption{{CCR6 SHAP dependence}}
\\\\end{{subfigure}}
\\\\caption{{SHAP dependence plots for selected candidate genes UBE2S and CCR6 showing threshold transition inflection points.}}
\\\\label{{fig:supp_shap_dependence}}
\\\\end{{figure}}

\\\\end{{document}}
\"\"\"
    with open(filepath, "w") as f:
        f.write(content)
    print(f"Generated {filepath}")

def generate_latex_zh(is_revised=False):
    filename = "pdac_biosensor_report_zh_revised.tex" if is_revised else "pdac_biosensor_report_zh.tex"
    filepath = os.path.join(LATEX_DIR, "zh", filename)
    
    zh_sections_esc = {k: (escape_text_for_latex(v) if isinstance(v, str) else v) for k, v in globals()['zh_sections'].items()}
    zh_sections = zh_sections_esc

    def make_latex_table(headers, data, caption, label, spec=None, span_columns=False):
        actual_span = span_columns and is_revised
        if spec is None:
            spec = "l" + "c"*(len(headers)-1)
        env = "table*" if actual_span else "table"
        pos = "[t]" if actual_span else "[H]"
        latex = f"\\\\begin{{{env}}}{pos}\\n\\\\centering\\n"
        latex += f"\\\\caption{{{caption}}}\\\\label{{{label}}}\\n"
        latex += f"\\\\begin{{tabular}}{{{spec}}}\\n\\\\toprule\\n"
        escaped_headers = [h.replace("%", "\\\\%").replace("_", "\\\\_") for h in headers]
        latex += " & ".join([f"\\\\textbf{{{h}}}" for h in escaped_headers]) + " \\\\\\\\\\n\\\\midrule\\n"
        for row in data:
            escaped_row = [str(x).replace("%", "\\\\%").replace("_", "\\\\_").replace("&", "\\\\&").replace("+-", "$\\\\pm$") for x in row]
            latex += " & ".join(escaped_row) + " \\\\\\\\\\n"
        latex += f"\\\\bottomrule\\n\\\\end{{tabular}}\\n\\\\end{{{env}}}\\n"
        return latex

    doc_class = "\\\\documentclass[10pt, a4paper, twocolumn]{article}" if is_revised else "\\\\documentclass[12pt, a4paper]{article}"
    style_file = "report_style_revised.tex" if is_revised else "report_style.tex"
    fig_width = "0.95\\\\linewidth" if is_revised else "0.7\\\\textwidth"

    if is_revised:
        title_block = f\"\"\"
\\\\twocolumn[
  \\\\begin{{@twocolumnfalse}}
    \\\\maketitle
    \\\\begin{{abstract}}
      {zh_sections['Abstract']}
    \\\\end{{abstract}}
    \\\\vspace{{1.5em}}
  \\\\end{{@twocolumnfalse}}
]
\"\"\"
    else:
        title_block = f\"\"\"
\\\\maketitle
\\\\newpage

\\\\begin{{abstract}}
{zh_sections['Abstract']}
\\\\end{{abstract}}
\\\\newpage
\"\"\"

    content = f\"\"\"% Auto-generated Chinese report for PDAC Biosensor Project
{doc_class}

\\\\input{{../shared/{style_file}}}
\\\\input{{../shared/macros.tex}}
\\\\usepackage{{xeCJK}}
\\\\setCJKmainfont{{DFKai-SB}}
\\\\setCJKsansfont{{DFKai-SB}}

\\\\setmainfont{{Times New Roman}}
\\\\setsansfont{{Arial}}
\\\\setmonofont{{Courier New}}

\\\\usepackage[style=apa,backend=biber]{{biblatex}}
\\\\addbibresource{{references_zh.bib}}

\\\\title{{\\\\Large \\\\bfseries {zh_sections['Title']} \\\\\\\\ \\\\vspace{{0.5em}} \\\\large {zh_sections['Subtitle']}}}
\\\\author{{{zh_sections['Author']}}}
\\\\affil{{{zh_sections['Affiliation']}}}
\\\\date{{{zh_sections['Date']}}}

\\\\begin{{document}}

{title_block}
\\\\tableofcontents
\\\\newpage

\\\\section{{前言}}
{zh_sections['Introduction']}

\\\\section{{科學背景與未滿足之臨床需求}}
{zh_sections['Rationale']}

\\\\section{{資料來源}}
{zh_sections['DataSources']}

{make_latex_table(t1_headers, t1_data, "數據世代與樣本量分布", "tab:datasets_zh", span_columns=False)}

\\\\section{{運算分析管線}}
{zh_sections['Pipeline']}

\\\\section{{品質控制與批次效應評估}}
{zh_sections['QC']}

\\\\section{{差異表現分析}}
{zh_sections['DE']}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{volcano_discovery.png}}
\\\\caption{{發現世代中的火山圖，標記顯著差異表現基因。其中 UBE2S 與 CCR6 被註記為顯著上調之候選基因。}}
\\\\label{{fig:volcano_zh}}
\\\\end{{figure}}

{make_latex_table(t2_headers, t2_data, "前 10 個差異表現基因 (按特異性得分排序)", "tab:top_de_zh", span_columns=True)}

\\\\section{{機器學習分類器表現}}
{zh_sections['ML']}

{make_latex_table(t3_headers, t3_data, "機器學習分類器表現與五折交叉驗證結果摘要", "tab:ml_perf_zh", span_columns=True)}

\\\\section{{基於 SHAP 的可解釋型人工智慧分析}}
{zh_sections['SHAP']}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{shap_summary.png}}
\\\\caption{{SHAP 摘要圖，顯示推動 L1 邏輯斯迴歸分類器決策的前 15 個特徵重要性排名。}}
\\\\label{{fig:shap_summary_zh}}
\\\\end{{figure}}

{make_latex_table(t4_headers, t4_data, "前 10 個 SHAP 特徵重要性與模型推估表達閾值", "tab:shap_thresh_zh", span_columns=True)}

\\\\section{{候選基因組合篩選}}
{zh_sections['PairSelection']}

{make_latex_table(t5_headers, t5_data, "選定候選基因組合的詳細分子譜描述", "tab:candidate_pair_zh", span_columns=True)}

\\\\section{{正交性評估}}
{zh_sections['Orthogonality']}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{gene_pair_scatter_final.png}}
\\\\caption{{UBE2S 與 CCR6 在發現世代中的表達量散佈圖，呈現決策邊界與樣本分布象限。}}
\\\\label{{fig:scatter_zh}}
\\\\end{{figure}}

\\\\section{{單細胞與空間轉錄體驗證之必要性}}
{zh_sections['SingleCell']}

\\\\section{{基於希爾方程式的 AND gate 建模}}
{zh_sections['HillModeling']}

\\\\section{{電腦模擬驗證}}
{zh_sections['InSilico']}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{and_gate_heatmap_final.png}}
\\\\caption{{UBE2S 與 CCR6 及閘模擬輸出之二維等高線熱圖，並疊加腫瘤與健康樣本。}}
\\\\label{{fig:heatmap_zh}}
\\\\end{{figure}}

{make_latex_table(t6_headers, t6_data, "及閘模擬在發現世代中的分類效能指標摘要", "tab:and_perf_zh", span_columns=False)}

\\\\begin{{figure}}[H]
\\\\centering
\\\\includegraphics[width={fig_width}]{{roc_curves.png}}
\\\\caption{{比較單一基因輸入 (UBE2S, CCR6) 與雙輸入邏輯及閘輸出的 ROC 曲線。}}
\\\\label{{fig:roc_zh}}
\\\\end{{figure}}

{make_latex_table(t7_headers, t7_data, "GSE62452 外部驗證結果摘要", "tab:ext_val_zh", span_columns=True)}

\\\\section{{穩健性分析與負控制}}
{zh_sections['Robustness']}

\\\\section{{研究限制}}
\\\\begin{{enumerate}}
\"\"\"
    limitations_list_zh = [
        "本研究僅為電腦模擬之概念驗證，實際生化反應之動力學可能有所不同。",
        "SHAP 推估之表達閾值為統計學之拐點，無法直接對應至物理學上的生化解離常數。",
        "組織轉錄體 (Bulk RNA-seq) 表達數據易受到細胞組成、腫瘤純度及免疫細胞浸潤之干擾。",
        "儘管經過 TOIL 管線標準化，TCGA (腫瘤) 與 GTEx (健康) 的比較可能仍存在部分批次效應。",
        "外部驗證世代中顯示極低的敏感度 (4.3%)，顯示跨平台定量尺度的不匹配是閾值轉移的重大挑戰。",
        "選定的候選基因對 (UBE2S + CCR6) 統計上並非嚴格正交，在 bulk RNA-seq 中的相關係數達 0.714。",
        "轉錄體層級之豐度差異不代表合成感測器之可及性，亦不保證有同等程度的蛋白質翻譯表現。",
        "將此候選組合轉譯為合成基因線路，仍需進行啟動子工程或 RNA 轉錄調節器設計，程序較為複雜。",
        "任何診斷或治療性之臨床應用，皆須經過細胞、動物與安全性的嚴格檢驗。"
    ]
    for lim in limitations_list_zh:
        content += f"\\\\item {lim}\\n"
        
    supp_prefix = "\\\\newpage\\n\\\\onecolumn\\n" if is_revised else "\\\\newpage\\n"

    content += f\"\"\"\\\\end{{enumerate}}

\\\\section{{後續濕實驗驗證規劃}}
{zh_sections['WetLab']}

\\\\section{{結論}}
{zh_sections['Conclusion']}

\\\\newpage
\\\\printbibliography[title={{參考文獻}}]

{supp_prefix}\\\\section{{補充表格與圖}}
{zh_sections['Supplementary']}

\\\\subsection{{補充表格 1：參數掃描結果}}
{make_latex_table(["Hill 係數 (n)", "基底洩漏量 (P_basal)", "ROC-AUC", "準確度", "敏感度", "特異度"], 
                  [[str(x[0]), str(x[1]), f"{x[2]:.4f}", f"{x[4]*100:.1f}\\\\%", f"{x[5]*100:.1f}\\\\%", f"{x[6]*100:.1f}\\\\%"] for x in and_sweep.values[:6]] if and_sweep is not None else [], 
                  "希爾方程式網格掃描參數表現結果", "tab:supp_sweep_zh", span_columns=True)}

\\\\subsection{{補充表格 2：K 參數微擾對分類效能影響之敏感度分析}}
{make_latex_table(t8_headers, t8_data, "K 參數微擾對分類效能影響之敏感度分析", "tab:supp_sensitivity_zh", span_columns=True)}

\\\\subsection{{補充圖}}
\\\\begin{{figure}}[H]
\\\\centering
\\\\begin{{subfigure}}[b]{{0.48\\\\textwidth}}
\\\\centering
\\\\includegraphics[width=\\\\textwidth]{{shap_dependence_top_genes/UBE2S_shap_dependence.png}}
\\\\caption{{UBE2S SHAP 依賴性}}
\\\\end{{subfigure}}
\\\\hfill
\\\\begin{{subfigure}}[b]{{0.48\\\\textwidth}}
\\\\centering
\\\\includegraphics[width=\\\\textwidth]{{shap_dependence_top_genes/CCR6_shap_dependence.png}}
\\\\caption{{CCR6 SHAP 依賴性}}
\\\\end{{subfigure}}
\\\\caption{{選定候選基因 UBE2S 與 CCR6 之 SHAP 依賴圖，呈現出拐點轉換閾值。}}
\\\\label{{fig:supp_shap_dependence_zh}}
\\\\end{{figure}}

\\\\end{{document}}
\"\"\"
    with open(filepath, "w") as f:
        f.write(content)
    print(f"Generated {filepath}")

def add_word_paragraph(doc, text, style='Normal', space_after=6, line_spacing=1.5, bold=False, italic=False, font_name='Times New Roman', font_size=12, is_chinese=False):
    text = clean_math_for_word(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if is_chinese:
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    return p

def add_word_heading(doc, text, level, is_chinese=False):
    text = clean_math_for_word(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    font_size = 18 if level == 1 else (14 if level == 2 else 12)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial' if not is_chinese else 'DFKai-SB'
    run.font.size = Pt(font_size)
    if is_chinese:
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    return p

def add_word_table(doc, headers, data, is_chinese=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        title_clean = clean_math_for_word(title)
        hdr_cells[i].text = title_clean
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.name = 'Arial' if not is_chinese else 'DFKai-SB'
        run.font.size = Pt(10)
        if is_chinese:
            run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
            
    for row in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            val_clean = clean_math_for_word(val)
            row_cells[i].text = str(val_clean)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Times New Roman' if not is_chinese else 'DFKai-SB'
                run.font.size = Pt(10)
                if is_chinese:
                    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def add_word_image(doc, filepath, caption, is_chinese=False, width=Inches(5.0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    if os.path.exists(filepath):
        p.add_run().add_picture(filepath, width=width)
    else:
        r = p.add_run(f"[Image placeholder: {filepath}]")
        r.italic = True
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(f"Figure: {clean_math_for_word(caption)}")
    run_cap.italic = True
    run_cap.font.name = 'Times New Roman' if not is_chinese else 'DFKai-SB'
    run_cap.font.size = Pt(10)
    if is_chinese:
        run_cap._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')

def set_section_columns(section, num_cols, col_space=None):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    if num_cols > 1:
        cols.set(qn('w:equalWidth'), '1')
        if col_space is not None:
            cols.set(qn('w:space'), str(col_space))

def generate_word_en(is_revised=False):
    filename = "pdac_biosensor_report_en_revised.docx" if is_revised else "pdac_biosensor_report_en.docx"
    filepath = os.path.join(WORD_DIR, "en", filename)
    doc = docx.Document()
    
    sec1 = doc.sections[0]
    sec1.top_margin = Inches(0.98)
    sec1.bottom_margin = Inches(0.98)
    sec1.left_margin = Inches(0.98)
    sec1.right_margin = Inches(0.98)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(80)
    p_title.paragraph_format.space_after = Pt(18)
    run_title = p_title.add_run(clean_math_for_word(en_sections["Title"]))
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(6)
    run_author = p_author.add_run(clean_math_for_word(en_sections["Author"]))
    run_author.font.name = 'Times New Roman'
    run_author.font.size = Pt(12)
    
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(18)
    run_aff = p_aff.add_run(clean_math_for_word(en_sections["Affiliation"]))
    run_aff.font.name = 'Times New Roman'
    run_aff.font.size = Pt(10)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(clean_math_for_word(en_sections["Date"]))
    run_date.font.name = 'Times New Roman'
    run_date.font.size = Pt(11)
    
    doc.add_page_break()
    
    add_word_heading(doc, "Abstract", 1)
    add_word_paragraph(doc, en_sections["Abstract"])
    
    if is_revised:
        sec2 = doc.add_section(start_type=WD_SECTION.NEW_PAGE)
        set_section_columns(sec2, 2, col_space=720)
        sec2.top_margin = Inches(0.7)
        sec2.bottom_margin = Inches(0.7)
        sec2.left_margin = Inches(0.7)
        sec2.right_margin = Inches(0.7)
        main_fig_width = Inches(3.0)
    else:
        doc.add_page_break()
        main_fig_width = Inches(5.0)
        
    sections_order = [
        ("1. Introduction", en_sections["Introduction"]),
        ("2. Scientific Rationale and Unmet Need", en_sections["Rationale"]),
        ("3. Data Sources", en_sections["DataSources"]),
        ("4. Computational Pipeline", en_sections["Pipeline"]),
        ("5. Quality Control and Batch-Effect Assessment", en_sections["QC"]),
        ("6. Differential Expression Analysis", en_sections["DE"]),
        ("7. Machine Learning Classifier Performance", en_sections["ML"]),
        ("8. SHAP-Based Explainable AI Analysis", en_sections["SHAP"]),
        ("9. Candidate Gene Pair Selection", en_sections["PairSelection"]),
        ("10. Orthogonality Assessment", en_sections["Orthogonality"]),
        ("11. Need for Single-Cell and Spatial Validation", en_sections["SingleCell"]),
        ("12. Hill-Equation-Based AND Gate Modeling", en_sections["HillModeling"]),
        ("13. In Silico Validation", en_sections["InSilico"]),
        ("14. Robustness and Controls", en_sections["Robustness"])
    ]
    
    for title, text in sections_order:
        add_word_heading(doc, title, 1)
        add_word_paragraph(doc, text)
        
        if "3. Data Sources" in title:
            add_word_table(doc, t1_headers, t1_data)
        elif "6. Differential Expression" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "volcano_discovery.png"), "Volcano plot of discovery cohort", width=main_fig_width)
            add_word_table(doc, t2_headers, t2_data)
        elif "7. Machine Learning" in title:
            add_word_table(doc, t3_headers, t3_data)
        elif "8. SHAP-Based" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "shap_summary.png"), "SHAP Feature Importance Summary Bar Plot", width=main_fig_width)
            add_word_table(doc, t4_headers, t4_data)
        elif "9. Candidate Gene" in title:
            add_word_table(doc, t5_headers, t5_data)
        elif "10. Orthogonality" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "gene_pair_scatter_final.png"), "UBE2S vs CCR6 rescaled expression scatter plot", width=main_fig_width)
        elif "13. In Silico Validation" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "and_gate_heatmap_final.png"), "AND gate logical activation surface heatmap", width=main_fig_width)
            add_word_table(doc, t6_headers, t6_data)
            add_word_image(doc, os.path.join(FIGURES_DIR, "roc_curves.png"), "ROC curves comparison: single vs logical combination", width=main_fig_width)
            add_word_heading(doc, "Table 7: External Validation Results (GSE62452)", 2)
            add_word_table(doc, t7_headers, t7_data)
            
    add_word_heading(doc, "15. Limitations", 1)
    limitations_list = [
        "This is an in silico proof-of-concept; biochemical kinetics may differ.",
        "SHAP-inferred thresholds are statistical inflection points and do not map directly to biochemical dissociation constants.",
        "Bulk RNA-seq data reflects average cell populations and is highly influenced by stromal density and immune infiltration.",
        "The comparison between TCGA (tumor) and GTEx (normal) may contain subtle batch effects despite TOIL harmonization.",
        "External validation cohort demonstrated extremely low sensitivity (4.3%), indicating significant challenge in threshold transfer.",
        "The selected candidate pair (UBE2S + CCR6) is not strictly statistically orthogonal, exhibiting a correlation of 0.714.",
        "Transcriptomic abundance differences do not guarantee equivalent sensor accessibility or protein translation.",
        "The final candidates require promoter engineering or RNA sensor design, which introduces additional complexity.",
        "Any diagnostic or therapeutic application requires extensive wet-lab validation and safety testing in model organisms."
    ]
    for lim in limitations_list:
        add_word_paragraph(doc, f"- {lim}", space_after=4)
        
    add_word_heading(doc, "16. Proposed Wet-Lab Validation", 1)
    add_word_paragraph(doc, en_sections["WetLab"])
    
    add_word_heading(doc, "17. Conclusion", 1)
    add_word_paragraph(doc, en_sections["Conclusion"])
    
    add_word_heading(doc, "18. References", 1)
    refs = en_sections["References"].split("\\n")
    for ref in refs:
        add_word_paragraph(doc, ref, space_after=4)
        
    if is_revised:
        sec3 = doc.add_section(start_type=WD_SECTION.NEW_PAGE)
        set_section_columns(sec3, 1)
        sec3.top_margin = Inches(0.98)
        sec3.bottom_margin = Inches(0.98)
        sec3.left_margin = Inches(0.98)
        sec3.right_margin = Inches(0.98)
    else:
        doc.add_page_break()
        
    add_word_heading(doc, "19. Supplementary Tables and Figures", 1)
    add_word_paragraph(doc, en_sections["Supplementary"])
    
    add_word_heading(doc, "Supplementary Table 1: Grid Search Sweep", 2)
    add_word_table(doc, ["Hill Coefficient (n)", "Basal Leakiness (P_basal)", "ROC-AUC", "Accuracy", "Sensitivity", "Specificity"],
                   [[str(x[0]), str(x[1]), f"{x[2]:.4f}", f"{x[4]*100:.1f}%", f"{x[5]*100:.1f}%", f"{x[6]*100:.1f}%"] for x in and_sweep.values[:6]] if and_sweep is not None else [])
    
    add_word_heading(doc, "Supplementary Table 2: Sensitivity analysis", 2)
    add_word_table(doc, t8_headers, t8_data)
    
    add_word_heading(doc, "Supplementary Figures: SHAP dependence plots", 2)
    add_word_image(doc, os.path.join(FIGURES_DIR, "shap_dependence_top_genes/UBE2S_shap_dependence.png"), "UBE2S SHAP dependence plot", width=Inches(4.5))
    add_word_image(doc, os.path.join(FIGURES_DIR, "shap_dependence_top_genes/CCR6_shap_dependence.png"), "CCR6 SHAP dependence plot", width=Inches(4.5))
    
    doc.save(filepath)
    print(f"Generated {filepath}")

def generate_word_zh(is_revised=False):
    filename = "pdac_biosensor_report_zh_revised.docx" if is_revised else "pdac_biosensor_report_zh.docx"
    filepath = os.path.join(WORD_DIR, "zh", filename)
    doc = docx.Document()
    
    sec1 = doc.sections[0]
    sec1.top_margin = Inches(0.98)
    sec1.bottom_margin = Inches(0.98)
    sec1.left_margin = Inches(0.98)
    sec1.right_margin = Inches(0.98)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(80)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(clean_math_for_word(zh_sections["Title"]))
    run_title.bold = True
    run_title.font.name = 'DFKai-SB'
    run_title.font.size = Pt(20)
    run_title._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run(clean_math_for_word(zh_sections["Subtitle"]))
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(12)
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(6)
    run_author = p_author.add_run(clean_math_for_word(zh_sections["Author"]))
    run_author.font.name = 'DFKai-SB'
    run_author.font.size = Pt(12)
    run_author._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(18)
    run_aff = p_aff.add_run(clean_math_for_word(zh_sections["Affiliation"]))
    run_aff.font.name = 'DFKai-SB'
    run_aff.font.size = Pt(10)
    run_aff._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(clean_math_for_word(zh_sections["Date"]))
    run_date.font.name = 'DFKai-SB'
    run_date.font.size = Pt(11)
    run_date._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'DFKai-SB')
    
    doc.add_page_break()
    
    add_word_heading(doc, "摘要", 1, is_chinese=True)
    add_word_paragraph(doc, zh_sections["Abstract"], is_chinese=True, font_name='DFKai-SB')
    
    if is_revised:
        sec2 = doc.add_section(start_type=WD_SECTION.NEW_PAGE)
        set_section_columns(sec2, 2, col_space=720)
        sec2.top_margin = Inches(0.7)
        sec2.bottom_margin = Inches(0.7)
        sec2.left_margin = Inches(0.7)
        sec2.right_margin = Inches(0.7)
        main_fig_width = Inches(3.0)
    else:
        doc.add_page_break()
        main_fig_width = Inches(5.0)
        
    sections_order_zh = [
        ("一、前言", zh_sections["Introduction"]),
        ("二、科學背景與未滿足之臨床需求", zh_sections["Rationale"]),
        ("三、資料來源", zh_sections["DataSources"]),
        ("四、運算分析管線", zh_sections["Pipeline"]),
        ("五、品質控制與批次效應評估", zh_sections["QC"]),
        ("六、差異表現分析", zh_sections["DE"]),
        ("七、機器學習分類器表現", zh_sections["ML"]),
        ("八、基於 SHAP 的可解釋型人工智慧分析", zh_sections["SHAP"]),
        ("九、候選基因組合篩選", zh_sections["PairSelection"]),
        ("十、正交性評估", zh_sections["Orthogonality"]),
        ("十一、單細胞與空間轉錄體驗證之必要性", zh_sections["SingleCell"]),
        ("十二、基於希爾方程式的 AND gate 建模", zh_sections["HillModeling"]),
        ("十三、電腦模擬驗證", zh_sections["InSilico"]),
        ("十四、穩健性分析與負控制", zh_sections["Robustness"])
    ]
    
    for title, text in sections_order_zh:
        add_word_heading(doc, title, 1, is_chinese=True)
        add_word_paragraph(doc, text, is_chinese=True, font_name='DFKai-SB')
        
        if "三、資料來源" in title:
            add_word_table(doc, t1_headers, t1_data, is_chinese=True)
        elif "六、差異表現分析" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "volcano_discovery.png"), "發現世代之火山圖", is_chinese=True, width=main_fig_width)
            add_word_table(doc, t2_headers, t2_data, is_chinese=True)
        elif "七、機器學習" in title:
            add_word_table(doc, t3_headers, t3_data, is_chinese=True)
        elif "八、基於 SHAP" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "shap_summary.png"), "SHAP 特徵重要性摘要條形圖", is_chinese=True, width=main_fig_width)
            add_word_table(doc, t4_headers, t4_data, is_chinese=True)
        elif "九、候選基因" in title:
            add_word_table(doc, t5_headers, t5_data, is_chinese=True)
        elif "十、正交性評估" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "gene_pair_scatter_final.png"), "UBE2S 與 CCR6 表達量散佈圖", is_chinese=True, width=main_fig_width)
        elif "十三、電腦模擬驗證" in title:
            add_word_image(doc, os.path.join(FIGURES_DIR, "and_gate_heatmap_final.png"), "邏輯及閘活化熱圖", is_chinese=True, width=main_fig_width)
            add_word_table(doc, t6_headers, t6_data, is_chinese=True)
            add_word_image(doc, os.path.join(FIGURES_DIR, "roc_curves.png"), "ROC 曲線比較圖", is_chinese=True, width=main_fig_width)
            add_word_heading(doc, "表格 7：外部驗證結果 (GSE62452)", 2, is_chinese=True)
            add_word_table(doc, t7_headers, t7_data, is_chinese=True)
            
    add_word_heading(doc, "十五、研究限制", 1, is_chinese=True)
    limitations_list_zh = [
        "本研究僅為電腦模擬之概念驗證，實際生化反應之動力學可能有所不同。",
        "SHAP 推估之表達閾值為統計學之拐點，無法直接對應至物理學上的生化解離常數。",
        "組織轉錄體 (Bulk RNA-seq) 表達數據易受到細胞組成、腫瘤純度及免疫細胞浸潤之干擾。",
        "儘長經過 TOIL 管線標準化，TCGA (腫瘤) 與 GTEx (健康) 的比較可能仍存在部分批次效應。",
        "外部驗證世代中顯示極低的敏感度 (4.3%)，顯示跨平台定量尺度的不匹配是閾值轉移的重大挑戰。",
        "選定的候選基因對 (UBE2S + CCR6) 統計上並非嚴格正交，在 bulk RNA-seq 中的相關係數達 0.714。",
        "轉錄體層級之豐度差異不代表合成感測器之可及性，亦不保證有同等程度的蛋白質翻譯表現。",
        "將此候選組合轉譯為合成基因線路，仍需進行啟動子工程或 RNA 轉錄調節器設計，程序較為複雜。",
        "任何診斷或治療性之臨床應用，皆須經過細胞、動物與安全性的嚴格檢驗。"
    ]
    for lim in limitations_list_zh:
        add_word_paragraph(doc, f"- {lim}", is_chinese=True, font_name='DFKai-SB', space_after=4)
        
    add_word_heading(doc, "十六、後續濕實驗驗證規劃", 1, is_chinese=True)
    add_word_paragraph(doc, zh_sections["WetLab"], is_chinese=True, font_name='DFKai-SB')
    
    add_word_heading(doc, "十七、結論", 1, is_chinese=True)
    add_word_paragraph(doc, zh_sections["Conclusion"], is_chinese=True, font_name='DFKai-SB')
    
    add_word_heading(doc, "十八、參考文獻", 1, is_chinese=True)
    refs = zh_sections["References"].split("\\n")
    for ref in refs:
        add_word_paragraph(doc, ref, is_chinese=True, font_name='DFKai-SB', space_after=4)
        
    if is_revised:
        sec3 = doc.add_section(start_type=WD_SECTION.NEW_PAGE)
        set_section_columns(sec3, 1)
        sec3.top_margin = Inches(0.98)
        sec3.bottom_margin = Inches(0.98)
        sec3.left_margin = Inches(0.98)
        sec3.right_margin = Inches(0.98)
    else:
        doc.add_page_break()
        
    add_word_heading(doc, "十九、補充表格與圖", 1, is_chinese=True)
    add_word_paragraph(doc, zh_sections["Supplementary"], is_chinese=True, font_name='DFKai-SB')
    
    add_word_heading(doc, "補充表格 1：希爾方程式網格掃描參數表現結果", 2, is_chinese=True)
    add_word_table(doc, ["Hill 係數 (n)", "基底洩漏量 (P_basal)", "ROC-AUC", "準確度", "敏感度", "特異度"],
                   [[str(x[0]), str(x[1]), f"{x[2]:.4f}", f"{x[4]*100:.1f}%", f"{x[5]*100:.1f}%", f"{x[6]*100:.1f}%"] for x in and_sweep.values[:6]] if and_sweep is not None else [], is_chinese=True)
    
    add_word_heading(doc, "補充表格 2：K 參數微擾對分類效能影響之敏感度分析", 2, is_chinese=True)
    add_word_table(doc, t8_headers, t8_data, is_chinese=True)
    
    add_word_heading(doc, "補充圖：SHAP 依賴圖", 2, is_chinese=True)
    add_word_image(doc, os.path.join(FIGURES_DIR, "shap_dependence_top_genes/UBE2S_shap_dependence.png"), "UBE2S SHAP 依賴圖", is_chinese=True, width=Inches(4.5))
    add_word_image(doc, os.path.join(FIGURES_DIR, "shap_dependence_top_genes/CCR6_shap_dependence.png"), "CCR6 SHAP 依賴圖", is_chinese=True, width=Inches(4.5))
    
    doc.save(filepath)
    print(f"Generated {filepath}")

import shutil
import time

def make_backup(path):
    if os.path.exists(path):
        ts = time.strftime("%Y%m%d_%H%M%S")
        backup_path = f"{path}.backup_{ts}"
        shutil.copy2(path, backup_path)
        print(f"Created backup of {path} at {backup_path}")

def main():
    original_files = [
        os.path.join(LATEX_DIR, "en", "pdac_biosensor_report_en.tex"),
        os.path.join(LATEX_DIR, "zh", "pdac_biosensor_report_zh.tex"),
        os.path.join(WORD_DIR, "en", "pdac_biosensor_report_en.docx"),
        os.path.join(WORD_DIR, "zh", "pdac_biosensor_report_zh.docx")
    ]
    for path in original_files:
        make_backup(path)
        
    print("Writing English LaTeX files...")
    generate_latex_en(is_revised=False)
    generate_latex_en(is_revised=True)
    
    print("Writing Chinese LaTeX files...")
    generate_latex_zh(is_revised=False)
    generate_latex_zh(is_revised=True)
    
    print("Writing English Word documents...")
    generate_word_en(is_revised=False)
    generate_word_en(is_revised=True)
    
    print("Writing Chinese Word documents...")
    generate_word_zh(is_revised=False)
    generate_word_zh(is_revised=True)
    
    print("\\nAll files written successfully!")

if __name__ == "__main__":
    main()
''')
    
    print("Reconstruction and save completed successfully!")

if __name__ == '__main__':
    main()
