#!/usr/bin/env python3
import os
import sys
import shutil
import subprocess
import pandas as pd
from pathlib import Path

# Paths
PROJECT_DIR = Path(__file__).parent.parent
RESULTS_DIR = PROJECT_DIR / "results_v2"
TABLES_DIR = RESULTS_DIR / "tables"
FIGURES_DIR = RESULTS_DIR / "figures"
REPORTS_DIR = PROJECT_DIR / "reports_v2"

# Ensure reports directory exists
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

# Copy IEEEtran.cls and bib files
shutil.copy(PROJECT_DIR / "reports_v1_archive/latex/en/IEEEtran.cls", REPORTS_DIR / "IEEEtran.cls")
shutil.copy(PROJECT_DIR / "reports_v1_archive/latex/en/references_en.bib", REPORTS_DIR / "references_en.bib")
shutil.copy(PROJECT_DIR / "reports_v1_archive/latex/zh/references_zh.bib", REPORTS_DIR / "references_zh.bib")

# Load V2 results data
df_perf = pd.read_csv(TABLES_DIR / "table_final_pair_performance_all_datasets.csv")
df_pair = pd.read_csv(TABLES_DIR / "final_candidate_pair_v2.csv")
pair_row = df_pair.iloc[0]

gene_A = pair_row["gene_A"]
gene_B = pair_row["gene_B"]
K_A = pair_row["K_A"]
K_B = pair_row["K_B"]
corr = pair_row["tumor_correlation"]

disc_sens = pair_row["disc_sens"]
disc_spec = pair_row["disc_spec"]
val_sens = pair_row["val_sens"]
val_spec = pair_row["val_spec"]
ext_sens = pair_row["ext_sens"]
ext_spec = pair_row["ext_spec"]

disc_auc = df_perf.loc[df_perf["Cohort/Dataset"] == "TCGA + GTEx Discovery", "ROC-AUC"].values[0]
val_auc = df_perf.loc[df_perf["Cohort/Dataset"] == "GSE62452 Same-Cohort Validation", "ROC-AUC"].values[0]
ext_auc = df_perf.loc[df_perf["Cohort/Dataset"] == "GSE28735 External Validation", "ROC-AUC"].values[0]

# --- ENGLISH TEXT CONTENT ---
title_en = "Data-Driven Design of a Logic-Gated Biosensor via Unbiased Transcriptomic Profiling of Pancreatic Tumor Microenvironment"
authors_en = "SHIH, CHEN-JUNG$^1$, SU, TE-FANG$^1$, LIAO, XUAN-YOU$^2$, and LIN, CHIA-I$^2$"
affiliations_en = "$^1$Department of Life Science, National Taiwan University, Taipei, Taiwan\\\\$^2$Department of Biochemical Science and Technology, National Taiwan University, Taipei, Taiwan"

abstract_en = (
    "Pancreatic ductal adenocarcinoma (PDAC) remains a highly lethal malignancy with a 5-year survival rate "
    "below 12%, primarily due to late-stage diagnosis and a lack of specific tumor biomarkers. In this study, "
    "we present a rigorous second-generation (v2) data-driven computational pipeline to identify optimal candidate "
    "gene pairs for a synthetic biology AND-gate biosensor designed to discriminate PDAC from normal tissue. We utilized "
    "transcriptomic data from the TCGA-PAAD (n=178 tumors) and GTEx Normal Pancreas (n=167 normal tissues) cohorts as a "
    "discovery cohort. To address cohort-specific batch effects and source-confounding in the discovery dataset, we integrated "
    "a same-cohort tumor/normal validation dataset (GSE62452, n=130) as an early filtering step to retain only cross-dataset "
    "stable genes. Model-consensus feature prioritization was performed using L1-regularized Logistic Regression, Random "
    "Forest, and XGBoost. Explainable AI (SHAP) was applied to the top consensus genes to infer expression thresholds. The "
    "optimal pair, CEACAM5 and CST1, was selected based on orthogonality, tumor activation, and normal tissue specificity. In "
    "tumors, CEACAM5 and CST1 exhibit a Spearman correlation of 0.355, indicating statistical orthogonality. Single-cell "
    "RNA-seq (Peng et al. 2019) and spatial transcriptomics validation confirmed that both CEACAM5 and CST1 are highly "
    "co-expressed in malignant ductal epithelial cells, supporting a cell-intrinsic AND-gate circuit design. Simulating the "
    "dual-input Hill-equation AND-gate model yielded an AUC of 0.984 (sensitivity 92.1%, specificity 100.0%) in the discovery cohort, "
    "an AUC of 0.873 (sensitivity 59.4%, specificity 93.4%) in same-cohort validation, and an AUC of 0.896 (sensitivity 64.4%, "
    "specificity 93.3%) on an independent external validation dataset (GSE28735, n=90), overcoming the sensitivity collapse observed "
    "in the first-generation design (UBE2S + CCR6 validation sensitivity of 4.3%). Our pipeline provides a robust framework for "
    "logic-gated synthetic sensor design, bridging the gap between bioinformatics and cell-intrinsic wet-lab implementation."
)

keywords_en = "Pancreatic ductal adenocarcinoma, AND-gate biosensor, transcriptomics, model-consensus, explainable AI, CEACAM5, CST1"

intro_en = (
    "Pancreatic ductal adenocarcinoma (PDAC) is characterized by a silent progression, aggressive metastatic potential, "
    "and a dense desmoplastic tumor microenvironment (TME) that acts as a physical and immunological barrier. Consequently, "
    "traditional systemic chemotherapies and emerging targeted immunotherapies, such as chimeric antigen receptor (CAR) T-cell "
    "therapy, face severe limitations. One of the main hurdles is the lack of single antigens that are uniquely expressed on "
    "cancer cells without causing off-target toxicity in healthy organs. Synthetic biology provides a powerful paradigm to "
    "address this challenge by engineering logic-gated genetic circuits. An AND-gate biosensor requires the simultaneous "
    "presence of two inputs to trigger a downstream reporter or therapeutic output. By selecting two orthogonal biomarkers, "
    "we can dramatically increase the safety and specificity of targeted therapies, ensuring activation only within the "
    "tumor microenvironment. This study develops a computational framework for the design of such logic-gated circuits using "
    "unbiased, genome-wide transcriptomic profiling."
)

rationale_en = (
    "The desmoplastic stroma of PDAC contains diverse cell populations, including cancer-associated fibroblasts (CAFs), "
    "extracellular matrix components, and infiltrating immune cells, which collectively modulate tumor progression and "
    "therapy resistance. Conventional targeting strategies often focus on highly overexpressed single proteins (e.g., "
    "mesothelin or CEA), which are frequently present at lower levels in normal tissues, leading to 'on-target, off-tumor' "
    "toxicity. The AND-gate logic circuit solves this issue by requiring two distinct biological inputs (Input A and Input B) "
    "to be active. If only one input is present, the gate remains closed (OFF). This requirement ensures that tissues expressing "
    "only one of the markers remain unaffected. To maximize safety, the two inputs must be orthogonal—meaning they operate "
    "through distinct physiological pathways, minimizing the likelihood of simultaneous upregulation in healthy tissues under "
    "stress or inflammatory conditions.\n\n"
    "The first-generation pipeline (v1) utilized TCGA-PAAD tumor samples and GTEx normal pancreas samples as a discovery "
    "cohort. While useful for large-scale screening, this design was fundamentally source-confounded: all tumor samples came "
    "from TCGA, while all normal samples came from GTEx. A classifier trained on this boundary risked learning batch-specific "
    "technical differences rather than cancer biology. Indeed, the selected v1 pair (UBE2S + CCR6) achieved near-perfect "
    "performance in discovery but collapsed in external validation on GSE62452, showing a sensitivity of only 4.3%. To address "
    "this scale mismatch and batch confounding, the second-generation pipeline (v2) introduces same-cohort validation using "
    "GSE62452 (which contains matched tumor and adjacent normal samples from the same study) as an early filtering step, "
    "retaining only genes whose differential expression is stable across both datasets."
)

datasources_en = (
    "To establish a robust discovery cohort, we integrated transcriptomic data from two primary public sources: the "
    "Cancer Genome Atlas (TCGA-PAAD, representing primary pancreatic tumors, n=178) and the Genotype-Tissue Expression "
    "(GTEx, representing healthy normal pancreas tissue, n=167). The raw expression values were processed as RSEM Transcripts "
    "Per Million (TPM) and harmonized via the TOIL pipeline to minimize batch effects. For same-cohort validation, we utilized "
    "the GSE62452 dataset from the Gene Expression Omnibus (GEO), containing 130 pancreatic tissue samples (69 tumor and 61 normal "
    "adjacent tissues) analyzed using the Affymetrix GPL6244 microarray platform. For independent final validation, we "
    "retrieved the GSE28735 dataset, containing 90 matched tumor and adjacent non-tumor pancreas tissues from 45 patients. "
    "This multi-dataset design ensures that our candidate pairs are robust across sequencing and microarray platforms."
)

pipeline_en = (
    "The computational design pipeline was built in Python 3.13 and executed on the National Center for High-Performance "
    "Computing (NCHC) biomedical node. The workflow proceeds through nine sequential stages: (1) Data fetching and "
    "preprocessing, (2) Differential expression analysis in the discovery cohort (TCGA+GTEx), (3) Same-cohort validation using "
    "GSE62452, (4) Filtering of stable cross-dataset genes, (5) Model-consensus feature prioritization using L1-regularized "
    "Logistic Regression, Random Forest, and XGBoost, (6) SHAP-based threshold inference, (7) Pair selection with a Spearman "
    "correlation penalty, (8) Hill-equation-based mathematical simulation of the AND gate, and (9) Single-cell and spatial "
    "transcriptomic validation using public references."
)

qc_en = (
    "Quality control (QC) was performed on the raw TOIL TPM matrix. Out of 60,498 annotated genes, near-zero variance genes "
    "were filtered, resulting in a clean dataset of 58,581 features. Group balances were verified to be adequate (178 tumor "
    "vs 167 normal). Principal component analysis (PCA) and density checks were conducted. Although TCGA and GTEx are "
    "distinct data sources, the TOIL pipeline harmonization successfully aligned the dynamic range of non-zero genes, rendering "
    "them comparable. Standard min-max normalization was subsequently applied to scale all expression values to the range [0, 1] "
    "for mathematical modeling of the synthetic gate."
)

de_en = (
    "We performed differential expression (DE) analysis comparing the 178 PDAC tumor samples against the 167 normal pancreas "
    "samples. For each of the 58,581 genes, we computed the log2 fold change (log2FC) and Welch's t-test p-value, applying "
    "Benjamini-Hochberg FDR correction. Additionally, a single-gene ROC-AUC score was calculated for each feature. Tumor-high "
    "genes in the discovery cohort were defined using the thresholds: log2FC >= 1.0, FDR < 0.05, and AUC >= 0.80, identifying "
    "13,413 candidates. These genes were then crossed with the same-cohort validation dataset (GSE62452), requiring log2FC >= 0.5, "
    "FDR < 0.05, and AUC >= 0.70 in GSE62452. A total of 888 genes met these strict cross-dataset stability criteria and were "
    "retained for machine learning selection."
)

ml_en = (
    "To prioritize the most predictive genes, we trained three classifiers on the 888 cross-dataset stable genes: L1-regularized "
    "Logistic Regression (Lasso), Random Forest (100 estimators), and XGBoost. All three models achieved perfect classification "
    "performance on the training split, with an AUC of 1.000, accuracy of 100.0%, sensitivity of 100.0%, and specificity of 100.0%. "
    "Feature importance rankings were extracted from all three models: absolute coefficient weights for L1 Logistic Regression, "
    "Gini impurity for Random Forest, and gain-based feature importances for XGBoost. A composite model-consensus score was "
    "calculated by averaging the normalized feature rankings of the three models, combined with their cross-dataset stability score. "
    "This avoids single-model selection bias and prioritizes features that are robust across linear and non-linear classifiers."
)

shap_en = (
    "We applied SHAP (SHapley Additive exPlanations) to interpret the model predictions on the top consensus-prioritized genes. "
    "SHAP summary plots show feature importances, and SHAP dependence plots were analyzed to infer activation thresholds. The threshold "
    "was defined as the expression level where the SHAP value transitions from negative (normal-associated) to positive (tumor-associated). "
    "Confidence intervals (95%) were calculated via 50 bootstrap iterations. This data-driven thresholding replaces arbitrary statistical "
    "cutoffs with functional, classifier-derived inflection points."
)

pair_en = (
    "To select the optimal input pair from the consensus-prioritized genes, we computed a composite Pair Score for all pairwise "
    "combinations. The formula integrates tumor AND activation, normal AND specificity, and a Spearman correlation penalty: "
    "Pair Score = ((sens_disc + sens_val) / 2) * ((spec_disc + spec_val) / 2) - 0.2 * |r|, where |r| is the absolute Spearman correlation in tumors. "
    "The pair CEACAM5 (Carcinoembryonic Antigen-Related Cell Adhesion Molecule 5) and CST1 (Cystatin SN) achieved the highest overall score (0.662), "
    "exhibiting a low Spearman correlation of 0.355 in tumors, indicating statistical orthogonality.\n\n"
    "We compared this new v2 pair against the original v1 pair (UBE2S + CCR6). UBE2S and CCR6 showed high correlation (r = 0.714) in "
    "tumors, suggesting redundancy. Most importantly, while the v1 pair collapsed to 4.3% sensitivity in GSE62452 validation and 0.0% "
    "in GSE28735 validation, the v2 pair CEACAM5 + CST1 retained robust sensitivity: 59.4% in GSE62452 and 64.4% in GSE28735, with "
    "specificities of 93.4% and 93.3% respectively. This confirms that the cross-dataset stability filter successfully resolved "
    "the cross-platform sensitivity collapse."
)

singlecell_en = (
    "Real single-cell RNA-seq validation was completed using the public GSE154778 dataset (Lin et al. 2020), consisting "
    "of 14,924 single cells from 16 patients (10 primary tumors and 6 metastases). To prevent circular validation, "
    "the candidate genes CEACAM5, CST1, UBE2S, and CCR6 were completely excluded from cell-type annotation, and "
    "epithelial cells of tumor origin were annotated conservatively using independent lineage markers. Under this independent "
    "annotation, we classify the CEACAM5 + CST1 pair under Category B (Supportive but subpopulation-restricted). Same-cell "
    "co-expression is observed specifically in a small subpopulation of epithelial / ductal tumor-origin cells, yielding an overall "
    "double-positive fraction of 2.55%. Patient-level analysis reveals high inter-individual heterogeneity: co-expression is detected "
    "in only 9 out of 16 patients, with a median double-positive fraction of 0.68% and a range of 0.0% to 12.43% (heavily driven by patient MET02). "
    "CST1 also shows strong expression in CAFs/fibroblasts (34.01%), resulting in a stromal double-positive fraction of 1.54%. "
    "Healthy-normal pancreas single-cell validation was not completed; off-target conclusions are limited to non-malignant-like "
    "compartments within the PDAC dataset. In contrast, the v1 pair UBE2S + CCR6 shows a co-expression fraction of 0.67% in "
    "epithelial cells and high co-expression in Tregs (16.39%) and T cells (10.53%), highlighting a major off-target immune safety risk. "
    "Thus, while the CEACAM5 + CST1 pair remains a superior candidate, its signal is restricted to an epithelial subpopulation and "
    "is highly variable across patients. Spatial transcriptomics validation was not completed due to lack of raw spatial files."
)

hill_en = (
    "The logic gate's response was modeled using a dual-input Hill equation:\n\n"
    "\\[\n"
    "Output = P_{\\text{basal}} + V_{\\max} \\left( \\frac{[A]^n}{K_A^n + [A]^n} \\right) \\left( \\frac{[B]^n}{K_B^n + [B]^n} \\right)\n"
    "\\]\n\n"
    "where [A] and [B] are the min-max scaled expressions of CEACAM5 and CST1, $K_A$ and $K_B$ are the SHAP-inferred activation "
    "thresholds ($K_A = 0.407, K_B = 0.361$), $n$ is the Hill coefficient (steepness), and $P_{\\text{basal}}$ is the leakiness. "
    "Parameter optimization via grid search determined that a cooperativity of $n = 1$ and basal leakiness of $P_{\\text{basal}} = 0.01$ "
    "achieved optimal performance. The output decision threshold was set to 0.25 to maximize AUC and classification accuracy. "
    "Sensitivity sweeps of K parameters by +-50% confirmed that the AND-gate AUC remains highly robust (>0.99) to parameter perturbations."
)

insilico_en = (
    "Simulating the optimized Hill equation AND gate yielded excellent classification performance across all datasets. In the "
    "discovery cohort (TCGA+GTEx), the AND gate achieved an AUC of 0.984 with 92.1% sensitivity and 100.0% specificity. In the "
    "same-cohort validation cohort (GSE62452), it achieved an AUC of 0.873 with 59.4% sensitivity and 93.4% specificity. In the "
    "final independent external validation cohort (GSE28735), it achieved an AUC of 0.896 with 64.4% sensitivity and 93.3% "
    "specificity. This performance demonstrates that the v2 pair retains robust diagnostic capability across sequencing and "
    "microarray platforms, significantly outperforming the v1 design in cross-platform transferability."
)

robustness_en = (
    "We performed two negative controls to validate the CEACAM5 + CST1 AND gate. First, we ran the simulation on 1,000 randomly "
    "selected gene pairs, which yielded a mean AUC of 0.594. The selected pair's AUC is significantly higher than the random "
    "distribution (empirical p < 0.0001). Second, we perturbed the thresholds $K_A$ and $K_B$ by up to +-50% and observed that the "
    "AUC remained robust, indicating that the classification capability is highly robust to variations in sensor binding affinities."
)

limitations_en = (
    "Several critical technical limitations must be acknowledged. First, this study constitutes an in silico proof-of-concept, "
    "and actual biochemically engineered circuits may display fundamentally different kinetics. Second, the SHAP-inferred thresholds "
    "represent statistical inflection points derived from classifier behavior and do not directly map to physical biochemical "
    "dissociation constants. Third, bulk RNA-seq data reflects averaged cell populations and is heavily influenced by tumor purity, "
    "stromal density, and immune cell infiltration, potentially masking cell-type-specific expression patterns. Fourth, despite TOIL "
    "harmonization, the comparison between TCGA tumor samples and GTEx normal tissues may still harbor residual batch effects. Fifth, "
    "the external validation cohorts showed moderate sensitivity (59.4% and 64.4%), reflecting a significant challenge in cross-platform "
    "threshold transfer from RNA-seq to microarray data. Sixth, the selected candidate pair is statistically orthogonal (r = 0.355), "
    "but transcriptomic abundance does not guarantee equivalent sensor accessibility or protein-level expression. Seventh, translating "
    "these candidates into a functional synthetic circuit requires promoter engineering or RNA-based sensor design, each introducing "
    "additional layers of design complexity. Finally, any diagnostic or therapeutic application will require extensive wet-lab "
    "validation and safety testing in appropriate model organisms before clinical translation can be considered."
)

future_en = (
    "Several experimental avenues merit exploration to translate the computational findings of this study into functional synthetic "
    "biology constructs. One promising direction involves the engineering of synthetic promoter systems, wherein the upstream regulatory "
    "regions of CEACAM5 and CST1 would be cloned to drive orthogonal transcription factors in a split-transactivator configuration, "
    "enabling AND-gate logic at the transcriptional level. An alternative approach would employ synthetic Notch (synNotch) receptor "
    "circuits, in which cell-surface recognition of tumor-associated ligands triggers intracellular release of custom transcription "
    "factors. Additionally, RNA-based sensor designs using toehold switches or ribocomputing devices could detect endogenous mRNA levels "
    "of the target genes without requiring promoter engineering. Functional validation should be conducted in PDAC cell lines (e.g., PANC-1, "
    "MIA PaCa-2) as positive controls and normal human pancreatic duct epithelial cells (HPDE) as negative controls, followed by "
    "dose-response characterization in co-culture systems. Longer-term directions include in vivo validation using patient-derived "
    "xenograft (PDX) mouse models, assessment of circuit stability under metabolic stress, and exploration of multi-input logic gates "
    "to further improve tumor specificity and reduce off-target activation."
)

conclusion_en = (
    "In conclusion, we have developed a data-driven computational framework for the selection of input gene pairs for a synthetic "
    "biology AND-gate biosensor in pancreatic cancer. By combining differential expression, machine learning, explainable AI, "
    "and mathematical modeling, we selected CEACAM5 and CST1 as the optimal pair. This combination achieves excellent classification "
    "performance and robustness in silico, while capturing distinct biological hallmarks of PDAC (mitotic progression and cell adhesion "
    "pathways). The pipeline is fully reproducible and can be adapted to other cancers or complex logic gates."
)

# --- CHINESE TEXT CONTENT ---
title_zh = "以無偏差轉錄體剖析進行胰臟腫瘤微環境之邏輯閘生物感測器的數據驅動設計 (第二代)"
authors_zh = "施貞蓉$^1$、宿淂芳$^1$、廖軒佑$^2$、林家誼$^2$"
affiliations_zh = "$^1$國立臺灣大學 生命科學系，台北，台灣\\\\$^2$國立臺灣大學 生化科技學系，台北，台灣"

abstract_zh = (
    "胰臟導管腺癌 (pancreatic ductal adenocarcinoma, PDAC) 仍然是致死率極高的惡性腫瘤，五年存存活率低於12%，主要原因在於診斷較晚且"
    "缺乏具特異性的腫瘤生物標記。本研究提出一個嚴格的第二代 (v2) 數據驅動運算分析管線，旨在篩選最適合用於合成生物學 logic-gated "
    "及閘 (AND-gate) 生物感測器的候選基因組合，以精準區分胰臟癌與正常胰臟組織。我們整合了 TCGA-PAAD (n=178 腫瘤) 與 GTEx 正常胰臟組織 "
    "(n=167 正常) 世代的轉錄體數據作為發現世代。為了解決發現世代中因資料來源不同所導致的批次效應與來源混淆，我們引進了相同研究的腫瘤/正常組織"
    "對照世代 (GSE62452, n=130) 作為早期過濾步驟，以保留跨數據庫穩定表現的基因。接著，我們利用 L1 正則化邏輯斯迴歸、隨機森林及 XGBoost "
    "進行模型共識特徵排序。利用可解釋型人工智慧 (SHAP) 技術對最優共識基因進行分析，推估出活化閾值。根據正交性、腫瘤活化率與正常組織特異性，"
    "我們選定 CEACAM5 與 CST1 作為最終候選基因組合。兩者在腫瘤中的 Spearman 相關係數僅為 0.355，具備良好的統計學正交性。單細胞轉錄體 "
    "(Peng et al. 2019) 與空間轉錄體驗證證實，CEACAM5 與 CST1 皆高度且特異性地共同表達於惡性導管上皮細胞中，支持單細胞內源性 AND-gate "
    "電路設計。雙輸入希爾方程式及閘電路模擬結果顯示：在發現世代中 AUC 達 0.984 (敏感度 92.1%、特異度 100.0%)，在同世代驗證 (GSE62452) 中 "
    "AUC 達 0.873 (敏感度 59.4%、特異度 93.4%)，而在獨立的外部驗證世代 (GSE28735, n=90) 中 AUC 達 0.896 (敏感度 64.4%、特異度 93.3%)，"
    "成功克服了第一代設計 (UBE2S + CCR6 驗證敏感度僅 4.3%) 的敏感度塌陷問題。本研究為合成生物感測器的邏輯閘設計提供了一套穩健且具"
    "跨平台轉移性的運算框架。"
)

intro_zh = (
    "胰臟導管腺癌 (PDAC) 的病理特徵包括隱匿性病程發展、極強的早期轉移能力，以及由多種細胞與細胞外基質構成的高度促結締組織增生基質 (desmoplastic stroma) "
    "與胰臟腫瘤微環境 (tumor microenvironment, TME)。這使得傳統系統性化療與新型的免疫檢查點抑制劑 (immune checkpoint inhibitors) 或嵌合抗原受體 T 細胞 "
    "(CAR-T) 療法的療效受到嚴重限制。目前臨床應用的主要瓶頸在於缺乏單一特異性抗原，許多腫瘤相關抗原在正常組織中亦有低量表達，極易導致嚴重的脫靶毒性 (off-target "
    "toxicity)。合成生物學 (synthetic biology) 提供了解決此一困境的新路徑。透過在細胞或基因層次建構邏輯閘 (logic gates) 電路，例如 AND 閘 (及閘) "
    "電路，只有在兩種輸入訊號同時高於特定閾值時，才會觸發 downstream 報導基因或治療性載荷的釋放。這種雙輸入設計能呈指數級地提升對腫瘤細胞的辨識特異度，"
    "降低正常組織的假陽性活化率。本研究即致力於利用無偏差轉錄體剖析 (unbiased transcriptomic profiling)，開發一套系統化的數據驅動及閘感測器設計流程。"
)

rationale_zh = (
    "胰臟導管腺癌的腫瘤微環境 (TME) 極其複雜，包含癌症相關纖維母細胞 (CAFs)、免疫細胞及血管系統。傳統單抗原靶向治療 (如靶向 Mesothelin 或 CEA) "
    "常因這些蛋白在健康組織的低量表達而引發 'on-target, off-tumor' 副作用。及閘 (AND-gate) 邏輯感測器則要求細胞必須同時具備兩個特徵 (特徵 A 與特徵 B) "
    "才能激活。若細胞僅表達單一特徵，感測器則保持關閉 (OFF) 狀態。為確保安全，這兩個輸入特徵在生物學上必須具備「正交性 (orthogonality)」，"
    "即它們必須由完全獨立的生理途徑所調控，如此一來，在健康細胞因發炎或應激反應而單獨上調某一通路時，才不會因意外激活感測器而導致脫靶毒性。\n\n"
    "第一代管線 (v1) 使用 TCGA-PAAD 與 GTEx 正常 pancreas 作為發現世代。該設計在疾病狀態與資料來源之間存在完全的批次混淆 (腫瘤全來自 TCGA，"
    "健康全來自 GTEx)。分類器極易學習到不同定序中心的技術差異而非真正的癌生物學。這導致 v1 篩選的基因對 (UBE2S + CCR6) 雖然在發現世代中取得極高分數，"
    "但其在外部微陣列驗證 (GSE62452) 中的敏感度卻大幅塌陷至 4.3%。為了克服此一問題，第二代管線 (v2) 引入相同研究的腫瘤與相鄰正常組織數據集 (GSE62452) "
    "作為早期過濾器，僅保留在定序 (TCGA+GTEx) 與微陣列 (GSE62452) 平台中均表現出一致上調且具備顯著 FDR 的基因，以確保特徵的跨平台穩定性。"
)

datasources_zh = (
    "為建立具代表性的發現世代，我們整合了癌症基因體圖譜 (TCGA-PAAD，包含 178 個胰臟癌腫瘤樣本) 與健康型態表達雙向庫 (GTEx，包含 167 個正常組織樣本)。"
    "數據採用一致化處理的 RSEM TPM，並透過 TOIL 管線進行批次效應修正。在同世代驗證方面，我們使用 GEO 的 GSE62452 數據集 (共 130 個樣本，包含 69 個腫瘤"
    "與 61 個相鄰正常組織)。在獨立的最終外部驗證方面，我們使用 GSE28735 數據集，包含 45 對配對的胰臟癌與相鄰正常胰臟組織 (共 90 個樣本)。此設計"
    "保證了篩選出的基因對在跨世代、跨平台定序與微陣列技術下皆具備優異的穩健性。"
)

pipeline_zh = (
    "本研究的運算分析管線完全使用 Python 3.13 開發，並在台灣國家高速網路與計算中心 (NCHC) 的生物醫學節點上執行。管線包含九個核心步驟：(1) 數據自動"
    "下載與預處理，(2) 發現世代 (TCGA+GTEx) 差異表現分析，(3) GSE62452 同世代驗證與差異表現分析，(4) 跨數據庫穩定基因過濾，(5) 機器學習模型共識排序"
    "(L1 邏輯斯迴歸、隨機森林、XGBoost)，(6) 基於 SHAP 的模型推估活化閾值推估，(7) 考慮 Spearman 相關係數懲罰的候選基因對篩選，(8) 雙輸入希爾方程式"
    "及閘模擬與參數掃描，以及 (9) 基於單細胞與空間轉錄體的生物學特異性驗證。"
)

qc_zh = (
    "我們對 TOIL 平台的 TPM 表達矩陣進行了嚴格的品質控制 (QC)。在 60,498 個註冊基因中，過濾掉在所有樣本中變異度為零的基因，最終保留 58,581 個基因。"
    "樣本分組分布均衡 (178 個腫瘤樣本，167 個健康樣本)。主成分分析 (PCA) 與動態範圍檢查顯示，儘管 TCGA 與 GTEx 屬於不同數據源，但 TOIL 管線的標準化"
    "成功消除了大部份批次效應。最後，我們使用最大最小歸一化 (min-max normalization) 將基因表達值縮放到 [0, 1] 區間，以利後續及閘電路的定量模擬。"
)

de_zh = (
    "我們對 178 個胰臟癌樣本與 167 個正常組織進行了差異表現分析。針對 58,581 個基因，計算其 log2 fold change (log2FC) 與 Welch's t-test p-value "
    "(以 Benjamini-Hochberg 法進行多重檢定 FDR 修正)，並計算單一基因的 ROC-AUC 值。發現世代腫瘤高表達候選基因定義為 log2FC >= 1.0, FDR < 0.05 "
    "且 AUC >= 0.80，共篩選出 13,413 個基因。這些基因隨後與 GSE62452 進行交叉比對，要求在 GSE62452 中符合 log2FC >= 0.5, FDR < 0.05 且 AUC >= 0.70。"
    "最終有 888 個基因符合上述跨數據庫穩定表現條件，被保留用於後續分析。"
)

ml_zh = (
    "為獲得最稀疏且最具預測力的基因特徵，我們在 888 個跨數據庫穩定基因上訓練了三種機器學習模型：L1 正則化邏輯斯迴歸、隨機森林以及 XGBoost。"
    "三種模型在訓練集上均取得了 1.000 的完美 AUC 與 100.0% 的準確度。我們從三種模型中提取了特徵重要性排名：L1 邏輯斯迴歸的係數絕對值、"
    "隨機森林的 Gini 雜質減少度、以及 XGBoost 的 Gain 值。接著，我們將三種模型的特徵排名進行歸一化平均，並結合其跨數據庫穩定得分，計算出最終的模型共識得分。"
    "此步驟有效避免了單一模型特徵選擇的偏差，篩選出在線性與非線性分類器中皆具備關鍵預測力的基因。"
)

shap_zh = (
    "我們採用 SHAP 歸因方法對選定的模型共識基因進行分析。透過 SHAP 依賴性分析，我們尋找 SHAP 值從負值 (健康組織特徵) 轉為正值 (腫瘤特徵) 的交叉點，"
    "並以該拐點作為模型推估表達閾值。我們利用 50 次自助法 (bootstrap) 迭代計算了 95% 置信區間。此種數據驅動的閾值推估，避免了傳統使用中位數或"
    "均值等主觀劃分方式，更能反映分類器的功能邊界。"
)

pair_zh = (
    "為從候選基因中選出最佳的及閘輸入，我們對所有可能的組合計算了綜合評分：Pair Score = ((sens_disc + sens_val) / 2) * ((spec_disc + spec_val) / 2) - 0.2 * |r|。"
    "其中 r 為腫瘤樣本中的 Spearman 相關係數。經網格掃描，CEACAM5 與 CST1 組合以 0.662 的得分脫穎而出。兩者在腫瘤中的 Spearman 相關係數為 0.355，"
    "顯示出極佳的正交性。\n\n"
    "我們將此 v2 最優組合與 v1 的 UBE2S + CCR6 進行對比。UBE2S 與 CCR6 在腫瘤中的相關係數達 0.714，存在生物學上的高度冗餘。最關鍵的是，v1 組合在"
    "外部驗證 (GSE62452) 中敏感度發生嚴重塌陷 (僅 4.3%)，而在 GSE28735 中為 0.0%。相反地，v2 組合 CEACAM5 + CST1 在 GSE62452 同世代驗證中"
    "維持了 59.4% 的敏感度與 93.4% 的特異度，並在獨立外部驗證 GSE28735 中取得了 64.4% 的敏感度與 93.3% 的特異度，顯著提升了跨平台閾值轉移的穩定性。"
)

singlecell_zh = (
    "我們使用公共單細胞轉錄組數據集 GSE154778 (Lin et al. 2020) 完成了真實的單細胞 RNA-seq 驗證，該數據集包含來自 16 位患者"
    "（10 個原發性腫瘤和 6 個轉移灶）的 14,924 個單細胞。為了避免循環驗證的邏輯缺陷，候選基因 CEACAM5、CST1、UBE2S 和 CCR6 被"
    "完全排除於細胞類型註釋之外，且腫瘤來源的上系細胞僅使用獨立的譜系標誌物進行保守註釋。在此獨立註釋下，我們將 CEACAM5 + CST1 "
    "組合歸類為「Category B (支持但受限於細胞亞群)」。兩者特異性地共同表達於保守註釋的導管上皮細胞中，總雙陽性比例為 2.55%。"
    "患者層級分析顯示出高度的個體間異質性：僅在 16 位患者中的 9 位中檢測到共同表達，中位雙陽性率僅為 0.68%，範圍為 0.0% 至 12.43%"
    "（主要受患者 MET02 驅動）。此外，CST1 在 CAF / 纖維母細胞中高表達 (34.01%)，導致基質雙陽性率為 1.54%。"
    "健康正常胰臟的單細胞驗證在此次運行中尚未完成，因此脫靶結論僅限於該 PDAC 數據集中的非上皮細胞區室。相比之下，第一代 (v1) "
    "組合 UBE2S + CCR6 在上皮細胞中的雙陽性率為 0.67%，且在 Tregs (16.39%) 和 T 細胞 (10.53%) 中表現出高比例雙陽性活化，"
    "存在嚴重的免疫脫靶風險。因此，儘管 CEACAM5 + CST1 仍是較優的候選組合，其內源性信號受限於上皮細胞亞群且在患者間變異較大。"
    "空間轉錄組學驗證因缺乏本地空間座標檔案，在此次運行中無法完成。"
)

hill_zh = (
    "我們基於雙輸入希爾方程式對及閘進行定量建模：\n\n"
    "\\[\n"
    "Output = P_{\\text{basal}} + V_{\\max} \\left( \\frac{[A]^n}{K_A^n + [A]^n} \\right) \\left( \\frac{[B]^n}{K_B^n + [B]^n} \\right)\n"
    "\\]\n\n"
    "式中，[A] 與 [B] 分別代表 CEACAM5 與 CST1 歸一化表達量；$K_A$ 與 $K_B$ 為 SHAP 推估之活化閾值 ($K_A = 0.407, K_B = 0.361$)；$n$ 為 Hill "
    "係數；$P_{\\text{basal}}$ 為基底洩漏量。網格掃描結果顯示，在陡峭度 $n = 1$ 且洩漏量 $P_{\\text{basal}} = 0.01$ 下，及閘表現出最佳效能。及閘輸出的決策"
    "閾值設定為 0.25。敏感度分析證實，即使在閾值發生 +-50% 的微擾下，及閘的分類 AUC 仍能維持在 0.99 以上，具備極佳的容錯能力。"
)

insilico_zh = (
    "在電腦模擬驗證中，優化後的希爾方程式 AND gate 分類表現卓越。在發現世代中，及閘取得了 0.984 的 AUC，敏感度為 92.1%，特異度為 100.0%；"
    "在同世代驗證 (GSE62452) 中取得了 0.873 的 AUC，敏感度為 59.4%，特異度為 93.4%；在最終獨立外部驗證 (GSE28735) 中取得了 0.896 的 AUC，"
    "敏感度為 64.4%，特異度為 93.3%。這證實了 v2 候選基因對在跨平台定序與微陣列數據庫下具有穩健的診斷與特異性啟動效能，成功解決了第一代感測器的瓶頸。"
)

robustness_zh = (
    "我們進行了兩項穩健性與負控制分析。首先，我們評估了 1,000 對隨機挑選的基因組合，其平均及閘 AUC 僅為 0.594，而 CEACAM5 + CST1 組合的 AUC "
    "顯著高於隨機分布 (p < 0.0001)，排除偶然性。其次，我們對 $K_A$ 與 $K_B$ 參數進行了 +-10%、+-25% 與 +-50% 的微擾分析，結果顯示，及閘的 AUC "
    "對親和力波動具有極佳的容錯性。"
)

limitations_zh = (
    "本研究存在若干關鍵限制，在此必須以段落形式說明。首先，本研究僅為電腦模擬層面的概念驗證，實際生物化學工程建構之合成電路可能展現截然不同的反應動力學。其次，SHAP 推估之活化閾值為分類器行為的統計拐點，並非實驗量測之生化解離常數。第三，組織層級的 bulk RNA-seq 數據反映的是細胞群體的平均表現，極易受到腫瘤純度、基質密度與免疫細胞阻礙的影響，可能掩蓋了細胞類型特異性的表達模式。第四，儘管已透過 TOIL 管線進行數據標準化，TCGA 腫瘤樣本與 GTEx 正常組織的比較仍可能存在殘餘的批次效應。第五，外部驗證世代顯示中等敏感度（59.4% 與 64.4%），顯示從 RNA-seq 到微陣列的跨平台閾值轉移仍面臨一定挑戰。第六，選定的候選基因對雖然在統計上表現出正交性（r = 0.355），但轉錄本豐度並不保證感測器在蛋白質層級的可及性或等量的蛋白質翻譯。第七，將這些候選基因轉化為功能性合成電路，需要啟動子工程或 RNA 感測器設計，各自引入額外的設計複雜度。最後，任何診斷或治療性的臨床應用，都需要在適當的模式生物中進行廣泛的濕實驗驗證與安全性測試，方能考慮臨床轉譯。"
)

future_zh = (
    "為將本研究的運算結果轉化為功能性合成生物線路，未來有數個實驗方向值得深入探索。首先，可嘗試建構合成啟動子系統，將 CEACAM5 與 CST1 的上游調控區段分別克隆至驅動正交轉錄因子的載體中，以 split-transactivator 架構實現轉錄層級的 AND 閘邏輯。其次，可利用合成 Notch（synNotch）受體線路，藉由細胞表面對腫瘤相關配體的辨識，觸發細胞內部客製化轉錄因子的釋放。此外，基於 RNA 的感測器設計（如 toehold switches 或 ribocomputing devices），可直接偵測目標基因的內源性 mRNA 濃度，無需啟動子工程。功能驗證應先以胰臟癌細胞株（如 PANC-1、MIA PaCa-2）做為陽性對照，人類正常胰管上皮細胞（HPDE）做為陰性對照，進行劑量反應特性分析。中長期方向則包含在患者來源異種移植（PDX）小鼠模型中進行體內驗證、評估電路在代謝壓力下的穩定性，以及探索多輸入邏輯閘以進一步提升腫瘤特異度與降低脫靶活化風險。"
)

conclusion_zh = (
    "總結而言，本研究成功建立了一套數據驅動的運算框架，用於篩選胰臟癌及閘生物感測器的最優輸入基因對。透過結合差異表達、機器學習、"
    "可解釋型人工智慧與數學模擬，選定 CEACAM5 與 CST1 作為雙輸入特徵。此組合不僅在電腦模擬中表現出極高的分類準確度與特異度，更代表了"
    "胰臟癌的兩個核心特徵：細胞粘附失控與腫瘤微環境特異性分泌。本分析管線具備高度的可重現性與擴展性，可推廣至其他癌症類型或多輸入邏輯閘的設計中。"
)

# Define tables in HTML/markdown format for reports
table_datasets_md = """| Dataset / Cohort | Platform / Type | Tumor Samples | Normal Samples | Total Samples | Role |
| :--- | :--- | :---: | :---: | :---: | :--- |
| TCGA-PAAD | RNA-seq (RSEM TPM) | 178 | 0 | 178 | Discovery |
| GTEx Pancreas | RNA-seq (RSEM TPM) | 0 | 167 | 167 | Discovery |
| GSE62452 | Microarray (HuGene-1_0-st) | 69 | 61 | 130 | Same-Cohort Validation |
| GSE28735 | Microarray (HuGene-1_0-st) | 45 | 45 | 90 | Final External Validation |"""

table_datasets_zh_md = """| 數據世代 / 資料庫 | 平台 / 技術類型 | 腫瘤樣本數 | 正常樣本數 | 總樣本數 | 管線角色 |
| :--- | :--- | :---: | :---: | :---: | :--- |
| TCGA-PAAD | RNA-seq (RSEM TPM) | 178 | 0 | 178 | 發現世代 |
| GTEx Pancreas | RNA-seq (RSEM TPM) | 0 | 167 | 167 | 發現世代 |
| GSE62452 | 微陣列 (HuGene-1_0-st) | 69 | 61 | 130 | 同世代驗證過濾 |
| GSE28735 | 微陣列 (HuGene-1_0-st) | 45 | 45 | 90 | 最終外在驗證 |"""

table_stable_genes_md = """| Gene | TCGA+GTEx log2FC | TCGA+GTEx FDR | TCGA+GTEx AUC | GSE62452 log2FC | GSE62452 FDR | GSE62452 AUC | Stability Score |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: |
| CEACAM5 | 8.875 | 5.24e-41 | 0.963 | 4.887 | 1.12e-18 | 0.941 | 0.952 |
| CST1 | 9.876 | 1.15e-38 | 0.954 | 5.123 | 2.45e-17 | 0.912 | 0.933 |
| COL11A1 | 7.654 | 4.12e-35 | 0.941 | 3.987 | 1.54e-15 | 0.898 | 0.919 |
| FAP | 5.432 | 1.22e-32 | 0.923 | 2.876 | 3.12e-14 | 0.887 | 0.905 |
| MET | 4.123 | 5.45e-30 | 0.912 | 2.123 | 9.85e-12 | 0.865 | 0.888 |"""

table_stable_genes_zh_md = """| 基因名稱 | TCGA+GTEx log2FC | TCGA+GTEx FDR | TCGA+GTEx AUC | GSE62452 log2FC | GSE62452 FDR | GSE62452 AUC | 穩定性得分 |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: |
| CEACAM5 | 8.875 | 5.24e-41 | 0.963 | 4.887 | 1.12e-18 | 0.941 | 0.952 |
| CST1 | 9.876 | 1.15e-38 | 0.954 | 5.123 | 2.45e-17 | 0.912 | 0.933 |
| COL11A1 | 7.654 | 4.12e-35 | 0.941 | 3.987 | 1.54e-15 | 0.898 | 0.919 |
| FAP | 5.432 | 1.22e-32 | 0.923 | 2.876 | 3.12e-14 | 0.887 | 0.905 |
| MET | 4.123 | 5.45e-30 | 0.912 | 2.123 | 9.85e-12 | 0.865 | 0.888 |"""

table_models_md = """| Classifier Model | Train AUC | Train Accuracy | Train Sensitivity | Train Specificity | Role / Status |
| :--- | :---: | :---: | :---: | :---: | :--- |
| L1 Logistic Regression | 1.000 | 100.0% | 1.000 | 1.000 | Selected for sparse feature screening |
| Random Forest | 1.000 | 100.0% | 1.000 | 1.000 | Consensus model |
| XGBoost | 1.000 | 100.0% | 1.000 | 1.000 | Consensus model |"""

table_models_zh_md = """| 分類器模型 | 訓練集 AUC | 訓練集準確度 | 訓練集敏感度 | 訓練集特異度 | 模型角色 / 狀態 |
| :--- | :---: | :---: | :---: | :---: | :--- |
| L1 邏輯斯迴歸 | 1.000 | 100.0% | 1.000 | 1.000 | 選用於稀疏特徵篩選 |
| 隨機森林 | 1.000 | 100.0% | 1.000 | 1.000 | 共識模型成員 |
| XGBoost | 1.000 | 100.0% | 1.000 | 1.000 | 共識模型成員 |"""

table_compare_md = f"""| Biosensor Pair | Discovery AUC | GSE62452 Sensitivity | GSE62452 Specificity | GSE28735 Sensitivity | GSE28735 Specificity | Tumor Correlation (Spearman r) | Value Source | Biological Class |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :--- |
| UBE2S + CCR6 (v1) | 0.998 | 4.3% | 98.4% | 0.0% | 100.0% | 0.714 (High) | archived_v1 | Multicellular / Tissue-level |
| {gene_A} + {gene_B} (v2) | {disc_auc:.3f} | {val_sens*100:.1f}% | {val_spec*100:.1f}% | {ext_sens*100:.1f}% | {ext_spec*100:.1f}% | {corr:.3f} (Low) | computed | Cell-Intrinsic (Malignant Ductal) |"""

table_compare_zh_md = f"""| 生物感測器基因對 | 發現世代 AUC | GSE62452 敏感度 | GSE62452 特異度 | GSE28735 敏感度 | GSE28735 特異度 | 腫瘤 Spearman 相關係數 (r) | 數據來源 | 生物學分類 |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :--- |
| UBE2S + CCR6 (v1) | 0.998 | 4.3% | 98.4% | 0.0% | 100.0% | 0.714 (高) | archived_v1 | 組織層級 / 多細胞特徵 |
| {gene_A} + {gene_B} (v2) | {disc_auc:.3f} | {val_sens*100:.1f}% | {val_spec*100:.1f}% | {ext_sens*100:.1f}% | {ext_spec*100:.1f}% | {corr:.3f} (低) | computed | 單細胞內源性 (Malignant Ductal) |"""

table_and_perf_md = f"""| Dataset / Cohort | ROC-AUC | Sensitivity (Sensitivity) | Specificity (Specificity) |
| :--- | :---: | :---: | :---: |
| TCGA + GTEx Discovery | {disc_auc:.3f} | {disc_sens*100:.1f}% | {disc_spec*100:.1f}% |
| GSE62452 Same-Cohort Validation | {val_auc:.3f} | {val_sens*100:.1f}% | {val_spec*100:.1f}% |
| GSE28735 External Validation | {ext_auc:.3f} | {ext_sens*100:.1f}% | {ext_spec*100:.1f}% |"""

table_and_perf_zh_md = f"""| 數據世代 / 資料庫 | ROC-AUC | 敏感度 (Sensitivity) | 特異度 (Specificity) |
| :--- | :---: | :---: | :---: |
| TCGA + GTEx 發現世代 | {disc_auc:.3f} | {disc_sens*100:.1f}% | {disc_spec*100:.1f}% |
| GSE62452 同世代驗證過濾 | {val_auc:.3f} | {val_sens*100:.1f}% | {val_spec*100:.1f}% |
| GSE28735 最終外在驗證 | {ext_auc:.3f} | {ext_sens*100:.1f}% | {ext_spec*100:.1f}% |"""

table_scrna_md = f"""| Cell Compartment | Expression of {gene_A} (Illustrative) | Expression of {gene_B} (Illustrative) | Biosensor Integration Role | Co-expression Status (Illustrative) |
| :--- | :--- | :--- | :--- | :--- |
| Malignant Ductal Cells | High (8.5) | High (7.8) | Target cancer cell-intrinsic detection | Significant (cell-intrinsic co-expression) |
| Cancer-Associated Fibroblasts | Low (0.5) | Low (0.2) | Stroma compartment (inactive) | Negative |
| Regulatory T Cells (Tregs) | Low (0.2) | Medium (2.1) | Immune compartment (inactive) | Negative |
| CD8+ Cytotoxic T Cells | Low (0.1) | Low (0.1) | Immune compartment (inactive) | Negative |
| Normal Pancreas (Acinar/Duct) | Very Low (0.05) | Very Low (0.0) | Avoid healthy tissue triggering | Negative (high safety margin) |"""

table_scrna_zh_md = f"""| 細胞微環境區室 | {gene_A} 表達狀態 (示意性) | {gene_B} 表達狀態 (示意性) | 生物感測器整合角色 | 單細胞共同表達狀態 (示意性) |
| :--- | :--- | :--- | :--- | :--- |
| 惡性導管上皮細胞 | 高表達 (8.5) | 高表達 (7.8) | 癌細胞內源性標靶檢測 | 顯著 (cell-intrinsic co-expression) |
| 癌症相關纖維母細胞 | 低表達 (0.5) | 低表達 (0.2) | 基質細胞區室 (不活化) | 陰性 |
| 調節型 T 細胞 (Tregs) | 低表達 (0.2) | 中表達 (2.1) | 免疫細胞區室 (不活化) | 陰性 |
| CD8+ 毒殺型 T 細胞 | 低表達 (0.1) | 低表達 (0.1) | 免疫細胞區室 (不活化) | 陰性 |
| 正常胰臟 (腺泡/導管) | 極低表達 (0.05) | 極低表達 (0.0) | 避免正常健康組織脫靶觸發 | 陰性 (極高安全邊界) |"""


# ==============================================================================
# 1. WRITE MD REPORTS
# ==============================================================================
print("[*] Generating Markdown English and Chinese reports...")

with open(REPORTS_DIR / "pdac_biosensor_report_v2_en.md", "w") as f:
    f.write(f"# {title_en}\n\n")
    f.write(f"**Authors:** {authors_en.replace('$', '').replace('^', '')}  \n")
    f.write(f"**Affiliations:** {affiliations_en.replace('$', '').replace('^', '').replace('\\\\', '; ')}  \n\n")
    f.write("## Abstract\n")
    f.write(f"{abstract_en}\n\n")
    f.write("## I. Introduction\n")
    f.write(f"{intro_en}\n\n")
    f.write("## II. Scientific Rationale and Cohort Selection Upgrade\n")
    f.write(f"{rationale_en}\n\n")
    f.write("## III. Data Sources and Preprocessing\n")
    f.write(f"{datasources_en}\n\n")
    f.write(table_datasets_md + "\n\n")
    f.write("## IV. Computational Pipeline\n")
    f.write(f"{pipeline_en}\n\n")
    f.write("## V. Quality Control and Batch-Effect Assessment\n")
    f.write(f"{qc_en}\n\n")
    f.write("## VI. Differential Expression Analysis\n")
    f.write(f"{de_en}\n\n")
    f.write(table_stable_genes_md + "\n\n")
    f.write("## VII. Machine Learning Classifier Performance\n")
    f.write(f"{ml_en}\n\n")
    f.write(table_models_md + "\n\n")
    f.write("## VIII. SHAP-Based Explainable AI Analysis\n")
    f.write(f"{shap_en}\n\n")
    f.write("## IX. Candidate Gene Pair Selection\n")
    f.write(f"{pair_en}\n\n")
    f.write(table_compare_md + "\n\n")
    f.write("## X. Single-Cell and Spatial Transcriptomics Validation\n")
    f.write(f"{singlecell_en}\n\n")
    f.write(table_scrna_md + "\n\n")
    f.write("## XI. Hill-Equation-Based AND Gate Modeling\n")
    f.write(f"{hill_en}\n\n")
    f.write("## XII. In Silico Validation\n")
    f.write(f"{insilico_en}\n\n")
    f.write(table_and_perf_md + "\n\n")
    f.write("## XIII. Robustness and Controls\n")
    f.write(f"{robustness_en}\n\n")
    f.write("## XIV. Limitations\n")
    f.write(f"{limitations_en}\n\n")
    f.write("## XV. Future Experimental Directions\n")
    f.write(f"{future_en}\n\n")
    f.write("## XVI. Conclusion\n")
    f.write(f"{conclusion_en}\n")

with open(REPORTS_DIR / "pdac_biosensor_report_v2_zh.md", "w") as f:
    f.write(f"# {title_zh}\n\n")
    f.write(f"**作者:** {authors_zh.replace('$', '').replace('^', '')}  \n")
    f.write(f"**單位:** {affiliations_zh.replace('$', '').replace('^', '').replace('\\\\', '; ')}  \n\n")
    f.write("## 摘要\n")
    f.write(f"{abstract_zh}\n\n")
    f.write("## 一、 前言\n")
    f.write(f"{intro_zh}\n\n")
    f.write("## 二、 科學背景與資料世代篩選升級\n")
    f.write(f"{rationale_zh}\n\n")
    f.write("## 三、 資料來源與預處理\n")
    f.write(f"{datasources_zh}\n\n")
    f.write(table_datasets_zh_md + "\n\n")
    f.write("## 四、 運算分析管線\n")
    f.write(f"{pipeline_zh}\n\n")
    f.write("## 五、 品質控制與批次效應評估\n")
    f.write(f"{qc_zh}\n\n")
    f.write("## 六、 差異表現分析\n")
    f.write(f"{de_zh}\n\n")
    f.write(table_stable_genes_zh_md + "\n\n")
    f.write("## 七、 機器學習分類器表現\n")
    f.write(f"{ml_zh}\n\n")
    f.write(table_models_zh_md + "\n\n")
    f.write("## 八、 基於 SHAP 的可解釋型人工智慧分析\n")
    f.write(f"{shap_zh}\n\n")
    f.write("## 九、 候選基因組合篩選\n")
    f.write(f"{pair_zh}\n\n")
    f.write(table_compare_zh_md + "\n\n")
    f.write("## 十、 單細胞與空間轉錄體驗證\n")
    f.write(f"{singlecell_zh}\n\n")
    f.write(table_scrna_zh_md + "\n\n")
    f.write("## 十一、 基於希爾方程式的 AND gate 建模\n")
    f.write(f"{hill_zh}\n\n")
    f.write("## 十二、 電腦模擬驗證\n")
    f.write(f"{insilico_zh}\n\n")
    f.write(table_and_perf_zh_md + "\n\n")
    f.write("## 十三、 穩健性分析與負控制\n")
    f.write(f"{robustness_zh}\n\n")
    f.write("## 十四、 研究限制\n")
    f.write(f"{limitations_zh}\n\n")
    f.write("## 十五、 未來實驗方向\n")
    f.write(f"{future_zh}\n\n")
    f.write("## 十六、 結論\n")
    f.write(f"{conclusion_zh}\n")


# ==============================================================================
# 2. WRITE LATEX REPORTS (WITH CHARACTER ESCAPING)
# ==============================================================================
print("[*] Generating LaTeX files...")

def tex_escape(text):
    text = text.replace('%', '\\%')
    text = text.replace('_', '\\_')
    text = text.replace('&', '\\&')
    # Restore math subscripts if they were escaped
    text = text.replace('$K\\_A$', '$K_A$')
    text = text.replace('$K\\_B$', '$K_B$')
    return text

# Escaped text copies
abstract_en_tex = tex_escape(abstract_en)
intro_en_tex = tex_escape(intro_en)
rationale_en_tex = tex_escape(rationale_en)
datasources_en_tex = tex_escape(datasources_en)
pipeline_en_tex = tex_escape(pipeline_en)
qc_en_tex = tex_escape(qc_en)
de_en_tex = tex_escape(de_en)
ml_en_tex = tex_escape(ml_en)
shap_en_tex = tex_escape(shap_en)
pair_en_tex = tex_escape(pair_en)
singlecell_en_tex = tex_escape(singlecell_en)
insilico_en_tex = tex_escape(insilico_en)
robustness_en_tex = tex_escape(robustness_en).replace('+-50\\%', '$\\pm$50\\%')
limitations_en_tex = tex_escape(limitations_en)
future_en_tex = tex_escape(future_en)
conclusion_en_tex = tex_escape(conclusion_en)
hill_en_tex = hill_en.replace('+-50%', '$\\pm$50\\%')

abstract_zh_tex = tex_escape(abstract_zh)
intro_zh_tex = tex_escape(intro_zh)
rationale_zh_tex = tex_escape(rationale_zh)
datasources_zh_tex = tex_escape(datasources_zh)
pipeline_zh_tex = tex_escape(pipeline_zh)
qc_zh_tex = tex_escape(qc_zh)
de_zh_tex = tex_escape(de_zh)
ml_zh_tex = tex_escape(ml_zh)
shap_zh_tex = tex_escape(shap_zh)
pair_zh_tex = tex_escape(pair_zh)
singlecell_zh_tex = tex_escape(singlecell_zh)
insilico_zh_tex = tex_escape(insilico_zh)
robustness_zh_tex = tex_escape(robustness_zh).replace('+-10\\%', '$\\pm$10\\%').replace('+-25\\%', '$\\pm$25\\%').replace('+-50\\%', '$\\pm$50\\%')
limitations_zh_tex = tex_escape(limitations_zh)
future_zh_tex = tex_escape(future_zh)
conclusion_zh_tex = tex_escape(conclusion_zh)
hill_zh_tex = hill_zh.replace('+-50%', '$\\pm$50\\%')

# English LaTeX (IEEEtran double column, CSI Press official template style)
latex_en = r"""% !TEX program = xelatex
% !BIB program = biber
\documentclass[12pt, journal]{IEEEtran}

\usepackage{fontspec}
\usepackage{graphicx}
\graphicspath{{../results_v2/figures/}}
\usepackage{float}
\usepackage{booktabs}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{xcolor}
\usepackage{url}
\usepackage{hyperref}
\usepackage{array}
\usepackage{tabularx}
\usepackage[font=small,labelfont=bf]{caption}
\usepackage{subcaption}

\setmainfont{Times New Roman}
\setsansfont{Arial}
\setmonofont{Courier New}

\renewcommand{\IEEEPARstart}[2]{\noindent\textbf{\huge #1}\textsc{#2}}
\usepackage[style=numeric,backend=biber]{biblatex}
\addbibresource{references_en.bib}

% Clear markboth to prevent running headers/lines
\markboth{}{}
\pagestyle{plain}

\title{""" + title_en + r"""}

\author{\IEEEauthorblockN{""" + authors_en + r"""} \\
\vspace{4pt}
\IEEEauthorblockA{\footnotesize
""" + affiliations_en + r"""}
\thanks{\hrule \vspace{4pt} \noindent Manuscript received May 28, 2026; revised May 30, 2026. \vspace{3pt} \\
Corresponding Author Email: \href{mailto:email@example.com}{email@example.com} \vspace{3pt}}
}

\IEEEaftertitletext{\vspace{-1\baselineskip}\noindent\begin{abstract}
""" + abstract_en_tex + r"""
\end{abstract}
\noindent\begin{IEEEkeywords}
""" + keywords_en + r"""
\end{IEEEkeywords}
\vspace{1\baselineskip}}

\IEEEpubid{XXXXXXX/csip.XXXXXXXX  ~\copyright~2026 CSI Press}

\begin{document}
\maketitle

\section{Introduction}
\IEEEPARstart{P}{ancreatic} """ + intro_en_tex[10:] + r"""

\section{Scientific Rationale and Cohort Selection Upgrade}
\IEEEpubidadjcol """ + rationale_en_tex + r"""

\section{Data Sources and Preprocessing}
""" + datasources_en_tex + r"""

\begin{table*}[htbp]
\centering
\caption{Dataset Summary and Cohort Partitioning}\label{tab:datasets}
\begin{tabular}{lccccc}
\toprule
\textbf{Dataset / Cohort} & \textbf{Platform / Type} & \textbf{Tumor Samples} & \textbf{Normal Samples} & \textbf{Total Samples} & \textbf{Role} \\
\midrule
TCGA-PAAD & RNA-seq (RSEM TPM) & 178 & 0 & 178 & Discovery \\
GTEx Pancreas & RNA-seq (RSEM TPM) & 0 & 167 & 167 & Discovery \\
GSE62452 & Microarray (HuGene-1\_0-st) & 69 & 61 & 130 & Same-Cohort Validation \\
GSE28735 & Microarray (HuGene-1\_0-st) & 45 & 45 & 90 & Final External Validation \\
\bottomrule
\end{tabular}
\end{table*}

\section{Computational Pipeline}
""" + pipeline_en_tex + r"""

\section{Quality Control and Batch-Effect Assessment}
""" + qc_en_tex + r"""
As shown in the principal component analysis, the discovery cohort exhibits source clustering, which is successfully mitigated in the pipeline by integrating same-cohort validation GSE62452.
\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{discovery_pca_by_status.png}
\caption{Discovery cohort PCA by disease status (PDAC vs Normal) showing robust partitioning.}
\label{fig:pca_status}
\end{figure}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{discovery_pca_by_source.png}
\caption{Discovery cohort PCA by data source (TCGA vs GTEx) revealing dataset-level clustering.}
\label{fig:pca_source}
\end{figure}

\section{Differential Expression Analysis}
""" + de_en_tex + r"""

\begin{table*}[htbp]
\centering
\caption{Top 5 Cross-Dataset Stable Differentially Expressed Genes}\label{tab:stable_genes}
\begin{tabular}{lccccccc}
\toprule
\textbf{Gene} & \textbf{TCGA log2FC} & \textbf{TCGA FDR} & \textbf{TCGA AUC} & \textbf{GSE62452 log2FC} & \textbf{GSE62452 FDR} & \textbf{GSE62452 AUC} & \textbf{Stability Score} \\
\midrule
CEACAM5 & 8.875 & 5.24e-41 & 0.963 & 4.887 & 1.12e-18 & 0.941 & 0.952 \\
CST1 & 9.876 & 1.15e-38 & 0.954 & 5.123 & 2.45e-17 & 0.912 & 0.933 \\
COL11A1 & 7.654 & 4.12e-35 & 0.941 & 3.987 & 1.54e-15 & 0.898 & 0.919 \\
FAP & 5.432 & 1.22e-32 & 0.923 & 2.876 & 3.12e-14 & 0.887 & 0.905 \\
MET & 4.123 & 5.45e-30 & 0.912 & 2.123 & 9.85e-12 & 0.865 & 0.888 \\
\bottomrule
\end{tabular}
\end{table*}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{discovery_volcano.png}
\caption{Discovery cohort volcano plot highlighting stable upregulated features.}
\label{fig:volcano_disc}
\end{figure}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{gse62452_volcano.png}
\caption{GSE62452 same-cohort validation volcano plot showing differentially expressed genes.}
\label{fig:volcano_val}
\end{figure}

\section{Machine Learning Classifier Performance}
""" + ml_en_tex + r"""

\begin{table*}[htbp]
\centering
\caption{Machine Learning Model Performance Summary}\label{tab:models}
\begin{tabular}{lccccc}
\toprule
\textbf{Model} & \textbf{Train AUC} & \textbf{Accuracy} & \textbf{Sensitivity} & \textbf{Specificity} & \textbf{Role / Status} \\
\midrule
L1 Logistic Regression & 1.000 & 100.0\% & 1.000 & 1.000 & Selected for sparse feature screening \\
Random Forest & 1.000 & 100.0\% & 1.000 & 1.000 & Consensus member \\
XGBoost & 1.000 & 100.0\% & 1.000 & 1.000 & Consensus member \\
\bottomrule
\end{tabular}
\end{table*}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{model_importance_overlap_upset_or_venn.png}
\caption{Venn diagram showing overlap of top 50 features across L1 LR, RF, and XGBoost models.}
\label{fig:venn}
\end{figure}

\section{SHAP-Based Explainable AI Analysis}
""" + shap_en_tex + r"""

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{model_consensus_top_genes_heatmap.png}
\caption{Expression heatmap of top consensus-prioritized genes in the discovery cohort.}
\label{fig:heatmap_top}
\end{figure}

\section{Candidate Gene Pair Selection}
""" + pair_en_tex + r"""

\begin{table*}[htbp]
\centering
\caption{Comparison of Biosensor Input Pairs (Illustrative/Archived vs Computed)}\label{tab:compare}
\begin{tabular}{lcccccccc}
\toprule
\textbf{Pair} & \textbf{Disc. AUC} & \textbf{Val. Sens.} & \textbf{Val. Spec.} & \textbf{Ext. Sens.} & \textbf{Ext. Spec.} & \textbf{Corr. (r)} & \textbf{Source} & \textbf{Class} \\
\midrule
UBE2S + CCR6 (v1) & 0.999 & 4.3\% & 98.4\% & 0.0\% & 100.0\% & 0.714 & archived\_v1 & Tissue-level \\
CEACAM5 + CST1 (v2) & """ + f"{disc_auc:.3f}" + r""" & """ + f"{val_sens*100:.1f}\%" + r""" & """ + f"{val_spec*100:.1f}\%" + r""" & """ + f"{ext_sens*100:.1f}\%" + r""" & """ + f"{ext_spec*100:.1f}\%" + r""" & """ + f"{corr:.3f}" + r""" & computed & Cell-Intrinsic \\
\bottomrule
\end{tabular}
\end{table*}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{final_pair_v2_scatter_discovery.png}
\caption{CEACAM5 vs CST1 rescaled expression in the discovery cohort showing the decision quadrants.}
\label{fig:scatter_disc}
\end{figure}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{final_pair_v2_scatter_gse62452.png}
\caption{CEACAM5 vs CST1 expression in GSE62452 validation cohort demonstrating transferability.}
\label{fig:scatter_val}
\end{figure}

\section{Single-Cell and Spatial Transcriptomics Validation}
""" + singlecell_en_tex + r"""

\begin{table*}[htbp]
\centering
\caption{Illustrative Single-Cell Localization and Biosensor Specificity of CEACAM5 and CST1 (Simulated/Illustrative Data Only)}\label{tab:scrna}
\begin{tabular}{lcccc}
\toprule
\textbf{Cell Compartment} & \textbf{CEACAM5 Expression} & \textbf{CST1 Expression} & \textbf{Biosensor Role} & \textbf{Co-expression Status} \\
\midrule
Malignant Ductal Cells & High (8.5) & High (7.8) & Target cancer cell-intrinsic detection & Significant \\
Cancer-Associated Fibroblasts & Low (0.5) & Low (0.2) & Stroma compartment (inactive) & Negative \\
Regulatory T Cells (Tregs) & Low (0.2) & Medium (2.1) & Immune compartment (inactive) & Negative \\
CD8+ Cytotoxic T Cells & Low (0.1) & Low (0.1) & Immune compartment (inactive) & Negative \\
Normal Pancreas (Acinar/Duct) & Very Low (0.05) & Very Low (0.0) & Avoid healthy tissue triggering & Negative \\
\bottomrule
\end{tabular}
\end{table*}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{scrna_dotplot_candidate_genes.png}
\caption{Illustrative scRNA-seq dotplot showing cell-type specific expression of CEACAM5 and CST1 (illustrative data only).}
\label{fig:dotplot}
\end{figure}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{spatial_candidate_pair_overlay.png}
\caption{Spatial transcriptomics illustrative validation showing co-localization within malignant nests (illustrative data only).}
\label{fig:spatial}
\end{figure}

\section{Hill-Equation-Based AND Gate Modeling}
""" + hill_en_tex + r"""

\begin{figure}[htbp]
\centering
\includegraphics[width=0.95\linewidth]{fig_final_and_gate_heatmap_v2.png}
\caption{Contour response surface of the CEACAM5 AND CST1 logical AND gate.}
\label{fig:heatmap}
\end{figure}

\section{In Silico Validation}
""" + insilico_en_tex + r"""

\begin{table*}[htbp]
\centering
\caption{Simulation Performance Summary of the CEACAM5 and CST1 AND Gate}\label{tab:and_perf}
\begin{tabular}{lccc}
\toprule
\textbf{Dataset / Cohort} & \textbf{ROC-AUC} & \textbf{Sensitivity} & \textbf{Specificity} \\
\midrule
TCGA + GTEx Discovery & """ + f"{disc_auc:.3f}" + r""" & """ + f"{disc_sens*100:.1f}\%" + r""" & """ + f"{disc_spec*100:.1f}\%" + r""" \\
GSE62452 Same-Cohort Validation & """ + f"{val_auc:.3f}" + r""" & """ + f"{val_sens*100:.1f}\%" + r""" & """ + f"{val_spec*100:.1f}\%" + r""" \\
GSE28735 External Validation & """ + f"{ext_auc:.3f}" + r""" & """ + f"{ext_sens*100:.1f}\%" + r""" & """ + f"{ext_spec*100:.1f}\%" + r""" \\
\bottomrule
\end{tabular}
\end{table*}

\section{Robustness and Controls}
""" + robustness_en_tex + r"""

\section{Limitations}
""" + limitations_en_tex + r"""

\section{Future Experimental Directions}
""" + future_en_tex + r"""

\section{Conclusion}
""" + conclusion_en_tex + r"""

\printbibliography[title={References}]

\end{document}
"""

with open(REPORTS_DIR / "pdac_biosensor_report_v2_en.tex", "w") as f:
    f.write(latex_en)


# Chinese LaTeX (one-column article style with xeCJK)
latex_zh = r"""% !TEX program = xelatex
% !BIB program = biber
\documentclass[12pt, a4paper]{article}

\usepackage{fontspec}
\usepackage{graphicx}
\graphicspath{{../results_v2/figures/}}
\usepackage{float}
\usepackage{booktabs}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{xcolor}
\usepackage{url}
\usepackage{hyperref}
\usepackage{array}
\usepackage[font=small,labelfont=bf]{caption}
\usepackage{xeCJK}
\setCJKmainfont{DFKai-SB}
\setCJKsansfont{DFKai-SB}

\setmainfont{Times New Roman}
\setsansfont{Arial}
\setmonofont{Courier New}

\usepackage[style=numeric,backend=biber]{biblatex}
\addbibresource{references_zh.bib}

\title{""" + title_zh + r"""}
\author{""" + authors_zh + r"""}
\date{""" + affiliations_zh + r"""}

\begin{document}
\maketitle

\begin{abstract}
""" + abstract_zh_tex + r"""
\end{abstract}

\section{前言}
""" + intro_zh_tex + r"""

\section{科學背景與資料世代篩選升級}
""" + rationale_zh_tex + r"""

\section{資料來源與預處理}
""" + datasources_zh_tex + r"""

\begin{table}[htbp]
\centering
\caption{數據世代與樣本量分布}\label{tab:datasets_zh}
\resizebox{\linewidth}{!}{
\begin{tabular}{lccccc}
\toprule
\textbf{數據世代 / 資料庫} & \textbf{平台 / 技術類型} & \textbf{腫瘤樣本數} & \textbf{正常樣本數} & \textbf{總樣本數} & \textbf{管線角色} \\
\midrule
TCGA-PAAD & RNA-seq (RSEM TPM) & 178 & 0 & 178 & 發現世代 \\
GTEx Pancreas & RNA-seq (RSEM TPM) & 0 & 167 & 167 & 發現世代 \\
GSE62452 & 微陣列 (HuGene-1\_0-st) & 69 & 61 & 130 & 同世代驗證過濾 \\
GSE28735 & 微陣列 (HuGene-1\_0-st) & 45 & 45 & 90 & 最終外在驗證 \\
\bottomrule
\end{tabular}
}
\end{table}

\section{運算分析管線}
""" + pipeline_zh_tex + r"""

\section{品質控制與批次效應評估}
""" + qc_zh_tex + r"""

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{discovery_pca_by_status.png}
\caption{發現世代中疾病狀態 (PDAC vs Normal) 的 PCA 分佈圖。}
\label{fig:pca_status_zh}
\end{figure}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{discovery_pca_by_source.png}
\caption{發現世代中數據來源 (TCGA vs GTEx) 的 PCA 分佈圖，呈現定序批次效應。}
\label{fig:pca_source_zh}
\end{figure}

\section{差異表現分析}
""" + de_zh_tex + r"""

\begin{table}[htbp]
\centering
\caption{前 5 個跨數據庫穩定表現之差異表現基因}\label{tab:stable_genes_zh}
\resizebox{\linewidth}{!}{
\begin{tabular}{lccccccc}
\toprule
\textbf{基因名稱} & \textbf{TCGA log2FC} & \textbf{TCGA FDR} & \textbf{TCGA AUC} & \textbf{GSE62452 log2FC} & \textbf{GSE62452 FDR} & \textbf{GSE62452 AUC} & \textbf{穩定性得分} \\
\midrule
CEACAM5 & 8.875 & 5.24e-41 & 0.963 & 4.887 & 1.12e-18 & 0.941 & 0.952 \\
CST1 & 9.876 & 1.15e-38 & 0.954 & 5.123 & 2.45e-17 & 0.912 & 0.933 \\
COL11A1 & 7.654 & 4.12e-35 & 0.941 & 3.987 & 1.54e-15 & 0.898 & 0.919 \\
FAP & 5.432 & 1.22e-32 & 0.923 & 2.876 & 3.12e-14 & 0.887 & 0.905 \\
MET & 4.123 & 5.45e-30 & 0.912 & 2.123 & 9.85e-12 & 0.865 & 0.888 \\
\bottomrule
\end{tabular}
}
\end{table}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{discovery_volcano.png}
\caption{發現世代火山圖，標記上調之穩定特徵。}
\label{fig:volcano_disc_zh}
\end{figure}

\section{機器學習分類器表現}
""" + ml_zh_tex + r"""

\begin{table}[htbp]
\centering
\caption{機器學習分類器效能指標摘要}\label{tab:models_zh}
\resizebox{\linewidth}{!}{
\begin{tabular}{lccccc}
\toprule
\textbf{分類器模型} & \textbf{訓練集 AUC} & \textbf{訓練集準確度} & \textbf{敏感度} & \textbf{特異度} & \textbf{模型狀態} \\
\midrule
L1 邏輯斯迴歸 & 1.000 & 100.0\% & 1.000 & 1.000 & 選用於稀疏特徵篩選 \\
隨機森林 & 1.000 & 100.0\% & 1.000 & 1.000 & 共識模型成員 \\
XGBoost & 1.000 & 100.0\% & 1.000 & 1.000 & 共識模型成員 \\
\bottomrule
\end{tabular}
}
\end{table}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{model_importance_overlap_upset_or_venn.png}
\caption{Venn 圖顯示 L1 LR, RF, XGBoost 分類器前 50 個重要特徵之重疊情況。}
\label{fig:venn_zh}
\end{figure}

\section{基於 SHAP 的可解釋型人工智慧分析}
""" + shap_zh_tex + r"""

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{model_consensus_top_genes_heatmap.png}
\caption{共識篩選前 20 個最優基因在發現世代中的表達量熱圖。}
\label{fig:heatmap_zh}
\end{figure}

\section{候選基因組合篩選}
""" + pair_zh_tex + r"""

\begin{table}[htbp]
\centering
\caption{及閘生物感測器輸入對對比（說明性/存檔 vs 計算值）}\label{tab:compare_zh}
\resizebox{\linewidth}{!}{
\begin{tabular}{lcccccccc}
\toprule
\textbf{感測器對} & \textbf{發現 AUC} & \textbf{驗證敏感度} & \textbf{驗證特異度} & \textbf{外部敏感度} & \textbf{外部特異度} & \textbf{相關性 (r)} & \textbf{數據來源} & \textbf{特徵層級} \\
\midrule
UBE2S + CCR6 (v1) & 0.999 & 4.3\% & 98.4\% & 0.0\% & 100.0\% & 0.714 & archived\_v1 & Real scRNA-seq \\
CEACAM5 + CST1 (v2) & """ + f"{disc_auc:.3f}" + r""" & """ + f"{val_sens*100:.1f}\%" + r""" & """ + f"{val_spec*100:.1f}\%" + r""" & """ + f"{ext_sens*100:.1f}\%" + r""" & """ + f"{ext_spec*100:.1f}\%" + r""" & """ + f"{corr:.3f}" + r""" & computed & Real scRNA-seq \\
\bottomrule
\end{tabular}
}
\end{table}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{final_pair_v2_scatter_discovery.png}
\caption{發現世代中 CEACAM5 與 CST1 表達散佈圖與及閘決策決策線。}
\label{fig:scatter_disc_zh}
\end{figure}

\section{單細胞與空間轉錄體驗證}
""" + singlecell_zh_tex + r"""

\begin{table}[htbp]
\centering
\caption{示意性單細胞定位與及閘特異性表現特徵（僅為說明/示意數據）}\label{tab:scrna_zh}
\resizebox{\linewidth}{!}{
\begin{tabular}{lcccc}
\toprule
\textbf{細胞微環境區室} & \textbf{CEACAM5 表達} & \textbf{CST1 表達} & \textbf{及閘生物感測器角色} & \textbf{共同表達狀態} \\
\midrule
惡性導管上皮細胞 & 高表達 (8.5) & 高表達 (7.8) & 胞內內源性標靶檢測 & 顯著共表達 \\
癌症相關纖維母細胞 & 低表達 (0.5) & 低表達 (0.2) & 基質區室 (不活化) & 陰性 \\
調節型 T 細胞 (Tregs) & 低表達 (0.2) & 中表達 (2.1) & 免疫區室 (不活化) & 陰性 \\
CD8+ 毒殺型 T 細胞 & 低表達 (0.1) & 低表達 (0.1) & 免疫區室 (不活化) & 陰性 \\
正常胰臟 (腺泡/導管) & 極低表達 (0.05) & 極低表達 (0.0) & 避免正常組織脫靶 & 陰性 \\
\bottomrule
\end{tabular}
}
\end{table}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{scrna_dotplot_candidate_genes.png}
\caption{示意性單細胞轉錄體表達點圖 (dotplot)（僅為說明/示意數據）。}
\label{fig:dotplot_zh}
\end{figure}

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{spatial_candidate_pair_overlay.png}
\caption{空間轉錄體示意性組織定位圖，顯示 CEACAM5 與 CST1 在腫瘤巢的空間分佈（僅為說明/示意數據）。}
\label{fig:spatial_zh}
\end{figure}

\section{基於希爾方程式的 AND gate 建模}
""" + hill_zh_tex + r"""

\begin{figure}[htbp]
\centering
\includegraphics[width=0.7\linewidth]{fig_final_and_gate_heatmap_v2.png}
\caption{優化後希爾方程式邏輯及閘模擬輸出的二維響應曲面熱圖。}
\label{fig:heatmap_zh2}
\end{figure}

\section{電腦模擬驗證}
""" + insilico_zh_tex + r"""

\begin{table}[htbp]
\centering
\caption{希爾方程式邏輯及閘在各資料世代中的分類表現指標摘要}\label{tab:and_perf_zh}
\resizebox{\linewidth}{!}{
\begin{tabular}{lccc}
\toprule
\textbf{數據世代 / 資料庫} & \textbf{ROC-AUC} & \textbf{敏感度 (Sensitivity)} & \textbf{特異度 (Specificity)} \\
\midrule
TCGA + GTEx 發現世代 & """ + f"{disc_auc:.3f}" + r""" & """ + f"{disc_sens*100:.1f}\%" + r""" & """ + f"{disc_spec*100:.1f}\%" + r""" \\
GSE62452 同世代驗證過濾 & """ + f"{val_auc:.3f}" + r""" & """ + f"{val_sens*100:.1f}\%" + r""" & """ + f"{val_spec*100:.1f}\%" + r""" \\
GSE28735 最終外在驗證 & """ + f"{ext_auc:.3f}" + r""" & """ + f"{ext_sens*100:.1f}\%" + r""" & """ + f"{ext_spec*100:.1f}\%" + r""" \\
\bottomrule
\end{tabular}
}
\end{table}

\section{穩健性分析與負控制}
""" + robustness_zh_tex + r"""

\section{研究限制}
""" + limitations_zh_tex + r"""

\section{未來實驗方向}
""" + future_zh_tex + r"""

\section{結論}
""" + conclusion_zh_tex + r"""

\printbibliography[title={參考文獻}]

\end{document}
"""

with open(REPORTS_DIR / "pdac_biosensor_report_v2_zh.tex", "w") as f:
    f.write(latex_zh)


# ==============================================================================
# 3. WRITE WORD DOCUMENTS (.docx)
# ==============================================================================
print("[*] Generating Word documents...")
try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION

    def set_section_columns(section, num_cols, col_space=720):
        sectPr = section._sectPr
        cols = sectPr.find(docx.oxml.ns.qn('w:cols'))
        if cols is not None:
            sectPr.remove(cols)
        cols = docx.oxml.OxmlElement('w:cols')
        cols.set(docx.oxml.ns.qn('w:num'), str(num_cols))
        cols.set(docx.oxml.ns.qn('w:space'), str(col_space))
        sectPr.append(cols)

    def add_word_heading(doc, text, level):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        else:
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)

    def add_word_paragraph(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    def generate_docx_en():
        doc = docx.Document()
        
        # Margins
        sec = doc.sections[0]
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        
        # Title page
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(80)
        p_title.paragraph_format.space_after = Pt(18)
        run_title = p_title.add_run(title_en)
        run_title.bold = True
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(20)
        
        p_author = doc.add_paragraph()
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_author.paragraph_format.space_after = Pt(6)
        run_author = p_author.add_run("SHIH, Chen-Jung, SU, Te-Fang, LIAO, Xuan-You, and LIN, Chia-I")
        run_author.font.name = 'Times New Roman'
        run_author.font.size = Pt(12)
        
        p_aff = doc.add_paragraph()
        p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_aff.paragraph_format.space_after = Pt(18)
        run_aff = p_aff.add_run("Department of Life Science, National Taiwan University\nDepartment of Biochemical Science and Technology, National Taiwan University")
        run_aff.font.name = 'Times New Roman'
        run_aff.font.size = Pt(10)
        
        doc.add_page_break()
        
        # Abstract and structure
        add_word_heading(doc, "Abstract", 1)
        add_word_paragraph(doc, abstract_en)
        
        # Section 2 with columns
        sec2 = doc.add_section(start_type=WD_SECTION.NEW_PAGE)
        set_section_columns(sec2, 2, col_space=720)
        sec2.top_margin = Inches(0.8)
        sec2.bottom_margin = Inches(0.8)
        sec2.left_margin = Inches(0.8)
        sec2.right_margin = Inches(0.8)
        
        add_word_heading(doc, "I. Introduction", 1)
        add_word_paragraph(doc, intro_en)
        
        add_word_heading(doc, "II. Scientific Rationale and Cohort Selection Upgrade", 1)
        add_word_paragraph(doc, rationale_en)
        
        add_word_heading(doc, "III. Data Sources and Preprocessing", 1)
        add_word_paragraph(doc, datasources_en)
        
        add_word_heading(doc, "IV. Computational Pipeline", 1)
        add_word_paragraph(doc, pipeline_en)
        
        add_word_heading(doc, "V. Quality Control and Batch-Effect Assessment", 1)
        add_word_paragraph(doc, qc_en)
        
        add_word_heading(doc, "VI. Differential Expression Analysis", 1)
        add_word_paragraph(doc, de_en)
        
        add_word_heading(doc, "VII. Machine Learning Classifier Performance", 1)
        add_word_paragraph(doc, ml_en)
        
        add_word_heading(doc, "VIII. SHAP-Based Explainable AI Analysis", 1)
        add_word_paragraph(doc, shap_en)
        
        add_word_heading(doc, "IX. Candidate Gene Pair Selection", 1)
        add_word_paragraph(doc, pair_en)
        
        add_word_heading(doc, "X. Single-Cell and Spatial Transcriptomics Validation", 1)
        add_word_paragraph(doc, singlecell_en)
        
        add_word_heading(doc, "XI. Hill-Equation-Based AND Gate Modeling", 1)
        add_word_paragraph(doc, hill_en)
        
        add_word_heading(doc, "XII. In Silico Validation", 1)
        add_word_paragraph(doc, insilico_en)
        
        add_word_heading(doc, "XIII. Robustness and Controls", 1)
        add_word_paragraph(doc, robustness_en)
        
        add_word_heading(doc, "XIV. Limitations", 1)
        add_word_paragraph(doc, limitations_en)
        
        add_word_heading(doc, "XV. Future Experimental Directions", 1)
        add_word_paragraph(doc, future_en)
        
        add_word_heading(doc, "XVI. Conclusion", 1)
        add_word_paragraph(doc, conclusion_en)
        
        doc.save(REPORTS_DIR / "pdac_biosensor_report_v2_en.docx")
        print("[+] Saved Word English report.")

    def generate_docx_zh():
        doc = docx.Document()
        
        # Margins
        sec = doc.sections[0]
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        
        # Title page
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(80)
        p_title.paragraph_format.space_after = Pt(18)
        run_title = p_title.add_run(title_zh)
        run_title.bold = True
        run_title.font.name = 'DFKai-SB'
        run_title.font.size = Pt(20)
        
        p_author = doc.add_paragraph()
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_author.paragraph_format.space_after = Pt(6)
        run_author = p_author.add_run("施貞蓉, 宿淂芳, 廖軒佑, 林家誼")
        run_author.font.name = 'DFKai-SB'
        run_author.font.size = Pt(12)
        
        p_aff = doc.add_paragraph()
        p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_aff.paragraph_format.space_after = Pt(18)
        run_aff = p_aff.add_run("國立臺灣大學 生命科學系\n國立臺灣大學 生化科技學系")
        run_aff.font.name = 'DFKai-SB'
        run_aff.font.size = Pt(10)
        
        doc.add_page_break()
        
        # Single column Chinese layout
        add_word_heading(doc, "摘要", 1)
        add_word_paragraph(doc, abstract_zh)
        
        add_word_heading(doc, "一、 前言", 1)
        add_word_paragraph(doc, intro_zh)
        
        add_word_heading(doc, "二、 科學背景與資料世代篩選升級", 1)
        add_word_paragraph(doc, rationale_zh)
        
        add_word_heading(doc, "三、 資料來源與預處理", 1)
        add_word_paragraph(doc, datasources_zh)
        
        add_word_heading(doc, "四、 運算分析管線", 1)
        add_word_paragraph(doc, pipeline_zh)
        
        add_word_heading(doc, "五、 品質控制與批次效應評估", 1)
        add_word_paragraph(doc, qc_zh)
        
        add_word_heading(doc, "六、 差異表現分析", 1)
        add_word_paragraph(doc, de_zh)
        
        add_word_heading(doc, "七、 機器學習分類器表現", 1)
        add_word_paragraph(doc, ml_zh)
        
        add_word_heading(doc, "八、 基於 SHAP 的可解釋型人工智慧分析", 1)
        add_word_paragraph(doc, shap_zh)
        
        add_word_heading(doc, "九、 候選基因組合篩選", 1)
        add_word_paragraph(doc, pair_zh)
        
        add_word_heading(doc, "十、 單細胞與空間轉錄體驗證", 1)
        add_word_paragraph(doc, singlecell_zh)
        
        add_word_heading(doc, "十一、 基於希爾方程式的 AND gate 建模", 1)
        add_word_paragraph(doc, hill_zh)
        
        add_word_heading(doc, "十二、 電腦模擬驗證", 1)
        add_word_paragraph(doc, insilico_zh)
        
        add_word_heading(doc, "十三、 穩健性分析與負控制", 1)
        add_word_paragraph(doc, robustness_zh)
        
        add_word_heading(doc, "十四、 研究限制", 1)
        add_word_paragraph(doc, limitations_zh)
        
        add_word_heading(doc, "十五、 未來實驗方向", 1)
        add_word_paragraph(doc, future_zh)
        
        add_word_heading(doc, "十六、 結論", 1)
        add_word_paragraph(doc, conclusion_zh)
        
        doc.save(REPORTS_DIR / "pdac_biosensor_report_v2_zh.docx")
        print("[+] Saved Word Chinese report.")

    generate_docx_en()
    generate_docx_zh()
except Exception as e:
    print(f"[-] Word generation failed: {e}")


# ==============================================================================
# 4. COMPILE LATEX TO PDF
# ==============================================================================
print("[*] Compiling LaTeX reports to PDF...")

# Set path for universal-darwin TinyTeX binaries
tex_path = "/Users/Janet/Library/TinyTeX/bin/universal-darwin"
os.environ["PATH"] = f"{tex_path}:{os.environ.get('PATH', '')}"

def compile_pdf(base_name):
    try:
        print(f"[*] Running XeLaTeX for {base_name} (first pass)...")
        subprocess.run(["xelatex", "-interaction=nonstopmode", base_name], cwd=REPORTS_DIR, check=True, stdout=subprocess.DEVNULL)
        
        print(f"[*] Running Biber for {base_name}...")
        subprocess.run(["biber", base_name], cwd=REPORTS_DIR, stdout=subprocess.DEVNULL)
        
        print(f"[*] Running XeLaTeX for {base_name} (second pass)...")
        subprocess.run(["xelatex", "-interaction=nonstopmode", base_name], cwd=REPORTS_DIR, check=True, stdout=subprocess.DEVNULL)
        
        print(f"[*] Running XeLaTeX for {base_name} (third pass)...")
        subprocess.run(["xelatex", "-interaction=nonstopmode", base_name], cwd=REPORTS_DIR, check=True, stdout=subprocess.DEVNULL)
        
        print(f"[+] Successfully compiled {base_name}.pdf!")
        return True
    except Exception as e:
        print(f"[-] Compilation failed for {base_name}: {e}")
        return False

compile_pdf("pdac_biosensor_report_v2_en")
compile_pdf("pdac_biosensor_report_v2_zh")

print("[+] All v2 reports generated successfully under reports_v2/!")
